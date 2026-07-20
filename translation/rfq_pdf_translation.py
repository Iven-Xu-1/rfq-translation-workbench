from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import urllib.error
import urllib.parse
import urllib.request
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path, PurePosixPath
from threading import Lock, local
from typing import Iterable

import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

from pdf_runtime.ocr import OCR_CONTRACT_VERSION
from pdf_runtime.output_naming import (
    OUTPUT_NAMING_CONTRACT_VERSION,
    OutputNamingPlan,
    plan_translated_outputs,
    translated_output_extension,
)
from pdf_runtime.preflight import (
    DEFAULT_THRESHOLDS as PDF_PREFLIGHT_THRESHOLDS,
    PDF_FALLBACK_CONTRACT_VERSION,
    PDF_PAGE_RANGE_CONTRACT_VERSION,
    PDF_PREFLIGHT_CONTRACT_VERSION,
    configured_preflight_thresholds,
    inspect_pdf_preflight,
    runtime_component_versions,
)

try:
    import requests
except ModuleNotFoundError:
    requests = None

try:
    from deep_translator import GoogleTranslator
except ModuleNotFoundError:
    GoogleTranslator = None


FONT_NAME = "STSong-Light"
pdfmetrics.registerFont(UnicodeCIDFont(FONT_NAME))
TRANSLATED_TEXT_DARKEN_OFFSETS = (0.0, 0.10, 0.18)
DIVIDER_WIDTH = 54.0
PROJECT_SOURCE_DIRNAME = "01_原始询价文件"
PROJECT_TRANSLATED_DIRNAME = "02_中文翻译文件"
PROJECT_SYSTEM_DIRNAME = "系统数据"
SUPPORTED_PROJECT_MODES = {"平衡": "平衡", "balanced": "平衡"}
PROJECT_TRANSLATION_CONFIG_SIGNATURE = "b-stage13-output-naming-office-isolation-v1"
B_TRANSLATION_COMPONENT_VERSION = "13.0.0"
PDF_PROTECTION_CONTRACT_VERSION = "pdf-protected-token-v2-translatable-compound"
PDF_OCR_DPI = 180
PDF_OCR_MIN_PAGE_CHARS = 8
PDF_OCR_MIN_AVERAGE_CONFIDENCE = 0.70
PDF_OCR_LOW_CONFIDENCE_THRESHOLD = 0.70
SELECTED_UPLOAD_MANIFEST_NAME = "selected_upload_files_manifest.json"
PDF_ENGINE_LEGACY = "legacy"
PDF_ENGINE_PDFMATHTRANSLATE_NEXT = "pdfmathtranslate_next"
SUPPORTED_PDF_ENGINES = {
    PDF_ENGINE_LEGACY,
    "old",
    "b_legacy",
    PDF_ENGINE_PDFMATHTRANSLATE_NEXT,
    "pdfmathtranslate-next",
    "pdf2zh_next",
}
PDF_RUNTIME_SOURCE_DIR = Path(__file__).resolve().parent / "pdf_runtime"
PDF_RUNTIME_WRAPPER = PDF_RUNTIME_SOURCE_DIR / "wrapper.py"
PDF_RUNTIME_DEPLOY_DIR = Path(__file__).resolve().parent / "deploy"
PDF_PREFLIGHT_STATE = local()
OUTPUT_NAMING_STATE = local()
DEFAULT_PDF_TRANSLATION_SERVICE = "openaicompatbatch"
DEFAULT_VECTOR_ENGINE_BASE_URL = "https://api.vectorengine.ai/v1"
DEFAULT_VECTOR_ENGINE_MODEL = "gemini-2.5-flash-lite"
DEFAULT_VECTOR_ENGINE_REPAIR_MODEL = "gemini-2.5-flash"
VECTOR_ENGINE_API_KEY_ENV = "VECTOR_ENGINE_API_KEY"
OFFICE_TRANSLATION_TARGET_LANGUAGE = "zh-CN"
OFFICE_PROMPT_CONTRACT_VERSION = "rfq-office-mechanical-v3-private-glossary"
OFFICE_PROTECTION_CONTRACT_VERSION = "protect-technical-terms-v1"
OFFICE_BATCH_FAILURE_ISOLATION_CONTRACT_VERSION = (
    "office-batch-recursive-isolation-bounded-single-retry-v1"
)
OFFICE_CACHE_NAMESPACES_KEY = "__office_translation_namespaces__"
PRIVATE_GLOSSARY_CACHE_KEY = "__private_glossary_terms__"
PRIVATE_GLOSSARY_ENV = "B_PDF_TRANSLATION_PRIVATE_GLOSSARIES"
SUPPORTED_TRANSLATION_SUFFIXES = {
    ".pdf",
    ".docx",
    ".doc",
    ".xlsx",
    ".xlsm",
    ".xls",
}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
XLSX_SUFFIXES = {".xlsx", ".xlsm"}
LEGACY_DOC_SUFFIXES = {".doc"}
LEGACY_XLS_SUFFIXES = {".xls"}
LAYOUT_STRATEGY_TEMPLATE = "template_overlay"
LAYOUT_STRATEGY_TABLE_REPAINT = "table_repaint"
UNDERLINED_SECTION_HEADINGS = {"SCOPE", "PURPOSE", "DEFINITIONS"}
DATASHEET_FORM_SPLIT_MARKERS = (
    "APPLICABLE TO",
    "SERVICE",
    "MODEL",
    "SIZE AND TYPE",
    "SITE",
    "NO. OF PUMPS REQUIRED",
    "MOTOR PROVIDED BY",
    "DRIVER PROVIDED BY",
    "GEAR PROVIDED BY",
    "MOTOR MOUNTED BY",
    "DRIVER MOUNTED BY",
    "GEAR MOUNTED BY",
    "MOTOR DATA SHEET NO.",
    "DRIVER DATA SHEET NO.",
    "GEAR DATA SHEET NO.",
)

EXACT_TRANSLATIONS = {
    "PUMP REQUIREMENTS": "泵要求",
    "TAG NO.": "位号",
    "TAG NO": "位号",
    "APPLICABLE STANDARD": "适用标准",
    "API 682 SEAL SYSTEM": "API 682 密封系统",
    "DETAIL DESIGN": "详细设计",
    "DOCUMENT TYPE": "文件类型",
    "DOCUMENT TITLE": "文件标题",
    "DOCUMENT NUMBER": "文件编号",
    "SITE": "现场",
    "TABULATION OF REVISED PAGES": "修订页清单",
    "MECHANICAL DATA SHEET FOR HP": "高压机械数据表",
    "DATA SHEET": "数据表",
    "PROCESS DATA SHEET /SPECIFICATION": "工艺数据表/技术规格",
    "MATERIAL REQUISITION FOR QUOTATION": "询价材料请购文件",
    "APPLICABLE TO": "适用于",
    "PROPOSAL": "报价阶段",
    "PURCHASE": "采购阶段",
    "AS BUILT": "竣工版",
    "NO. OF PUMPS REQUIRED": "所需泵数量",
    "SIZE AND TYPE": "规格和型式",
    "MODEL": "型号",
    "SERIAL NO.": "序列号",
    "NO. MOTOR DRIVEN": "电机驱动数量",
    "PUMP ITEM NO'S": "泵位号",
    "OTHER DRIVER TYPE": "其他驱动机类型",
    "DRIVER ITEM NO'S": "驱动机位号",
    "GEAR ITEM NO'S": "齿轮箱位号",
    "PUMP PROVIDED BY": "泵供货方",
    "MOTOR PROVIDED BY": "电机供货方",
    "DRIVER PROVIDED BY": "驱动机供货方",
    "GEAR PROVIDED BY": "齿轮箱供货方",
    "MOTOR MOUNTED BY": "电机安装方",
    "DRIVER MOUNTED BY": "驱动机安装方",
    "GEAR MOUNTED BY": "齿轮箱安装方",
    "MOTOR DATA SHEET NO.": "电机数据表编号",
    "DRIVER DATA SHEET NO.": "驱动机数据表编号",
    "GEAR DATA SHEET NO.": "齿轮箱数据表编号",
    "OPERATING CONDITIONS": "运行工况",
    "TYPE OR NAME OF LIQUID": "液体类型或名称",
    "COLD CONDENSATE": "冷凝液",
    "PUMPING TEMPERATURE": "泵送温度",
    "SPECIFIC GRAVITY": "比重",
    "SPECIFIC HEAT": "比热",
    "CORROSIVE/EROSIVE AGENTS": "腐蚀/冲蚀介质",
    "CHLORIDE CONCENTRATION": "氯离子浓度",
    "H2S CONCENTRATION": "H2S 浓度",
    "PERFORMANCE": "性能",
    "NUMBER OF FEEDS": "进料点数量",
    "RATED CAPACITY": "额定流量",
    "NPSH REQUIRED": "必需汽蚀余量",
    "KW RATED": "额定功率",
    "AT RELIEF SETTING": "在安全阀整定压力下",
    "PLUNGER SPEED": "柱塞速度",
    "STROKES/MIN": "次/分",
    "LENGTH OF STROKE": "行程长度",
    "PUMP HEAD": "泵头",
    "MAXIMUM PRESSURE": "最大压力",
    "HYDRO TEST PRESSURE": "水压试验压力",
    "MAX DISCH PRESS W/JOB DRIVER": "配套驱动机最大出口压力",
    "MAX KW BASIS GEAR STRENGTH": "按齿轮强度计最大功率",
    "CONNECTIONS": "接口",
    "LIQUID END": "液力端",
    "JACKET REQ'D": "需夹套",
    "DIAPHRAGM": "隔膜",
    "VALVES/FEED": "每个进料点阀数量",
    "INSTRUMENT AIR": "仪表空气",
    "ELECTRICAL AREA CLASS": "电气区域防爆等级",
    "TROPICALIZATION REQ'D": "需热带化",
    "RANGE OF AMBIENT TEMPS": "环境温度范围",
    "UNUSUAL CONDITIONS": "特殊环境条件",
    "SALT ATMOSPHERE": "盐雾环境",
    "RELATIVE HUMIDITY (NORM/MAX)": "相对湿度（正常/最大）",
    "UTILITY CONDITIONS": "公用工程条件",
    "ELECTRICITY": "电源",
    "DRIVERS": "驱动机",
    "HEATING": "加热",
    "SHUTDOWN": "停机",
    "HERTZ": "频率",
    "PHASE": "相数",
    "COOLING WATER": "冷却水",
    "PRESS.": "压力",
    "SOURCE": "来源",
    "APPLICABLE SPECIFICATIONS": "适用规范",
    "POSITIVE DISPLACEMENT PUMPS": "容积式泵",
    "CONTROLLED VOLUME": "控制容积式",
    "GOVERNING SPECIFICATION": "执行规范",
    "ACCORDING TO THE FOLLOWING SCHEDULE": "按以下时间安排",
    "ACCORDING TO THE FOLLOWING SCHEDULE:": "按以下时间安排：",
    "PUMP FRAME": "泵架",
    "ONTO THE PUMP FRAME": "安装在泵架上",
    "FRAME": "机架",
    "BE FURNISHED BY VENDOR": "由卖方提供",
    "BE FURNISHED BY VENDOR.": "由卖方提供。",
}

TERM_TRANSLATIONS = {
    "PUMP REQUIREMENTS": "泵要求",
    "TAG NO.": "位号",
    "TAG NO": "位号",
    "APPLICABLE STANDARD": "适用标准",
    "API 682 SEAL SYSTEM": "API 682 密封系统",
    "SEAL SYSTEM": "密封系统",
    "MATERIAL REQUISITION FOR QUOTATION": "询价材料请购文件",
    "PROCESS DATA SHEET": "工艺数据表",
    "MECHANICAL DATA SHEET": "机械数据表",
    "RECIPROCATING PUMPS": "往复泵",
    "RECIPROCATING PUMP": "往复泵",
    "METERING PUMPS": "计量泵",
    "METERING PUMP": "计量泵",
    "WATER MAKE-UP PUMPS": "补水泵",
    "WATER MAKE-UP PUMP": "补水泵",
    "CENTRIFUGAL PUMP": "离心泵",
    "AMINE CIRCULATION PUMPS": "胺液循环泵",
    "HORIZONTAL SURFACE PUMP": "卧式地面泵",
    "HIGH PRESSURE": "高压",
    "SITE DATA": "现场及公用工程数据",
    "DETAIL DESIGN": "详细设计",
    "DOCUMENT TYPE": "文件类型",
    "DOCUMENT TITLE": "文件标题",
    "DOCUMENT NUMBER": "文件编号",
    "CONTRACT NO.": "合同号",
    "CONTRACT NO": "合同号",
    "DOC. NO.": "文件号",
    "DOC. NO": "文件号",
    "REV.": "修订版",
    "CLASS": "等级",
    "SITE": "现场",
    "PAGE": "页",
    "PAGES": "页",
    "REVISION": "修订",
    "APPROVE": "批准",
    "APPROVED": "已批准",
    "ISSUED FOR": "发布用于",
    "DESCRIPTION": "说明",
    "PREPARED": "编制",
    "CHECKED": "校核",
    "APPROVED BY": "批准人",
    "TECHNICAL SPECIFICATION": "技术规格书",
    "SPECIFICATION": "技术规格",
    "DATA SHEET": "数据表",
    "PROCESS": "工艺",
    "MECHANICAL": "机械",
    "PUMP": "泵",
    "PUMPS": "泵",
    "UNIT": "单元",
    "TRAIN": "线",
    "GAS TREATING": "天然气处理",
    "FACILITIES": "设施",
    "INFORMATION": "信息",
    "CONTAINED HEREIN": "本文件所含",
    "PROPERTY": "财产",
    "PURCHASER": "买方",
    "VENDOR": "卖方",
    "MANUFACTURER": "制造商",
    "SUPPLIER": "供应商",
    "SCOPE OF SUPPLY": "供货范围",
    "SCOPE": "范围",
    "PROPOSAL": "报价阶段",
    "PURCHASE": "采购阶段",
    "AS BUILT": "竣工版",
    "PURPOSE": "目的",
    "DEFINITIONS": "定义",
    "GENERAL NOTES": "通用说明",
    "NOTES": "说明",
    "NOTE": "说明",
    "REMARKS": "备注",
    "OPERATING": "操作",
    "DESIGN": "设计",
    "NORMAL": "正常",
    "MIN": "最小",
    "MAX": "最大",
    "MINIMUM": "最小",
    "MAXIMUM": "最大",
    "PRESSURE": "压力",
    "TEMPERATURE": "温度",
    "FLOW RATE": "流量",
    "CAPACITY": "流量",
    "DIFFERENTIAL PRESSURE": "差压",
    "SUCTION PRESSURE": "入口压力",
    "DISCHARGE PRESSURE": "出口压力",
    "LIQUID": "液体",
    "FLUID": "介质",
    "SERVICE": "工况/用途",
    "DENSITY": "密度",
    "VISCOSITY": "黏度",
    "CORROSION ALLOWANCE": "腐蚀裕量",
    "MATERIAL": "材料",
    "CASING": "泵壳",
    "PLUNGER": "柱塞",
    "PISTON": "活塞",
    "SEAL": "密封",
    "GASKET": "垫片",
    "MOTOR": "电机",
    "MOTOR ITEM NO'S": "电机位号",
    "POWER": "功率",
    "VOLTAGE": "电压",
    "FREQUENCY": "频率",
    "DRIVER": "驱动机",
    "GEAR BOX": "齿轮箱",
    "GEAR": "齿轮箱",
    "INSTRUMENT": "仪表",
    "ELECTRICITY": "电源",
    "HEATING": "加热",
    "CONTROL": "控制",
    "SHUTDOWN": "停机",
    "COOLING WATER": "冷却水",
    "PRESS.": "压力",
    "CONTROL VOLUME RECIPROCATING PUMP": "控制容积式往复泵",
    "MOTOR DRIVEN": "电机驱动",
    "RELIEF SETTING": "安全阀整定压力",
    "DISCH PRESS": "出口压力",
    "STANDARD": "标准",
    "INSPECTION": "检验",
    "TEST": "试验",
    "PAINTING": "涂漆",
    "PACKING": "包装",
    "DELIVERY": "交货",
    "QUOTATION": "报价",
    "REQUIREMENT": "要求",
    "REQUIRED": "要求",
    "YES": "是",
    "NO": "否",
    "BY VENDOR": "由卖方提供",
    "BY PURCHASER": "由买方提供",
    "NOT APPLICABLE": "不适用",
    "N/A": "不适用",
}

MECHANICAL_TERM_TRANSLATIONS = {
    "SITE AND UTILITY DATA": "现场及公用工程数据",
    "SITE DATA": "现场及公用工程数据",
    "FOR": "供",
    "HP WATER MAKE-UP PUMPS": "高压补水泵",
    "APPLICABLE TO: PROPOSAL PURCHASE AS BUILT": "适用于：报价阶段 采购阶段 竣工版",
    "MOTOR PROVIDED BY PUMP MANUFACTURER": "电机供货方 泵制造商",
    "MOTOR MOUNTED BY PUMP MANUFACTURER": "电机安装方 泵制造商",
    "PUMP MANUFACTURER": "泵制造商",
    "APPROVED FOR CONSTRUCTION": "批准用于施工",
    "ISSUED FOR APPROVAL": "发布用于批准",
    "APPLICABLE TO": "适用于",
    "NO. OF PUMPS REQUIRED": "所需泵数量",
    "SIZE AND TYPE": "规格和型式",
    "SERIAL NO.": "序列号",
    "NO. MOTOR DRIVEN": "电机驱动数量",
    "PUMP ITEM NO'S": "泵位号",
    "OTHER DRIVER TYPE": "其他驱动机类型",
    "DRIVER ITEM NO'S": "驱动机位号",
    "GEAR ITEM NO'S": "齿轮箱位号",
    "PUMP PROVIDED BY": "泵供货方",
    "MOTOR PROVIDED BY": "电机供货方",
    "DRIVER PROVIDED BY": "驱动机供货方",
    "GEAR PROVIDED BY": "齿轮箱供货方",
    "MOTOR MOUNTED BY": "电机安装方",
    "DRIVER MOUNTED BY": "驱动机安装方",
    "GEAR MOUNTED BY": "齿轮箱安装方",
    "MOTOR DATA SHEET NO.": "电机数据表编号",
    "DRIVER DATA SHEET NO.": "驱动机数据表编号",
    "GEAR DATA SHEET NO.": "齿轮箱数据表编号",
    "OPERATING CONDITIONS": "运行工况",
    "NORMAL CAPACITY": "正常流量",
    "TYPE OR NAME OF LIQUID": "液体类型或名称",
    "PUMPING TEMPERATURE": "泵送温度",
    "SPECIFIC GRAVITY": "比重",
    "SPECIFIC HEAT": "比热",
    "CORROSIVE/EROSIVE AGENTS": "腐蚀/冲蚀介质",
    "CHLORIDE CONCENTRATION": "氯离子浓度",
    "H2S CONCENTRATION": "H2S 浓度",
    "NUMBER OF FEEDS": "进料点数量",
    "RATED CAPACITY": "额定流量",
    "DIFFERENTIAL HEAD": "差压头",
    "NPSH REQUIRED": "必需汽蚀余量",
    "NPSH AVAILABLE": "可用汽蚀余量",
    "VOLUMETRIC EFFICIENCY": "容积效率",
    "MECHANICAL EFFICIENCY": "机械效率",
    "KW RATED": "额定功率",
    "AT RELIEF SETTING": "在安全阀整定压力下",
    "VARIABLE SPEED": "可变转速",
    "PLUNGER SPEED (STROKES/MIN)": "柱塞速度（次/分）",
    "PLUNGER SPEED": "柱塞速度",
    "LENGTH OF STROKE": "行程长度",
    "PUMP HEAD": "泵头",
    "PUMP CASING": "泵壳",
    "PACKING GLAND": "填料函",
    "PULSATION DAMPENER": "脉动阻尼器",
    "HYDRO TEST PRESSURE": "水压试验压力",
    "MAX DISCH PRESS W/JOB DRIVER": "配套驱动机最大出口压力",
    "MAX DISCH PRESS. W/ JOB DRIVER": "配套驱动机最大出口压力",
    "MAX KW BASIS GEAR STRENGTH": "按齿轮强度计最大功率",
    "LIQUID END": "液力端",
    "JACKET REQ'D": "需夹套",
    "VALVES/FEED": "每个进料点阀数量",
    "INSTRUMENT AIR": "仪表空气",
    "ELECTRICAL AREA CLASS": "电气区域防爆等级",
    "HAZARDOUS AREA": "危险区域",
    "WINTERIZATION": "冬季防护",
    "TROPICALIZATION REQ'D": "需热带化",
    "TROPICALIZATION": "湿热带防护",
    "RANGE OF AMBIENT TEMPS": "环境温度范围",
    "UNUSUAL CONDITIONS": "特殊环境条件",
    "SALT ATMOSPHERE": "盐雾环境",
    "RELATIVE HUMIDITY (NORM/MAX)": "相对湿度（正常/最大）",
    "UTILITY CONDITIONS": "公用工程条件",
    "POSITIVE DISPLACEMENT PUMPS": "容积式泵",
    "CONTROLLED VOLUME": "控制容积式",
    "GOVERNING SPECIFICATION": "执行规范",
    "INDOOR": "室内",
    "OUTDOOR": "户外",
    "UNDER ROOF": "棚下",
    "UNHEATED": "不采暖",
    "HEATED": "采暖",
    "DUST": "粉尘",
    "FUMES": "烟气",
    "SALT ATMOSPHERE": "盐雾环境",
    "NORMAL": "正常",
    "MINIMUM": "最小",
    "MAXIMUM": "最大",
    "RATED": "额定",
    "ACTUAL": "实际",
    "TOXIC": "有毒",
    "FLAMMABLE": "易燃",
    "OTHER": "其他",
    "CLIENT": "客户",
    "LOCATION": "地点",
    "JOB NO.": "工作号",
    "PREPARED BY": "编制",
    "ITEM NUMBER": "位号",
    "CURVE NO.": "曲线号",
    "FLUID TYPE": "介质类型",
    "SUCTION FROM": "入口来源",
    "DISCHARGE TO": "出口去向",
    "NORM / MIN / MAX TEMP": "正常/最小/最大温度",
    "VAPOR PRESSURE @ P&T": "蒸汽压 @ P&T",
    "LUBE QUALITY": "润滑油品质",
    "CAPACITY (FOR EACH PUMP )": "流量（每台泵）",
    "DISCH PRESSURE": "出口压力",
    "PRESSURE INCREASE": "升压",
    "HYDRAULIC POWER": "液压功率",
    "VERTICAL PUMP SUMP DEPTH": "立式泵坑深度",
    "VERTICAL PUMP SUBMERGENCE": "立式泵淹没深度",
    "SUCTION NOZZLE SIZE & RATING": "入口管口尺寸及等级",
    "DISCHARGE NOZZLE SIZE & RATING": "出口管口尺寸及等级",
    "PUMP TYPE": "泵型",
    "FURNISHED BY": "供货方",
    "MOUNTED BY": "安装方",
    "SERVICE FACTOR": "服务系数",
    "KW RATING": "功率等级",
    "AREA CLASSIFICATION": "区域防爆等级",
    "VOLT/PHASE/CYCLE": "电压/相数/频率",
    "LIFTING LUGS": "吊耳",
    "SLING & SPREADER BAR": "吊索和撑杆",
    "EXPORT CRATING": "出口包装箱",
    "SUPPLY TEMP": "供给温度",
    "SUPPLY PRES": "供给压力",
    "SIGHT FLOW INDICATOR": "视流指示器",
    "PIPE MATERIAL": "管道材料",
    "COOLING WATER PLAN": "冷却水方案",
    "SEAL FLUSH PLAN": "密封冲洗方案",
    "AUXILIARY SEAL PLAN": "辅助密封方案",
    "TESTING (NON-HOLD)": "试验（非停检点）",
    "FULL LOAD": "满负荷",
    "HYDROTEST": "水压试验",
    "MECH RUN-IN": "机械试运行",
    "FINAL INSPECTION": "最终检验",
    "WITNESSED": "见证",
    "UNWITNESSED": "非见证",
    "TURNDOWN RATIO": "调节比",
    "WITHOUT ACCELERATION HEAD": "不含加速压头",
    "DRIVE": "驱动",
    "GENERAL": "总则",
    "SUCTION": "入口",
    "DISCHARGE": "出口",
    "FLUSH": "冲洗",
    "RATING": "等级",
    "DIAMETER": "直径",
    "SIZE": "尺寸",
    "FACING": "密封面",
    "POSITION": "位置",
    "LOCATION": "地点",
    "INLET": "入口",
    "RETURN": "回水",
    "TEMP": "温度",
    "CONSTRUCTION (NOTE 21)": "结构 (Note 21)",
    "ELECTRICITY DRIVERS HEATING CONTROL SHUTDOWN": "电源 驱动机 加热 控制 停机",
    "COOLING WATER INLET RETURN DESIGN MAX \u0394": "冷却水 入口 回水 设计 最大 \u0394",
    "MAX DISCH PRESS. W/ JOB DRIVER (BARG)": "配套驱动机最大出口压力 (BARG)",
    "COLD CONDENSATE (NOTE 12)": "冷凝液 (NOTE 12)",
    "CONTOUR PLATE": "轮廓板",
    "HYDRAULIC DIAPHRAGM": "液压隔膜",
    "PROCESS DIAPHRAGM": "工艺隔膜",
    "LANTERN RING": "灯笼环",
    "VALVE SEAT": "阀座",
    "VALVE GUIDE": "阀导向",
    "VALVE BODY": "阀体",
    "VALVE GASKET": "阀垫片",
    "VALVE": "阀",
    "SPECIAL MATERIAL TESTS": "特殊材料试验",
    "MINIMUM DESIGN METAL TEMPERATURE": "最低设计金属温度",
    "LOW AMBIENT TEMPERATURE MATERIALS TESTS": "低环境温度材料试验",
    "COMPLIANCE WITH INSPECTORS CHECK LIST": "符合检验员检查清单",
    "CERTIFICATION OF MATERIALS": "材料证明",
    "FINAL ASSEMBLY CLEARANCES": "最终装配间隙",
    "SURFACE AND SUBSURFACE EXAMINATIONS": "表面及近表面检查",
    "RADIOGRAPHY": "射线检测",
    "WELDS FOR PRESSURE CASING": "压力壳体焊缝",
    "TESTS": "试验",
    "ULTRASONIC": "超声检测",
    "MAGNETIC PARTICLE": "磁粉检测",
    "LIQUID PENETRANT": "渗透检测",
    "CLEANLINESS PRIOR TO FINAL ASSEMBLY": "最终装配前清洁度",
    "FURNISH PROCEDURES FOR OPTIONAL TESTS": "提供可选试验程序",
    "HYDROSTATIC": "水压试验",
    "STEADY STATE ACCURACY": "稳态精度",
    "REPEATABILITY": "重复性",
    "LINEARITY": "线性度",
    "NPSH TEST (5 POINT)": "NPSH 试验 (5 点)",
    "LUBRICATION FLUID": "润滑介质",
    "CRANKCASE": "曲轴箱",
    "INTERMEDIATE": "中间",
    "HYDRAULIC FLUID": "液压油",
    "ACCESSORIES": "附件",
    "SPEED REDUCER MANUFACTURER": "减速机制造商",
    "INTEGRAL": "整体式",
    "SEPARATE": "分体式",
    "RATIO": "传动比",
    "BASEPLATE UNDER": "底板在下",
    "COUPLING MANUFACTURER": "联轴器制造商",
    "CONTROLS": "控制",
    "SIGNAL": "信号",
    "MANUAL": "手动",
    "REMOTE": "远程",
    "PNEUMATIC": "气动",
    "AUTOMATIC": "自动",
    "LOCAL": "就地",
    "ELECTRONIC": "电子",
    "FLOW CONTROL RANGE": "流量控制范围",
    "OTHER PURCHASE REQUIREMENTS": "其他采购要求",
    "NAMEPLATE UNITS": "铭牌单位",
    "CUSTOMARY": "英制",
    "VENDOR FURNISHED PROCESS PIPING": "卖方提供工艺管道",
    "VENDOR REVIEW PIPING DRAWINGS": "卖方审查管道图",
    "VENDOR FURNISHED RELIEF VALVE": "卖方提供安全阀",
    "INTERNAL": "内置",
    "EXTERNAL": "外置",
    "RELIEF VALVE SETTING, INTERNAL/EXTERNAL": "安全阀整定，内置/外置",
    "RELIEF VALVE SETTING , INTERNAL/EXTERNAL": "安全阀整定，内置/外置",
    "VENDOR FURNISHED BACK-PRESSURE VALVE": "卖方提供背压阀",
    "DOUBLE CHECK VALVES REQUIRED": "需双止回阀",
    "OIL-FILLED PRESSURE GAUGES REQUIRED": "需充油压力表",
    "VENDOR FURNISHED CONTROL PANEL": "卖方提供控制柜",
    "BASEPLATE PREPARED FOR EPOXY GROUT": "底板适用于环氧灌浆",
    "PROVIDE TECHNICAL DATA MANUAL": "提供技术资料手册",
    "CALIBRATION POT ON PUMP SUCTION SIDE": "泵入口侧校验柱",
    "PREPARATION FOR SHIPMENT": "装运准备",
    "DOMESTIC": "国内",
    "EXPORT": "出口",
    "EXPORT BOXING": "出口箱装",
    "OUTDOOR STORAGE MORE THAN 12 MONTHS": "户外储存超过 12 个月",
    "WEIGHTS": "重量",
    "BASE": "底座",
    "MANUFACTURER": "制造商",
    "FRAME NO.": "机架号",
    "CONSTANT SPEED": "恒速",
    "VOLTS": "电压",
    "HERTZ": "频率",
    "ENCLOSURE": "防护等级",
    "SEE SEPARATE DATA SHEETS": "见单独数据表",
    "GAS DRIVEN": "燃气驱动",
    "STEAM TURBINE": "汽轮机",
    "SQUIRREL CAGE INDUCTION MOTORS": "鼠笼式感应电动机",
    "SEE ALSO SPECIFICATIONS FOR \"SQUIRREL CAGE": "另见“鼠笼式",
    "NOTE 5- VENDOR TO CONFIRM.": "说明 5- 卖方确认。",
    "NON - SPARKING COUPLING GUARD.": "无火花联轴器护罩。",
    "NOTE 17- PUMP IS IN CONTINIOUS OPERATION.": "说明 17- 泵连续运行。",
    "NOTE 18- DESIGN CASE: WINTER DESIGN CASE.": "说明 18- 设计工况：冬季设计工况。",
    "NOTE 22- DELETED": "说明 22- 删除",
    "NOTE 24- ACCURACY ACCORDING TO RELEVANT API": "说明 24- 精度按相关 API。",
    "REQ'D": "需",
    "RELATIVE HUMIDITY (NOR/MAX)": "相对湿度（正常/最大）",
    "CONTRACTOR APPD": "承包商批准",
}
EXACT_TRANSLATIONS.update(MECHANICAL_TERM_TRANSLATIONS)
TERM_TRANSLATIONS.update(MECHANICAL_TERM_TRANSLATIONS)

PAGE_PHRASE_TRANSLATIONS = {
    "FURNISHED: BY PUMP MFGR.": "供货：由泵制造商",
    "BY PUMP MFGR.": "由泵制造商",
    "BY PUMP MANUFACTURER": "由泵制造商",
    "BY MANUFACTURER": "由制造商",
    "MOUNTED:": "安装：",
    "NOT BY PUMP MANUFACTURER": "非泵制造商提供",
    "SUITABLE FOR EXPORT PACKING": "适合出口包装",
    "PUMPED FLUID": "泵送介质",
    "FABRICATED STEEL": "焊接钢结构",
    "HORIZONTAL, ON PUMP MFGRS. SKID": "卧式，安装在泵制造商底座上",
    "INCL. IN PUMP WT.": "包含在泵重量内",
    "□ PERFORMANCE": "□ 性能",
    "○ WINTERIZATION REQ'D": "○ 需冬季防护",
    "ZONE 2, GAS GROUP IIB, T3": "2 区，气体组别 IIB，T3",
    "ELECTRICITY DRIVERS HEATING CONTROL": "电源 驱动机 加热 控制",
    "70 MAXIMUM": "70 最大",
    "34 MAXIMUM": "34 最大",
    "36 MAXIMUM": "36 最大",
    "47.9 / 1.8 (AT SUCTION FLG. CENTERLINE) M": "47.9 / 1.8（入口法兰中心线处）m",
}
EXACT_TRANSLATIONS.update(PAGE_PHRASE_TRANSLATIONS)
TERM_TRANSLATIONS.update(PAGE_PHRASE_TRANSLATIONS)

PUBLIC_SHORT_PHRASE_TRANSLATIONS = {
    "DELIVERY SCHEDULE": "交付进度",
    "DIMENSIONAL OUTLINE DRAWING": "外形尺寸图",
    "MECHANICAL RUNNING TEST PROCEDURE": "机械运转试验程序",
    "PERFORMANCE TEST PROCEDURE": "性能试验程序",
    "DIRECTION OF ROTATION": "旋转方向",
    "MECHANICAL RUNNING TEST REPORT": "机械运转试验报告",
    "DRIVER TEST & CHARACTERISTIC REPORT": "驱动机试验及特性报告",
    "RUST PREVENTION / PAINTING": "防锈/涂漆",
    "PACKING LIST": "装箱单",
    "INSTRUMENT LIST": "仪表清单",
    "INSTRUMENT DRAWING": "仪表图纸",
    "PAINTING PROCEDURE / SCHEDULE": "涂漆程序/计划",
    "MATERIAL LIST": "材料清单",
    "INSPECTION REPORT": "检验报告",
    "FINAL DOCUMENTS": "最终文件",
}
EXACT_TRANSLATIONS.update(PUBLIC_SHORT_PHRASE_TRANSLATIONS)
TERM_TRANSLATIONS.update(PUBLIC_SHORT_PHRASE_TRANSLATIONS)

FIELD_LABEL_TRANSLATIONS = {
    "PROJECT": "项目",
    "JOB NO.": "工作号",
    "PREPARED BY": "编制",
    "CLIENT": "客户",
    "LOCATION": "地点",
    "SERVICE": "用途",
    "ITEM NUMBER": "位号",
    "ITEM NO.": "位号",
    "ITEM NO": "位号",
    "TAG NO.": "位号",
    "TAG NO": "位号",
    "TAG NUMBER": "位号",
    "PUMP ITEM NO.": "泵位号",
    "PUMP ITEM NO": "泵位号",
    "PUMP ITEM NO'S": "泵位号",
    "MODEL": "型号",
    "MODEL NO.": "型号",
    "MODEL NO": "型号",
    "MODEL NUMBER": "型号编号",
    "STANDARD": "标准",
    "DESIGN CODE": "设计标准",
    "CASE": "工况",
    "NOTE": "说明",
    "CURVE NO.": "曲线号",
    "PUMP VENDOR": "泵供应商",
    "FLUID TYPE": "介质类型",
    "SUCTION FROM": "入口来源",
    "DISCHARGE TO": "出口去向",
    "NORM / MIN / MAX TEMP": "正常/最小/最大温度",
    "VAPOR PRESSURE @ P&T": "蒸汽压 @ P&T",
    "LUBE QUALITY": "润滑油品质",
    "PUMP TYPE": "泵型",
    "FURNISHED BY": "供货方",
    "MOUNTED BY": "安装方",
    "SERVICE FACTOR": "服务系数",
    "KW RATING": "功率等级",
    "AREA CLASSIFICATION": "区域防爆等级",
    "VOLT/PHASE/CYCLE": "电压/相数/频率",
    "LIFTING LUGS": "吊耳",
    "EXPORT CRATING": "出口包装箱",
    "COOLING FLUID": "冷却介质",
    "SUPPLY TEMP": "供给温度",
    "SUPPLY PRES": "供给压力",
    "SIGHT FLOW INDICATOR": "视流指示器",
    "PIPE MATERIAL": "管道材料",
    "COOLING WATER PLAN": "冷却水方案",
    "SEAL FLUSH PLAN": "密封冲洗方案",
    "AUXILIARY SEAL PLAN": "辅助密封方案",
    "FULL LOAD": "满负荷",
    "HYDROTEST": "水压试验",
    "MECH RUN-IN": "机械试运行",
    "NPSH": "汽蚀余量",
    "FINAL INSPECTION": "最终检验",
    "TYPE": "类型",
    "SIZE": "尺寸",
    "MFGR.": "制造商",
    "NO. OF RINGS": "环数",
    "SPEED": "转速",
    "STAGES": "级数",
    "BHP": "制动功率",
    "RATED IMPELLER": "额定叶轮",
    "MAX HEAD": "最大扬程",
    "MAX BHP": "最大制动功率",
    "STAGE POWER": "单级功率",
    "COOLING WATER RATE": "冷却水流量",
    "SEAL FLUSH RATE": "密封冲洗流量",
    "PACKING RATE": "填料流量",
    "AUXILIARY TOTAL RATE": "辅助系统总流量",
    "MAX COOLING WATER PRESSURE": "最大冷却水压力",
    "CASE/BOWLS": "泵壳/碗体",
    "IMPELLER": "叶轮",
    "SHAFT": "轴",
    "SHAFT BEARINGS": "轴承",
    "THROAT BUSHING": "喉部衬套",
    "WEAR RINGS": "耐磨环",
    "SKID": "底座",
    "ROTATION, VIEW FROM COUPLING END": "旋转方向（从联轴器端看）",
    "WILL PASS": "可通过粒径",
    "DISCHARGE CASE": "出口壳体",
    "MOUNTING": "安装",
    "SPLIT": "剖分",
    "VOLUTE": "蜗壳",
    "CONNECTION": "连接",
    "MAWP": "最大允许工作压力",
    "MODEL NUMBER": "型号编号",
    "COLUMN PIPE": "立管",
    "LENGTH": "长度",
    "LINE SHAFT": "传动轴",
    "OPERATING THRUST": "运行推力",
    "PEAK THRUST": "峰值推力",
    "RATED THRUST": "额定推力",
    "PUMP DESIGN SPEC.": "泵设计规范",
    "EFF": "效率",
    "MIN SUBMER.": "最小浸没深度",
    "MIN FLOW": "最小流量",
    "SLEEVE": "轴套",
    "INTERNALS": "内件",
    "FLOAT & ROD": "浮子及杆",
    "COLUMN": "立管",
    "PIPE": "管道",
    "TUBING": "管路",
    "TOTAL": "总计",
    "RELATIVE HUMIDITY (NOR/MAX) (%)": "相对湿度（正常/最大）（%）",
}



TECHNICAL_ONLY_RE = re.compile(r"^[\s\d.,:/()\\+\-_%°A-Za-z]+$")
CODE_LIKE_RE = re.compile(
    r"^([A-Z]{1,8}[-/])?\d+[A-Z0-9./()\-:_ ]*$|^[A-Z]{1,8}[-/]\d+[A-Z0-9./()\-:_ ]*$"
)
CJK_RE = re.compile(r"[\u3400-\u9fff]")
LATIN_RE = re.compile(r"[A-Za-z]")
PROTECT_RE = re.compile(
    r"\b(?:API|ISO|IEC|ASME|ASTM|GB|DIN|EN|NACE|ANSI|NPSH)\s*[-A-Z0-9./]*\d[-A-Z0-9./]*\b"
    r"|\b[A-Z]{1,8}(?:[-/][A-Z0-9]+){1,}\b"
    r"|\b\d{2,4}(?:[-/][A-Z0-9]+){1,}\b"
    r"|\b\d+(?:\.\d+)?\s*(?:kg/h|m3/h|m³/h|bar|barg|bara|kW|KW|V|Hz|mm|cm|rpm|cP|MPa|kPa|Pa|%)\b"
    r"|\b(?:kg/h|m3/h|m³/h|bar|barg|bara|kW|KW|V|Hz|mm|cm|rpm|cP|MPa|kPa|Pa|%)\b"
)


@dataclass
class Segment:
    source_file: str
    page: int
    text: str
    translation: str
    bbox: tuple[float, float, float, float]


@dataclass
class TextRegion:
    page: int
    text: str
    translation: str
    bbox: tuple[float, float, float, float]
    align: str = "left"
    is_table: bool = False


def normalize_text(text: str) -> str:
    return " ".join((text or "").replace("\u00a0", " ").split())


def normalize_translated_text(text: str) -> str:
    normalized = normalize_text(text)
    return re.sub(r"(?<=[\u3400-\u9fff])\s+(?=[\u3400-\u9fff])", "", normalized)


def translated_pdf_name(source: Path) -> str:
    return f"{source.stem}-译.pdf"


def translated_office_name(source: Path) -> str:
    return f"{source.stem}-译{source.suffix.lower()}"


def active_output_file_name(default_name: str) -> str:
    return str(getattr(OUTPUT_NAMING_STATE, "file_name", None) or default_name)


def protect_terms(text: str) -> tuple[str, dict[str, str]]:
    replacements: dict[str, str] = {}

    def replace(match: re.Match) -> str:
        token = f"TKN{len(replacements):04d}X"
        replacements[token] = match.group(0)
        return token

    return PROTECT_RE.sub(replace, text), replacements


def restore_terms(text: str, replacements: dict[str, str]) -> str:
    restored = text
    for token, original in replacements.items():
        restored = restored.replace(token, original)
        restored = restored.replace(token.lower(), original)
    return restored


def needs_online_translation(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized or not LATIN_RE.search(normalized):
        return False
    if CJK_RE.search(normalized):
        return False
    if CODE_LIKE_RE.match(normalized):
        return False
    if len(normalized) <= 2:
        return False
    return True


def is_clause_or_sentence(text: str) -> bool:
    normalized = normalize_text(text)
    words = re.findall(r"[A-Za-z][A-Za-z']*", normalized)
    if len(words) >= 7:
        return True
    return bool(re.search(r"\b(?:shall|include|included|specified|provided|installed|recommended|submitted|required|quotation|approval)\b", normalized, flags=re.IGNORECASE))


def unprotected_latin_words(text: str) -> list[str]:
    protected, _ = protect_terms(normalize_text(text))
    protected = re.sub(r"TKN\d{4}X", "", protected)
    return re.findall(r"[A-Za-z][A-Za-z']*", protected)


def cached_translation_needs_refresh(source: str, translated: str) -> bool:
    source = normalize_text(source)
    translated = normalize_text(translated)
    if not translated:
        return True
    if not should_request_online_translation(source):
        return False
    residual_words = unprotected_latin_words(translated)
    if not residual_words:
        return False
    local_translation = normalize_text(translate_line(source, {}))
    if translated in {source, local_translation}:
        return True
    latin_count = len(re.findall(r"[A-Za-z]", " ".join(residual_words)))
    cjk_count = len(re.findall(r"[\u3400-\u9fff]", translated))
    return latin_count > max(24, cjk_count * 1.8)


def looks_like_protected_short_token(text: str) -> bool:
    normalized = normalize_text(text)
    if re.fullmatch(r"[A-Z]{2,}(?:\s+[A-Z]{2,}){0,2}", normalized):
        return True
    if re.fullmatch(r"(?:[A-Z]\.){1,4}[A-Za-z]+", normalized):
        return True
    if re.fullmatch(r"(?:[A-Z]\.){2,}[A-Z]?", normalized):
        return True
    return False


def translate_leading_bullet_exact(text: str) -> str | None:
    match = re.fullmatch(r"\s*-\s*(.+)", normalize_text(text))
    if not match:
        return None
    translated = EXACT_TRANSLATIONS.get(normalize_text(match.group(1)).upper())
    if not translated:
        return None
    return f"- {translated}"


def translate_leading_number_exact(text: str) -> str | None:
    match = re.fullmatch(r"(\d+)\s+(.+)", normalize_text(text))
    if not match:
        return None
    translated = EXACT_TRANSLATIONS.get(normalize_text(match.group(2)).upper())
    if not translated:
        return None
    return f"{match.group(1)} {translated}"


def should_request_online_translation(text: str) -> bool:
    normalized = normalize_text(text)
    if not needs_online_translation(normalized):
        return False
    if looks_like_protected_short_token(normalized):
        return False
    if normalized.upper() in EXACT_TRANSLATIONS:
        return False
    words = re.findall(r"[A-Za-z][A-Za-z']*", normalized)
    local_translation = translate_value(normalized)
    if local_translation != normalized:
        protected_local, _ = protect_terms(local_translation)
        protected_local = re.sub(r"TKN\d{4}X", "", protected_local)
        residual_words = re.findall(r"[A-Za-z][A-Za-z']*", protected_local)
        if not residual_words:
            return False
        return is_clause_or_sentence(normalized) or len(words) >= 4
    return is_clause_or_sentence(normalized) or len(words) >= 3


def translate_line(text: str, translation_cache: dict[str, str] | None = None) -> str:
    original = normalize_text(text)
    if not original:
        return ""
    exact_key = original.upper()
    if exact_key in EXACT_TRANSLATIONS:
        return EXACT_TRANSLATIONS[exact_key]
    if (
        translation_cache
        and original in translation_cache
        and should_request_online_translation(original)
    ):
        cached_translation = postprocess_translation(original, translation_cache[original])
        if not cached_translation_needs_refresh(original, cached_translation):
            return cached_translation
    bullet_exact = translate_leading_bullet_exact(original)
    if bullet_exact:
        return bullet_exact
    number_exact = translate_leading_number_exact(original)
    if number_exact:
        return number_exact
    field_match = re.fullmatch(r"(.+?):\s*(.*)", original)
    if field_match:
        label = normalize_text(field_match.group(1)).upper()
        value = field_match.group(2).strip()
        if label in FIELD_LABEL_TRANSLATIONS:
            translated_label = FIELD_LABEL_TRANSLATIONS[label]
            return f"{translated_label}：{translate_value(value)}" if value else f"{translated_label}："
    page_match = re.fullmatch(r"page\s+(\d+)\s+of\s+(\d+)", original, flags=re.IGNORECASE)
    if page_match:
        return f"第 {page_match.group(1)} 页，共 {page_match.group(2)} 页"
    field_match = re.fullmatch(r"contract\s+no\.?:?\s*(.*)", original, flags=re.IGNORECASE)
    if field_match:
        return f"合同号：{field_match.group(1).strip()}".rstrip("：")
    field_match = re.fullmatch(r"doc\.\s*no\.?:?\s*(.*)", original, flags=re.IGNORECASE)
    if field_match:
        return f"文件号：{field_match.group(1).strip()}".rstrip("：")
    field_match = re.fullmatch(r"rev\.?:?\s*(.*)", original, flags=re.IGNORECASE)
    if field_match:
        return f"版本：{field_match.group(1).strip()}".rstrip("：")
    field_match = re.fullmatch(r"class:?\s*(.*)", original, flags=re.IGNORECASE)
    if field_match:
        return f"等级：{field_match.group(1).strip()}".rstrip("：")
    for label, target in (
        ("DOCUMENT TYPE", "文件类型"),
        ("DOCUMENT TITLE", "文件标题"),
        ("DOCUMENT NUMBER", "文件编号"),
        ("SITE", "地点"),
    ):
        if exact_key.startswith(label + " "):
            tail = original[len(label) :].strip(" :")
            return f"{target}：{translate_value(tail)}" if tail else target
    if CODE_LIKE_RE.match(original):
        return original
    if ":" in original:
        left, right = original.split(":", 1)
        translated_left = translate_line(left)
        translated_right = translate_value(right)
        return f"{translated_left}：{translated_right}".strip()
    local_translation = translate_value(original)
    if local_translation != original:
        return local_translation
    return local_translation


def postprocess_translation(source: str, translated: str) -> str:
    result = normalize_translated_text(translated)
    if "TRAIN" in source.upper():
        result = result.replace("列车", "线")
        result = result.replace("火车", "线")
    if "DOC. NO" in source.upper():
        result = result.replace("博士。编号", "文件号").replace("博士编号", "文件号")
    result = result.replace("供应", "供货") if "SUPPLY" in source.upper() else result
    return result


def online_translate_text(text: str) -> str:
    protected, replacements = protect_terms(text)
    if GoogleTranslator is not None:
        translated = GoogleTranslator(source="en", target="zh-CN").translate(protected)
    else:
        translated = google_http_translate(protected)
    return postprocess_translation(text, restore_terms(normalize_text(translated), replacements))


def google_http_translate(text: str, timeout: int = 10) -> str:
    params = {"client": "gtx", "sl": "en", "tl": "zh-CN", "dt": "t", "q": text}
    if requests is not None:
        response = requests.get("https://translate.googleapis.com/translate_a/single", params=params, timeout=timeout)
        response.raise_for_status()
        data = response.json()
    else:
        url = "https://translate.googleapis.com/translate_a/single?" + urllib.parse.urlencode(params)
        with urllib.request.urlopen(url, timeout=timeout) as response:
            data = json.loads(response.read().decode("utf-8"))
    return "".join(part[0] for part in data[0] if part and part[0])


def build_online_translation_cache(segments: list[Segment], cache_path: Path) -> dict[str, str]:
    cache_path.parent.mkdir(parents=True, exist_ok=True)
    if cache_path.exists():
        cache = json.loads(cache_path.read_text(encoding="utf-8"))
    else:
        cache = {}

    candidates = []
    for segment in segments:
        text = normalize_text(segment.text)
        if should_request_online_translation(text) and (text not in cache or cached_translation_needs_refresh(text, cache.get(text, ""))):
            candidates.append(text)

    pending: list[tuple[str, str, dict[str, str]]] = []
    for text in candidates:
        protected, replacements = protect_terms(text)
        pending.append((text, protected, replacements))

    separator = " ZXQSEP000 "
    batches: list[list[tuple[str, str, dict[str, str]]]] = []
    current: list[tuple[str, str, dict[str, str]]] = []
    current_len = 0
    for item in pending:
        item_len = len(item[1]) + len(separator)
        if current and current_len + item_len > 2600:
            batches.append(current)
            current = []
            current_len = 0
        current.append(item)
        current_len += item_len
    if current:
        batches.append(current)

    for batch in batches:
        try:
            translated_joined = google_http_translate(separator.join(item[1] for item in batch), timeout=10)
            translated_batch = translated_joined.split(separator)
            if len(translated_batch) != len(batch):
                raise ValueError("translated batch separator mismatch")
            for (text, _protected, replacements), translated in zip(batch, translated_batch):
                restored = restore_terms(normalize_text(translated), replacements)
                cache[text] = postprocess_translation(text, restored)
        except Exception:
            for text, protected, replacements in batch:
                try:
                    translated = google_http_translate(protected, timeout=8)
                    restored = restore_terms(normalize_text(translated), replacements)
                    cache[text] = postprocess_translation(text, restored)
                except Exception:
                    cache[text] = translate_line(text, {})
        cache_path.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")

    cache_path.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")
    return cache


def office_requires_model_translation(text: str) -> bool:
    return should_request_online_translation(text)


def translate_office_text(
    original: str,
    translation_cache: dict[str, str] | None,
) -> tuple[str, str]:
    private_terms = (translation_cache or {}).get(PRIVATE_GLOSSARY_CACHE_KEY, {})
    if isinstance(private_terms, dict):
        private_translation = normalize_text(
            str(private_terms.get(normalize_text(original).casefold(), ""))
        )
        if private_translation:
            return private_translation, "private_glossary"
    if office_requires_model_translation(original):
        cached = normalize_text(str((translation_cache or {}).get(original, "")))
        if not cached or cached_translation_needs_refresh(original, cached):
            return original, "untranslated"
        return translate_line(original, translation_cache), "model_cache"

    translated = translate_line(original, translation_cache)
    return translated, "local_rules" if translated != original else "untranslated"


def request_openai_compatible_office_batch(
    items: list[tuple[str, str]],
    provider: dict,
) -> dict[str, str]:
    api_key_env = provider.get("api_key_env") or VECTOR_ENGINE_API_KEY_ENV
    api_key = user_environment_value(api_key_env)
    if not api_key:
        raise RuntimeError(f"未配置 {api_key_env}，Office 未知英文句子无法调用模型翻译")

    request_payload = {
        "model": provider["model"],
        "temperature": 0,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You translate mechanical-industry RFQ text into concise Simplified Chinese. "
                    "Return only a JSON object mapping every supplied id to its translation. "
                    "Preserve TKN####X placeholders, Tag numbers, models, standards and units exactly. "
                    f"Prompt contract: {OFFICE_PROMPT_CONTRACT_VERSION}; "
                    f"target language: {OFFICE_TRANSLATION_TARGET_LANGUAGE}."
                ),
            },
            {
                "role": "user",
                "content": json.dumps(
                    {item_id: protected for item_id, protected in items},
                    ensure_ascii=False,
                ),
            },
        ],
    }
    endpoint = provider["base_url"].rstrip("/") + "/chat/completions"
    request = urllib.request.Request(
        endpoint,
        data=json.dumps(request_payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    timeout = positive_env_int("B_OFFICE_TRANSLATION_TIMEOUT", 90, maximum=300)
    with urllib.request.urlopen(request, timeout=timeout) as response:
        payload = json.loads(response.read().decode("utf-8"))
    content = str(payload["choices"][0]["message"]["content"]).strip()
    fenced = re.fullmatch(r"```(?:json)?\s*(.*?)\s*```", content, flags=re.DOTALL | re.IGNORECASE)
    if fenced:
        content = fenced.group(1).strip()
    translated = json.loads(content)
    if not isinstance(translated, dict):
        raise ValueError("Office 批量翻译响应必须是 JSON 对象")
    return {
        item_id: normalize_text(str(translated[item_id]))
        for item_id, _ in items
        if normalize_text(str(translated.get(item_id, "")))
    }


def office_batch_failure_code(exc: Exception) -> str:
    """Return a stable diagnostic code without exposing request content."""

    message = normalize_text(str(exc)).casefold()
    if isinstance(exc, json.JSONDecodeError) or "json" in message:
        return "malformed_response"
    if isinstance(exc, KeyError) or "缺少" in message or "miss" in message:
        return "missing_response_item"
    if isinstance(exc, TimeoutError) or "timed out" in message or "timeout" in message:
        return "request_timeout"
    if isinstance(exc, urllib.error.HTTPError):
        return "http_error"
    if isinstance(exc, urllib.error.URLError):
        return "connection_error"
    return "batch_request_failed"


def build_office_translation_cache(
    texts: Iterable[str],
    cache_path: Path,
) -> tuple[dict[str, str], dict]:
    cache_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        cache_payload = json.loads(cache_path.read_text(encoding="utf-8")) if cache_path.exists() else {}
    except (OSError, json.JSONDecodeError):
        cache_payload = {}
    if not isinstance(cache_payload, dict):
        cache_payload = {}

    office_signature = office_config_signature()
    namespaces = cache_payload.setdefault(OFFICE_CACHE_NAMESPACES_KEY, {})
    if not isinstance(namespaces, dict):
        namespaces = {}
        cache_payload[OFFICE_CACHE_NAMESPACES_KEY] = namespaces
    cache = namespaces.setdefault(office_signature, {})
    if not isinstance(cache, dict):
        cache = {}
        namespaces[office_signature] = cache

    private_terms, private_glossary_summary = load_private_glossary_terms()
    requested = []
    seen = set()
    for raw_text in texts:
        text = normalize_text(raw_text)
        if text in seen:
            continue
        seen.add(text)
        if text.casefold() in private_terms:
            continue
        if not office_requires_model_translation(text):
            continue
        cached = normalize_text(str(cache.get(text, "")))
        if cached and not cached_translation_needs_refresh(text, cached):
            continue
        requested.append(text)

    prepared_items: list[tuple[str, str, dict[str, str], str]] = []
    for index, text in enumerate(requested, start=1):
        protected, replacements = protect_terms(text)
        prepared_items.append((f"s{index:04d}", protected, replacements, text))

    provider = office_model_provider_config()
    item_diagnostics = {
        item_id: {
            "item_id": item_id,
            "status": "pending",
            "request_attempts": 0,
            "max_split_depth": 0,
            "cache_written": False,
            "resolution": None,
            "last_failure_code": None,
        }
        for item_id, _protected, _replacements, _source in prepared_items
    }
    diagnostics = {
        "model_configured": bool(provider.get("api_key_configured")),
        "requested_count": len(requested),
        "translated_count": 0,
        "cache_hit_count": len(seen) - len(requested),
        "unresolved_count": 0,
        "batch_request_count": 0,
        "batch_split_count": 0,
        "single_item_retry_count": 0,
        "single_item_retry_success_count": 0,
        "local_fallback_count": 0,
        "item_diagnostics": [],
        "warnings": [],
        "errors": [],
        "config_signature": office_signature,
        "target_language": OFFICE_TRANSLATION_TARGET_LANGUAGE,
        "prompt_contract_version": OFFICE_PROMPT_CONTRACT_VERSION,
        "protection_contract_version": OFFICE_PROTECTION_CONTRACT_VERSION,
        "office_batch_isolation_version": (
            OFFICE_BATCH_FAILURE_ISOLATION_CONTRACT_VERSION
        ),
        "private_glossary": private_glossary_summary,
    }
    effective_cache = dict(cache)
    effective_cache[PRIVATE_GLOSSARY_CACHE_KEY] = private_terms
    if not requested:
        return effective_cache, diagnostics
    if not diagnostics["model_configured"]:
        diagnostics["unresolved_count"] = len(requested)
        for item in item_diagnostics.values():
            item["status"] = "unresolved"
            item["last_failure_code"] = "model_not_configured"
        diagnostics["item_diagnostics"] = list(item_diagnostics.values())
        diagnostics["warnings"].append(
            f"未配置 {provider.get('api_key_env') or VECTOR_ENGINE_API_KEY_ENV}；"
            f"{len(requested)} 条词典外 Office 英文句子保留原文"
        )
        return effective_cache, diagnostics

    batch_size = positive_env_int("B_OFFICE_TRANSLATION_BATCH_SIZE", 16, maximum=32)
    max_chars = positive_env_int("B_OFFICE_TRANSLATION_BATCH_MAX_CHARS", 6000, maximum=12000)
    single_item_attempts = positive_env_int(
        "B_OFFICE_TRANSLATION_SINGLE_ITEM_ATTEMPTS",
        2,
        maximum=3,
    )
    pending_batches: list[list[tuple[str, str, dict[str, str], str]]] = []
    current: list[tuple[str, str, dict[str, str], str]] = []
    current_chars = 0
    for item in prepared_items:
        _item_id, protected, _replacements, _source = item
        if current and (len(current) >= batch_size or current_chars + len(protected) > max_chars):
            pending_batches.append(current)
            current = []
            current_chars = 0
        current.append(item)
        current_chars += len(protected)
    if current:
        pending_batches.append(current)

    unresolved_ids: set[str] = set()

    def persist_cache() -> None:
        write_json_file(cache_path, cache_payload)

    def record_request_attempt(
        batch: list[tuple[str, str, dict[str, str], str]],
        depth: int,
    ) -> None:
        diagnostics["batch_request_count"] += 1
        for item_id, _protected, _replacements, _source in batch:
            item = item_diagnostics[item_id]
            item["request_attempts"] += 1
            item["max_split_depth"] = max(item["max_split_depth"], depth)

    def translated_item(
        item: tuple[str, str, dict[str, str], str],
        response: dict[str, str],
    ) -> tuple[str | None, str | None]:
        item_id, _protected, replacements, source = item
        if not isinstance(response, dict):
            return None, "malformed_response"
        raw_translation = normalize_text(str(response.get(item_id, "")))
        if not raw_translation:
            return None, "missing_response_item"
        missing_tokens = [
            token
            for token in replacements
            if token.casefold() not in raw_translation.casefold()
        ]
        if missing_tokens:
            return None, "protected_token_missing"
        translated = postprocess_translation(
            source,
            restore_terms(raw_translation, replacements),
        )
        if cached_translation_needs_refresh(source, translated):
            return None, "english_residual"
        return translated, None

    def commit_translation(
        item: tuple[str, str, dict[str, str], str],
        translated: str,
        resolution: str,
    ) -> None:
        item_id, _protected, _replacements, source = item
        cache[source] = translated
        persist_cache()
        item_diagnostic = item_diagnostics[item_id]
        item_diagnostic.update(
            {
                "status": "translated",
                "cache_written": True,
                "resolution": resolution,
            }
        )
        unresolved_ids.discard(item_id)
        diagnostics["translated_count"] += 1

    def mark_request_failure(
        batch: list[tuple[str, str, dict[str, str], str]],
        failure_code: str,
    ) -> None:
        for item_id, _protected, _replacements, _source in batch:
            item_diagnostics[item_id]["last_failure_code"] = failure_code

    def request_once(
        batch: list[tuple[str, str, dict[str, str], str]],
        depth: int,
    ) -> tuple[dict[str, str] | None, str | None]:
        record_request_attempt(batch, depth)
        try:
            response_payload = request_openai_compatible_office_batch(
                [(item_id, protected) for item_id, protected, _, _ in batch],
                provider,
            )
        except Exception as exc:
            failure_code = office_batch_failure_code(exc)
            mark_request_failure(batch, failure_code)
            return None, failure_code
        if not isinstance(response_payload, dict):
            mark_request_failure(batch, "malformed_response")
            return None, "malformed_response"
        return response_payload, None

    def try_local_fallback(
        item: tuple[str, str, dict[str, str], str],
    ) -> bool:
        item_id, _protected, _replacements, source = item
        item_diagnostics[item_id]["resolution"] = "local_fallback_checked"
        local_translation = normalize_text(translate_line(source, {}))
        if local_translation == source or cached_translation_needs_refresh(source, local_translation):
            return False
        commit_translation(item, local_translation, "local_fallback")
        diagnostics["local_fallback_count"] += 1
        return True

    def process_single_item(
        item: tuple[str, str, dict[str, str], str],
        depth: int,
    ) -> None:
        item_id = item[0]
        for _attempt in range(single_item_attempts):
            diagnostics["single_item_retry_count"] += 1
            response, failure_code = request_once([item], depth)
            if response is None:
                item_diagnostics[item_id]["last_failure_code"] = failure_code
                continue
            translated, failure_code = translated_item(item, response)
            if translated is not None:
                commit_translation(item, translated, "single_item_retry")
                diagnostics["single_item_retry_success_count"] += 1
                return
            item_diagnostics[item_id]["last_failure_code"] = failure_code
        if try_local_fallback(item):
            return
        unresolved_ids.add(item_id)
        item_diagnostics[item_id]["status"] = "unresolved"

    def process_batch(
        batch: list[tuple[str, str, dict[str, str], str]],
        depth: int = 0,
    ) -> None:
        if not batch:
            return
        if len(batch) == 1:
            process_single_item(batch[0], depth)
            return

        response, failure_code = request_once(batch, depth)
        if response is None:
            diagnostics["batch_split_count"] += 1
            midpoint = max(1, len(batch) // 2)
            process_batch(batch[:midpoint], depth + 1)
            process_batch(batch[midpoint:], depth + 1)
            return

        rejected: list[tuple[str, str, dict[str, str], str]] = []
        for item in batch:
            item_id = item[0]
            translated, item_failure_code = translated_item(item, response)
            if translated is None:
                item_diagnostics[item_id]["last_failure_code"] = item_failure_code
                rejected.append(item)
                continue
            commit_translation(item, translated, "batch")

        if rejected:
            if len(rejected) > 1:
                diagnostics["batch_split_count"] += 1
                midpoint = max(1, len(rejected) // 2)
                process_batch(rejected[:midpoint], depth + 1)
                process_batch(rejected[midpoint:], depth + 1)
            else:
                process_single_item(rejected[0], depth + 1)

    for batch in pending_batches:
        process_batch(batch)

    persist_cache()
    diagnostics["unresolved_count"] = len(unresolved_ids)
    diagnostics["item_diagnostics"] = list(item_diagnostics.values())
    if unresolved_ids:
        diagnostics["warnings"].append(
            f"Office 模型翻译仍有 {len(unresolved_ids)} 条未解决；"
            f"其余 {diagnostics['translated_count']} 条已独立写入缓存"
        )
        diagnostics["errors"].append(
            f"office_translation_unresolved_items:{len(unresolved_ids)}"
        )
    effective_cache = dict(cache)
    effective_cache[PRIVATE_GLOSSARY_CACHE_KEY] = private_terms
    return effective_cache, diagnostics


def build_layout_translation_cache(source_regions: list[TextRegion], source_file: str, cache_path: Path) -> dict[str, str]:
    segments = [
        Segment(
            source_file=source_file,
            page=region.page,
            text=region.text,
            translation="",
            bbox=region.bbox,
        )
        for region in source_regions
    ]
    return build_online_translation_cache(segments, cache_path)


def translate_value(text: str) -> str:
    result = normalize_text(text)
    if not result:
        return ""
    upper = result.upper()
    if upper in EXACT_TRANSLATIONS:
        return EXACT_TRANSLATIONS[upper]
    for source, target in sorted(TERM_TRANSLATIONS.items(), key=lambda item: len(item[0]), reverse=True):
        pattern = re.compile(rf"(?<![A-Z0-9]){re.escape(source)}(?![A-Z0-9])", re.IGNORECASE)
        result = pattern.sub(target, result)
    return normalize_translated_text(result)


def group_words_into_lines(words: list[dict], tolerance: float = 3.0) -> list[tuple[str, tuple[float, float, float, float]]]:
    sorted_words = sorted(words, key=lambda word: (float(word["top"]), float(word["x0"])))
    groups: list[list[dict]] = []
    for word in sorted_words:
        if not groups or abs(float(word["top"]) - float(groups[-1][0]["top"])) > tolerance:
            groups.append([word])
        else:
            groups[-1].append(word)

    lines = []
    for group in groups:
        group = sorted(group, key=lambda word: float(word["x0"]))
        text = normalize_text(" ".join(str(word["text"]) for word in group))
        if not text:
            continue
        x0 = min(float(word["x0"]) for word in group)
        top = min(float(word["top"]) for word in group)
        x1 = max(float(word["x1"]) for word in group)
        bottom = max(float(word["bottom"]) for word in group)
        lines.append((text, (x0, top, x1, bottom)))
    return lines


def text_width(text: str, size: float) -> float:
    return pdfmetrics.stringWidth(text, FONT_NAME, size)


def wrap_cjk_text(text: str, max_width: float, font_size: float) -> list[str]:
    normalized = normalize_text(text)
    if not normalized:
        return []
    lines: list[str] = []
    current = ""
    for char in normalized:
        candidate = current + char
        if current and text_width(candidate, font_size) > max_width:
            lines.append(current)
            current = char.lstrip()
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def fit_font_size(text: str, max_width: float, max_height: float) -> float:
    size = min(8.0, max(4.2, max_height * 0.82))
    while size > 4.2 and text_width(text, size) > max_width:
        size -= 0.4
    return max(size, 4.2)


def fit_wrapped_font_size(text: str, max_width: float, max_height: float, max_size: float = 8.6) -> tuple[float, list[str]]:
    size = max_size
    while size >= 4.2:
        lines = wrap_cjk_text(text, max_width, size)
        if lines and len(lines) * size * 1.18 <= max_height:
            return size, lines
        size -= 0.3
    return 4.2, wrap_cjk_text(text, max_width, 4.2)


def word_has_latin(word: dict) -> bool:
    text = str(word.get("text", ""))
    return bool(LATIN_RE.search(text)) and not CJK_RE.search(text)


def group_words_into_word_lines(words: list[dict], tolerance: float = 3.0) -> list[list[dict]]:
    sorted_words = sorted(words, key=lambda word: (float(word["top"]), float(word["x0"])))
    groups: list[list[dict]] = []
    for word in sorted_words:
        if not groups or abs(float(word["top"]) - float(groups[-1][0]["top"])) > tolerance:
            groups.append([word])
        else:
            groups[-1].append(word)
    return [sorted(group, key=lambda word: float(word["x0"])) for group in groups]


def bbox_from_words(words: list[dict]) -> tuple[float, float, float, float]:
    return (
        min(float(word["x0"]) for word in words),
        min(float(word["top"]) for word in words),
        max(float(word["x1"]) for word in words),
        max(float(word["bottom"]) for word in words),
    )


def bbox_inside(inner: tuple[float, float, float, float], outer: tuple[float, float, float, float], margin: float = 1.5) -> bool:
    ix0, itop, ix1, ibottom = inner
    ox0, otop, ox1, obottom = outer
    return ix0 >= ox0 - margin and ix1 <= ox1 + margin and itop >= otop - margin and ibottom <= obottom + margin


def word_center_inside_bbox(word: dict, bbox: tuple[float, float, float, float], margin: float = 1.0) -> bool:
    x0, top, x1, bottom = bbox
    center_x = (float(word["x0"]) + float(word["x1"])) / 2
    center_y = (float(word["top"]) + float(word["bottom"])) / 2
    return x0 - margin <= center_x <= x1 + margin and top - margin <= center_y <= bottom + margin


def words_inside_bbox(words: list[dict], bbox: tuple[float, float, float, float]) -> list[dict]:
    return [word for word in words if word_center_inside_bbox(word, bbox)]


def split_words_by_horizontal_gap(words: list[dict], gap_threshold: float = 16.0) -> list[list[dict]]:
    sorted_words = sorted(words, key=lambda word: float(word["x0"]))
    chunks: list[list[dict]] = []
    current: list[dict] = []
    previous_x1: float | None = None
    for word in sorted_words:
        x0 = float(word["x0"])
        if current and previous_x1 is not None and x0 - previous_x1 > gap_threshold:
            chunks.append(current)
            current = []
        current.append(word)
        previous_x1 = float(word["x1"])
    if current:
        chunks.append(current)
    return chunks


def datasheet_form_marker_count(text: str) -> int:
    upper = normalize_text(text).upper()
    count = 0
    for marker in DATASHEET_FORM_SPLIT_MARKERS:
        if re.search(rf"(?<![A-Z0-9]){re.escape(marker)}(?![A-Z0-9])", upper):
            count += 1
    return count


def should_keep_table_line_whole(line_words: list[dict]) -> bool:
    text = normalize_text(" ".join(str(word["text"]) for word in line_words))
    if not text:
        return True
    if datasheet_form_marker_count(text) >= 2:
        return False
    if text.startswith("-"):
        return True
    if re.search(r"[a-z]", text):
        return True
    if re.search(
        r"\b(?:shall|include|included|specified|provided|installed|recommended|submitted|required|quotation|approval)\b",
        text,
        flags=re.IGNORECASE,
    ):
        return True
    return False


def split_table_repaint_line_words(
    line_words: list[dict],
    cell_bbox: tuple[float, float, float, float],
) -> list[list[dict]]:
    chunks = split_words_by_horizontal_gap(line_words, gap_threshold=18.0)
    if len(chunks) <= 1:
        return [line_words]
    cell_width = cell_bbox[2] - cell_bbox[0]
    if cell_width < 100.0:
        return [line_words]
    if should_keep_table_line_whole(line_words):
        return [line_words]
    return chunks


def table_chunk_alignment(words: list[dict], cell_bbox: tuple[float, float, float, float]) -> str:
    x0, _top, x1, _bottom = bbox_from_words(words)
    cell_x0, _cell_top, cell_x1, _cell_bottom = cell_bbox
    cell_width = cell_x1 - cell_x0
    text = normalize_text(" ".join(str(word["text"]) for word in words))
    alpha_words = re.findall(r"[A-Za-z][A-Za-z']*", text)
    if cell_width > 70.0 and len(alpha_words) >= 2 and (text.startswith("-") or re.search(r"[a-z]", text)):
        return "left"
    chunk_center = (x0 + x1) / 2
    cell_center = (cell_x0 + cell_x1) / 2
    if abs(chunk_center - cell_center) < max(8.0, (cell_x1 - cell_x0) * 0.12):
        return "center"
    return "left"


def add_region_from_words(
    regions: list[TextRegion],
    page_number: int,
    source_words: list[dict],
    page_width: float,
    is_table: bool,
    cell_bbox: tuple[float, float, float, float] | None = None,
    align_override: str | None = None,
) -> None:
    latin_words = [word for word in source_words if word_has_latin(word)]
    if not latin_words:
        return
    all_text = normalize_text(" ".join(str(word["text"]) for word in source_words))
    if not is_table:
        upper_all = all_text.upper()
        for document_label in ("DOCUMENT TYPE", "DOCUMENT TITLE", "DOCUMENT NUMBER", "SITE"):
            if not upper_all.startswith(document_label):
                continue
            colon_index = next((index for index, word in enumerate(source_words) if str(word["text"]).strip() == ":"), None)
            if colon_index is None:
                continue
            label_words = source_words[: colon_index + 1]
            value_words = source_words[colon_index + 1 :]
            regions.append(
                TextRegion(
                    page=page_number,
                    text=f"{document_label}:",
                    translation="",
                    bbox=bbox_from_words(label_words),
                    align="left",
                    is_table=False,
                )
            )
            if value_words:
                regions.append(
                    TextRegion(
                        page=page_number,
                        text=normalize_text(" ".join(str(word["text"]) for word in value_words)),
                        translation="",
                        bbox=bbox_from_words(value_words),
                        align="left",
                        is_table=False,
                    )
                )
            return
    latin_text = normalize_text(" ".join(str(word["text"]) for word in latin_words))
    latin_upper = latin_text.upper()
    keep_context_numbers = (
        re.search(r"\bpage\b", all_text, flags=re.IGNORECASE)
        and re.search(r"\bof\b", all_text, flags=re.IGNORECASE)
        and re.search(r"\d", all_text)
    ) or (
        not is_table
        and re.search(r"\b(?:PHASE|TRAIN|UNIT)\b", all_text, flags=re.IGNORECASE)
        and re.search(r"\d", all_text)
    ) or (
        not is_table
        and re.search(r"\d", all_text)
        and latin_upper not in EXACT_TRANSLATIONS
        and latin_upper not in TERM_TRANSLATIONS
        and (is_clause_or_sentence(all_text) or all_text.upper() in EXACT_TRANSLATIONS)
    )
    if keep_context_numbers:
        text_words = source_words
        text = all_text
    else:
        text_words = latin_words
        text = latin_text
    text = text.replace("��", "").replace("�", "").strip()
    if not text:
        return
    if text.replace(" ", "").upper() == "PAGE":
        return
    upper = text.upper()
    if not needs_online_translation(text) and upper not in EXACT_TRANSLATIONS and upper not in TERM_TRANSLATIONS:
        return
    align = align_override or (table_chunk_alignment(latin_words, cell_bbox) if is_table and cell_bbox else line_alignment(latin_words, page_width))
    regions.append(
        TextRegion(
            page=page_number,
            text=text,
            translation="",
            bbox=bbox_from_words(text_words),
            align=align,
            is_table=is_table,
        )
    )


def word_has_repaintable_content(word: dict) -> bool:
    text = str(word.get("text", ""))
    return bool(re.search(r"[A-Za-z0-9]", text)) and not CJK_RE.search(text)


def add_repaint_region_from_words(
    regions: list[TextRegion],
    page_number: int,
    source_words: list[dict],
    page_width: float,
    is_table: bool,
    cell_bbox: tuple[float, float, float, float] | None = None,
    align_override: str | None = None,
    translation_cache: dict[str, str] | None = None,
) -> None:
    repaint_words = [word for word in source_words if word_has_repaintable_content(word)]
    if not repaint_words:
        return
    text = normalize_text(" ".join(str(word["text"]) for word in source_words))
    if not text:
        return
    translation = translate_line(text, translation_cache)
    if not translation or text.replace(" ", "").upper() == "PAGE":
        translation = text
    align = align_override or (
        table_chunk_alignment(repaint_words, cell_bbox) if is_table and cell_bbox else line_alignment(repaint_words, page_width)
    )
    regions.append(
        TextRegion(
            page=page_number,
            text=text,
            translation=translation,
            bbox=bbox_from_words(source_words),
            align=align,
            is_table=is_table,
        )
    )


def extract_table_bboxes(page: pdfplumber.page.Page) -> list[tuple[float, float, float, float]]:
    bboxes = []
    for table in page.find_tables():
        if table.bbox:
            bboxes.append(tuple(float(v) for v in table.bbox))
    return bboxes


def choose_page_layout_strategy(page: pdfplumber.page.Page) -> str:
    table_bboxes = extract_table_bboxes(page)
    page_area = max(1.0, float(page.width) * float(page.height))
    table_area = sum(max(0.0, x1 - x0) * max(0.0, bottom - top) for x0, top, x1, bottom in table_bboxes)
    table_area_ratio = table_area / page_area
    table_cell_count = sum(len(table.cells) for table in page.find_tables())
    vector_count = len(getattr(page, "lines", [])) + len(getattr(page, "rects", []))
    word_count = len(page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False))

    if table_cell_count >= 12 and table_area_ratio >= 0.12:
        return LAYOUT_STRATEGY_TABLE_REPAINT
    if table_bboxes and table_area_ratio >= 0.28:
        return LAYOUT_STRATEGY_TABLE_REPAINT
    if not table_bboxes and vector_count >= 10 and word_count >= 4:
        return LAYOUT_STRATEGY_TABLE_REPAINT
    return LAYOUT_STRATEGY_TEMPLATE


def collect_auto_layout_regions(
    pdf_path: Path,
    page_indices: list[int],
    translation_cache: dict[str, str] | None = None,
) -> list[TextRegion]:
    regions: list[TextRegion] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_index in page_indices:
            strategy = choose_page_layout_strategy(pdf.pages[page_index])
            if strategy == LAYOUT_STRATEGY_TABLE_REPAINT:
                regions.extend(collect_table_repaint_layout_regions(pdf_path, [page_index], translation_cache=translation_cache))
            else:
                regions.extend(collect_template_layout_regions(pdf_path, [page_index], translation_cache=translation_cache))
    return regions


def draw_translated_line(
    c: canvas.Canvas,
    right_offset: float,
    page_height: float,
    translated: str,
    bbox: tuple[float, float, float, float],
) -> None:
    x0, top, x1, bottom = bbox
    if not translated:
        return
    left = right_offset + x0
    bottom_y = page_height - bottom
    width = max(8.0, x1 - x0)
    height = max(5.0, bottom - top)
    c.setFillColorRGB(1, 1, 1)
    pad_x = 1.8
    pad_y = 2.2 if height < 9 else 1.4
    c.rect(left - pad_x, bottom_y - pad_y, width + pad_x * 2, height + pad_y * 2, stroke=0, fill=1)
    font_size = fit_font_size(translated, width, height)
    c.setFillColorRGB(0, 0, 0)
    c.setFont(FONT_NAME, font_size)
    draw_translated_text(c, left, bottom_y + max(0.4, (height - font_size) / 2), translated)


def draw_translated_text(c: canvas.Canvas, x: float, y: float, text: str, center: bool = False) -> None:
    for offset in TRANSLATED_TEXT_DARKEN_OFFSETS:
        if center:
            c.drawCentredString(x + offset, y, text)
        else:
            c.drawString(x + offset, y, text)


def draw_text_region(
    c: canvas.Canvas,
    right_offset: float,
    page_height: float,
    region: TextRegion,
) -> None:
    x0, top, x1, bottom = region.bbox
    left = right_offset + x0
    bottom_y = page_height - bottom
    width = max(8.0, x1 - x0)
    height = max(5.0, bottom - top)
    pad_x, pad_y = overlay_padding(region)
    c.setFillColorRGB(1, 1, 1)
    c.rect(left - pad_x, bottom_y - pad_y, width + pad_x * 2, height + pad_y * 2, stroke=0, fill=1)
    max_size = min(9.2, max(5.0, height * (0.75 if region.is_table else 0.88)))
    if region.is_table and height < 11:
        font_size = fit_font_size(region.translation, width, height)
        c.setFillColorRGB(0, 0, 0)
        c.setFont(FONT_NAME, font_size)
        y = bottom_y + max(0.3, (height - font_size) / 2)
        if region.align == "center":
            draw_translated_text(c, left + width / 2, y, region.translation, center=True)
        else:
            draw_translated_text(c, left, y, region.translation)
        return
    font_size, lines = fit_wrapped_font_size(region.translation, width, height + pad_y, max_size=max_size)
    c.setFillColorRGB(0, 0, 0)
    c.setFont(FONT_NAME, font_size)
    line_height = font_size * 1.18
    total_height = len(lines) * line_height
    y = bottom_y + max(0.2, (height - total_height) / 2) + total_height - font_size
    for line in lines:
        if region.align == "center":
            draw_translated_text(c, left + width / 2, y, line, center=True)
        else:
            draw_translated_text(c, left, y, line)
        y -= line_height


def overlay_padding(region: TextRegion) -> tuple[float, float]:
    if region.is_table:
        return 0.35, 0.25
    return 1.1, 0.8


def draw_pdf_vector_structure(
    c: canvas.Canvas,
    right_offset: float,
    page_height: float,
    page: pdfplumber.page.Page,
) -> None:
    c.setStrokeColorRGB(0, 0, 0)
    for line in getattr(page, "lines", []):
        x0 = float(line.get("x0", 0.0))
        x1 = float(line.get("x1", x0))
        y0 = float(line.get("y0", page_height - float(line.get("bottom", 0.0))))
        y1 = float(line.get("y1", page_height - float(line.get("top", 0.0))))
        width = abs(x1 - x0)
        height = abs(y1 - y0)
        if max(width, height) < 6.0:
            continue
        if width > 1.0 and height > 1.0:
            continue
        line_width = float(line.get("linewidth", 0.45) or 0.45)
        if line_width > 1.5:
            continue
        c.setLineWidth(max(0.25, min(line_width, 0.8)))
        c.line(right_offset + x0, y0, right_offset + x1, y1)

    for rect in getattr(page, "rects", []):
        x0 = float(rect.get("x0", 0.0))
        x1 = float(rect.get("x1", x0))
        y0 = float(rect.get("y0", page_height - float(rect.get("bottom", 0.0))))
        y1 = float(rect.get("y1", page_height - float(rect.get("top", 0.0))))
        width = x1 - x0
        height = y1 - y0
        if width < 4.0 or height < 4.0:
            continue
        line_width = float(rect.get("linewidth", 0.45) or 0.45)
        if line_width > 1.5:
            continue
        c.setLineWidth(max(0.25, min(line_width, 0.8)))
        c.rect(right_offset + x0, y0, width, height, stroke=1, fill=0)


def draw_table_grid(
    c: canvas.Canvas,
    right_offset: float,
    page_height: float,
    page: pdfplumber.page.Page,
) -> None:
    draw_pdf_vector_structure(c, right_offset, page_height, page)
    c.setStrokeColorRGB(0, 0, 0)
    c.setLineWidth(0.45)
    for table in page.find_tables():
        for cell in table.cells:
            if not cell:
                continue
            x0, top, x1, bottom = (float(v) for v in cell)
            if x1 <= x0 or bottom <= top:
                continue
            c.rect(right_offset + x0, page_height - bottom, x1 - x0, bottom - top, stroke=1, fill=0)


def collect_pdf_segments(pdf_path: Path) -> list[Segment]:
    segments: list[Segment] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            words = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False)
            for text, bbox in group_words_into_lines(words):
                segments.append(
                    Segment(
                        source_file=pdf_path.name,
                        page=page_index,
                        text=text,
                        translation=translate_line(text),
                        bbox=bbox,
                    )
                )
    return segments


def collect_pdf_segments_for_pages(pdf_path: Path, page_indices: list[int]) -> list[Segment]:
    segments: list[Segment] = []
    wanted = set(page_indices)
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_index in sorted(wanted):
            page = pdf.pages[page_index]
            words = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False)
            for line_words in group_words_into_word_lines(words):
                latin_words = [word for word in line_words if word_has_latin(word)]
                if not latin_words:
                    continue
                text = normalize_text(" ".join(str(word["text"]) for word in latin_words))
                if not text:
                    continue
                segments.append(
                    Segment(
                        source_file=pdf_path.name,
                        page=page_index + 1,
                        text=text,
                        translation=translate_line(text),
                        bbox=bbox_from_words(latin_words),
                    )
                )
    return segments


def line_alignment(line_words: list[dict], page_width: float) -> str:
    x0, _top, x1, _bottom = bbox_from_words(line_words)
    center = (x0 + x1) / 2
    text = line_words_text(line_words)
    left_margin = x0
    right_margin = page_width - x1
    line_width = x1 - x0
    balanced_margins = abs(left_margin - right_margin) <= max(12.0, page_width * 0.035)
    narrow_enough = line_width <= page_width * 0.62
    centered_on_page = abs(center - page_width / 2) <= page_width * 0.08
    has_lowercase = any(char.islower() for char in text)
    sentence_like = has_lowercase and (len(text) > 28 or text.endswith((".", ";", ",")))
    if centered_on_page and balanced_margins and narrow_enough and not sentence_like:
        return "center"
    return "left"


def collect_layout_regions(
    pdf_path: Path,
    page_indices: list[int],
    translation_cache: dict[str, str] | None = None,
) -> list[TextRegion]:
    source_regions = collect_layout_source_regions(pdf_path, page_indices)
    translated_regions: list[TextRegion] = []
    for region in source_regions:
        translation = translate_line(region.text, translation_cache)
        if not translation or translation == region.text:
            continue
        translated_regions.append(
            TextRegion(
                page=region.page,
                text=region.text,
                translation=translation,
                bbox=region.bbox,
                align=region.align,
                is_table=region.is_table,
            )
        )
    return translated_regions


def collect_layout_source_regions(
    pdf_path: Path,
    page_indices: list[int],
) -> list[TextRegion]:
    regions: list[TextRegion] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_index in page_indices:
            page = pdf.pages[page_index]
            table_bboxes = extract_table_bboxes(page)
            words = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False)
            for table in page.find_tables():
                for cell in table.cells:
                    if not cell:
                        continue
                    cell_bbox = tuple(float(v) for v in cell)
                    cell_words = words_inside_bbox(words, cell_bbox)
                    if not cell_words:
                        continue
                    for line_words in group_words_into_word_lines(cell_words, tolerance=2.2):
                        for chunk_words in split_words_by_horizontal_gap(line_words):
                            add_region_from_words(
                                regions,
                                page_index + 1,
                                chunk_words,
                                page.width,
                                is_table=True,
                                cell_bbox=cell_bbox,
                            )

            non_table_words = [
                word for word in words if not any(word_center_inside_bbox(word, table_bbox) for table_bbox in table_bboxes)
            ]
            for line_words in group_words_into_word_lines(non_table_words):
                add_region_from_words(regions, page_index + 1, line_words, page.width, is_table=False)
    return regions


def line_words_text(line_words: list[dict]) -> str:
    return normalize_text(" ".join(str(word["text"]) for word in line_words))


def is_template_break_line(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return True
    upper = normalized.upper()
    if re.fullmatch(r"\d+(?:\.\d+)?\s+[A-Z][A-Z0-9 /&().,'’:\-]+", upper):
        return True
    if re.match(r"^\d+\)", normalized):
        return True
    if re.match(r"^[A-Za-z][A-Za-z /&().,'’\-]{0,26}:", normalized):
        return True
    if upper in EXACT_TRANSLATIONS or upper in TERM_TRANSLATIONS:
        return True
    if re.fullmatch(r"PAGE\s+\d+\s+OF\s+\d+", upper):
        return True
    return False


def group_template_body_lines(line_groups: list[list[dict]]) -> list[list[dict]]:
    groups: list[list[dict]] = []
    current: list[list[dict]] = []
    previous_bbox: tuple[float, float, float, float] | None = None

    def flush_current() -> None:
        nonlocal current, previous_bbox
        if current:
            groups.append([word for line in current for word in line])
        current = []
        previous_bbox = None

    for line_words in line_groups:
        text = line_words_text(line_words)
        if not any(word_has_latin(word) for word in line_words):
            flush_current()
            continue
        bbox = bbox_from_words(line_words)
        if is_template_break_line(text):
            flush_current()
            groups.append(line_words)
            continue
        if previous_bbox is None:
            current = [line_words]
            previous_bbox = bbox
            continue
        previous_x0, _previous_top, _previous_x1, previous_bottom = previous_bbox
        x0, top, _x1, _bottom = bbox
        vertical_gap = top - previous_bottom
        same_paragraph_indent = abs(x0 - previous_x0) <= 38.0 or x0 >= previous_x0
        if vertical_gap <= 11.0 and same_paragraph_indent:
            current.append(line_words)
            previous_bbox = bbox
        else:
            flush_current()
            current = [line_words]
            previous_bbox = bbox
    flush_current()
    return groups


def collect_template_layout_source_regions(
    pdf_path: Path,
    page_indices: list[int],
) -> list[TextRegion]:
    regions: list[TextRegion] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_index in page_indices:
            page = pdf.pages[page_index]
            table_bboxes = extract_table_bboxes(page)
            words = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False)
            for table in page.find_tables():
                for cell in table.cells:
                    if not cell:
                        continue
                    cell_bbox = tuple(float(v) for v in cell)
                    cell_words = words_inside_bbox(words, cell_bbox)
                    if not cell_words:
                        continue
                    for line_words in group_words_into_word_lines(cell_words, tolerance=2.2):
                        for chunk_words in split_words_by_horizontal_gap(line_words):
                            add_region_from_words(
                                regions,
                                page_index + 1,
                                chunk_words,
                                page.width,
                                is_table=True,
                                cell_bbox=cell_bbox,
                            )

            non_table_words = [
                word for word in words if not any(word_center_inside_bbox(word, table_bbox) for table_bbox in table_bboxes)
            ]
            line_groups = group_words_into_word_lines(non_table_words)
            for grouped_words in line_groups:
                add_region_from_words(
                    regions,
                    page_index + 1,
                    grouped_words,
                    page.width,
                    is_table=False,
                )
    return regions


def collect_template_layout_regions(
    pdf_path: Path,
    page_indices: list[int],
    translation_cache: dict[str, str] | None = None,
) -> list[TextRegion]:
    source_regions = collect_template_layout_source_regions(pdf_path, page_indices)
    translated_regions: list[TextRegion] = []
    for region in source_regions:
        translation = translate_line(region.text, translation_cache)
        if not translation or translation == region.text:
            continue
        translated_regions.append(
            TextRegion(
                page=region.page,
                text=region.text,
                translation=translation,
                bbox=region.bbox,
                align=region.align,
                is_table=region.is_table,
            )
        )
    return translated_regions


def collect_table_repaint_layout_regions(
    pdf_path: Path,
    page_indices: list[int],
    translation_cache: dict[str, str] | None = None,
) -> list[TextRegion]:
    regions: list[TextRegion] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_index in page_indices:
            page = pdf.pages[page_index]
            table_bboxes = extract_table_bboxes(page)
            words = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False)
            for table in page.find_tables():
                for cell in table.cells:
                    if not cell:
                        continue
                    cell_bbox = tuple(float(v) for v in cell)
                    cell_words = words_inside_bbox(words, cell_bbox)
                    if not cell_words:
                        continue
                    for line_words in group_words_into_word_lines(cell_words, tolerance=2.2):
                        for chunk_words in split_table_repaint_line_words(line_words, cell_bbox):
                            add_repaint_region_from_words(
                                regions,
                                page_index + 1,
                                chunk_words,
                                page.width,
                                is_table=True,
                                cell_bbox=cell_bbox,
                                translation_cache=translation_cache,
                            )

            non_table_source_regions: list[TextRegion] = []
            non_table_words = [
                word for word in words if not any(word_center_inside_bbox(word, table_bbox) for table_bbox in table_bboxes)
            ]
            for line_words in group_words_into_word_lines(non_table_words):
                add_region_from_words(non_table_source_regions, page_index + 1, line_words, page.width, is_table=False)
            for region in non_table_source_regions:
                translation = translate_line(region.text, translation_cache)
                if not translation or translation == region.text:
                    continue
                regions.append(
                    TextRegion(
                        page=region.page,
                        text=region.text,
                        translation=translation,
                        bbox=region.bbox,
                        align=region.align,
                        is_table=False,
                    )
                )
    return regions


def draw_text_region_preserve_template(
    c: canvas.Canvas,
    right_offset: float,
    page_height: float,
    region: TextRegion,
) -> None:
    if region.is_table:
        draw_text_region(c, right_offset, page_height, region)
        return
    x0, top, x1, bottom = region.bbox
    left = right_offset + x0
    bottom_y = page_height - bottom
    width = max(8.0, x1 - x0)
    height = max(5.0, bottom - top)
    pad_x, pad_y = overlay_padding(region)
    c.setFillColorRGB(1, 1, 1)
    c.rect(left - pad_x, bottom_y - pad_y, width + pad_x * 2, height + pad_y * 2, stroke=0, fill=1)
    font_size, lines = fit_wrapped_font_size(region.translation, width, height + pad_y * 2, max_size=min(8.8, max(5.0, height * 0.78)))
    c.setFillColorRGB(0, 0, 0)
    c.setFont(FONT_NAME, font_size)
    line_height = font_size * 1.16
    y = page_height - top - font_size
    min_y = bottom_y + 0.4
    restore_underline = normalize_text(region.text).upper() in UNDERLINED_SECTION_HEADINGS
    for line_index, line in enumerate(lines):
        if y < min_y:
            break
        if region.align == "center":
            draw_translated_text(c, left + width / 2, y, line, center=True)
            underline_width = min(width, text_width(line, font_size))
            underline_x0 = left + (width - underline_width) / 2
        else:
            draw_translated_text(c, left, y, line)
            underline_width = min(width, text_width(line, font_size))
            underline_x0 = left
        if restore_underline and line_index == 0:
            c.setStrokeColorRGB(0, 0, 0)
            c.setLineWidth(0.45)
            underline_y = y - max(1.0, font_size * 0.18)
            c.line(underline_x0, underline_y, underline_x0 + underline_width, underline_y)
        y -= line_height


def generate_template_layout_sample_pdf(
    input_pdf: Path,
    output_pdf: Path,
    page_indices: list[int],
    translation_cache: dict[str, str] | None = None,
) -> list[TextRegion]:
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    regions = collect_template_layout_regions(input_pdf, page_indices, translation_cache=translation_cache)
    regions_by_page: dict[int, list[TextRegion]] = {}
    for region in regions:
        regions_by_page.setdefault(region.page, []).append(region)

    reader = PdfReader(str(input_pdf))
    pdf_doc = pdfium.PdfDocument(str(input_pdf))
    output_canvas = canvas.Canvas(str(output_pdf))
    try:
        with pdfplumber.open(str(input_pdf)) as plumber_pdf:
            for page_index in page_indices:
                page = reader.pages[page_index]
                plumber_page = plumber_pdf.pages[page_index]
                width = float(page.mediabox.width)
                height = float(page.mediabox.height)
                output_canvas.setPageSize((width * 2 + DIVIDER_WIDTH, height))
                rendered_page = pdf_doc[page_index]
                rendered_image = rendered_page.render(scale=1.9).to_pil()
                image_buffer = BytesIO()
                rendered_image.save(image_buffer, format="PNG")
                image_buffer.seek(0)
                output_canvas.drawImage(ImageReader(image_buffer), 0, 0, width=width, height=height)
                draw_divider(output_canvas, width, height)
                image_buffer.seek(0)
                right_offset = width + DIVIDER_WIDTH
                output_canvas.drawImage(ImageReader(image_buffer), right_offset, 0, width=width, height=height)
                for region in regions_by_page.get(page_index + 1, []):
                    draw_text_region_preserve_template(output_canvas, right_offset, height, region)
                draw_table_grid(output_canvas, right_offset, height, plumber_page)
                output_canvas.showPage()
    finally:
        pdf_doc.close()
    output_canvas.save()
    return regions


def generate_table_repaint_sample_pdf(
    input_pdf: Path,
    output_pdf: Path,
    page_indices: list[int],
    translation_cache: dict[str, str] | None = None,
) -> list[TextRegion]:
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    regions = collect_table_repaint_layout_regions(input_pdf, page_indices, translation_cache=translation_cache)
    regions_by_page: dict[int, list[TextRegion]] = {}
    for region in regions:
        regions_by_page.setdefault(region.page, []).append(region)

    reader = PdfReader(str(input_pdf))
    pdf_doc = pdfium.PdfDocument(str(input_pdf))
    output_canvas = canvas.Canvas(str(output_pdf))
    try:
        with pdfplumber.open(str(input_pdf)) as plumber_pdf:
            for page_index in page_indices:
                page = reader.pages[page_index]
                plumber_page = plumber_pdf.pages[page_index]
                width = float(page.mediabox.width)
                height = float(page.mediabox.height)
                output_canvas.setPageSize((width * 2 + DIVIDER_WIDTH, height))
                rendered_page = pdf_doc[page_index]
                rendered_image = rendered_page.render(scale=1.9).to_pil()
                image_buffer = BytesIO()
                rendered_image.save(image_buffer, format="PNG")
                image_buffer.seek(0)
                output_canvas.drawImage(ImageReader(image_buffer), 0, 0, width=width, height=height)
                draw_divider(output_canvas, width, height)
                image_buffer.seek(0)
                right_offset = width + DIVIDER_WIDTH
                output_canvas.drawImage(ImageReader(image_buffer), right_offset, 0, width=width, height=height)
                for region in regions_by_page.get(page_index + 1, []):
                    draw_text_region_preserve_template(output_canvas, right_offset, height, region)
                draw_table_grid(output_canvas, right_offset, height, plumber_page)
                output_canvas.showPage()
    finally:
        pdf_doc.close()
    output_canvas.save()
    return regions


def generate_auto_layout_sample_pdf(
    input_pdf: Path,
    output_pdf: Path,
    page_indices: list[int],
    translation_cache: dict[str, str] | None = None,
) -> list[TextRegion]:
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    regions = collect_auto_layout_regions(input_pdf, page_indices, translation_cache=translation_cache)
    regions_by_page: dict[int, list[TextRegion]] = {}
    for region in regions:
        regions_by_page.setdefault(region.page, []).append(region)

    reader = PdfReader(str(input_pdf))
    pdf_doc = pdfium.PdfDocument(str(input_pdf))
    output_canvas = canvas.Canvas(str(output_pdf))
    try:
        with pdfplumber.open(str(input_pdf)) as plumber_pdf:
            for page_index in page_indices:
                page = reader.pages[page_index]
                plumber_page = plumber_pdf.pages[page_index]
                width = float(page.mediabox.width)
                height = float(page.mediabox.height)
                output_canvas.setPageSize((width * 2 + DIVIDER_WIDTH, height))
                rendered_page = pdf_doc[page_index]
                rendered_image = rendered_page.render(scale=1.9).to_pil()
                image_buffer = BytesIO()
                rendered_image.save(image_buffer, format="PNG")
                image_buffer.seek(0)
                output_canvas.drawImage(ImageReader(image_buffer), 0, 0, width=width, height=height)
                draw_divider(output_canvas, width, height)
                image_buffer.seek(0)
                right_offset = width + DIVIDER_WIDTH
                output_canvas.drawImage(ImageReader(image_buffer), right_offset, 0, width=width, height=height)
                for region in regions_by_page.get(page_index + 1, []):
                    draw_text_region_preserve_template(output_canvas, right_offset, height, region)
                draw_table_grid(output_canvas, right_offset, height, plumber_page)
                output_canvas.showPage()
    finally:
        pdf_doc.close()
    output_canvas.save()
    return regions


def generate_layout_sample_pdf(
    input_pdf: Path,
    output_pdf: Path,
    page_indices: list[int],
    translation_cache: dict[str, str] | None = None,
) -> list[TextRegion]:
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    regions = collect_layout_regions(input_pdf, page_indices, translation_cache=translation_cache)
    regions_by_page: dict[int, list[TextRegion]] = {}
    for region in regions:
        regions_by_page.setdefault(region.page, []).append(region)

    reader = PdfReader(str(input_pdf))
    pdf_doc = pdfium.PdfDocument(str(input_pdf))
    output_canvas = canvas.Canvas(str(output_pdf))
    try:
        with pdfplumber.open(str(input_pdf)) as plumber_pdf:
            for page_index in page_indices:
                page = reader.pages[page_index]
                plumber_page = plumber_pdf.pages[page_index]
                width = float(page.mediabox.width)
                height = float(page.mediabox.height)
                output_canvas.setPageSize((width * 2 + DIVIDER_WIDTH, height))
                rendered_page = pdf_doc[page_index]
                rendered_image = rendered_page.render(scale=1.9).to_pil()
                image_buffer = BytesIO()
                rendered_image.save(image_buffer, format="PNG")
                image_buffer.seek(0)
                output_canvas.drawImage(ImageReader(image_buffer), 0, 0, width=width, height=height)
                draw_divider(output_canvas, width, height)
                image_buffer.seek(0)
                right_offset = width + DIVIDER_WIDTH
                output_canvas.drawImage(ImageReader(image_buffer), right_offset, 0, width=width, height=height)
                for region in regions_by_page.get(page_index + 1, []):
                    draw_text_region(output_canvas, right_offset, height, region)
                draw_table_grid(output_canvas, right_offset, height, plumber_page)
                output_canvas.showPage()
    finally:
        pdf_doc.close()
    output_canvas.save()
    return regions


def draw_divider(c: canvas.Canvas, page_width: float, page_height: float) -> None:
    c.setFillColorRGB(0.86, 0.86, 0.86)
    c.rect(page_width, 0, DIVIDER_WIDTH, page_height, stroke=0, fill=1)
    c.setStrokeColorRGB(0.25, 0.25, 0.25)
    c.setLineWidth(1.2)
    c.line(page_width + 0.5, 0, page_width + 0.5, page_height)
    c.line(page_width + DIVIDER_WIDTH - 0.5, 0, page_width + DIVIDER_WIDTH - 0.5, page_height)


def generate_side_by_side_pdf(
    input_pdf: Path,
    output_pdf: Path,
    translation_cache: dict[str, str] | None = None,
) -> list[Segment]:
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    segments = collect_pdf_segments(input_pdf)
    segments_by_page: dict[int, list[Segment]] = {}
    for segment in segments:
        segment.translation = translate_line(segment.text, translation_cache)
        segments_by_page.setdefault(segment.page, []).append(segment)

    reader = PdfReader(str(input_pdf))
    pdf_doc = pdfium.PdfDocument(str(input_pdf))
    output_canvas = canvas.Canvas(str(output_pdf))
    try:
        for page_index, page in enumerate(reader.pages, start=1):
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            output_canvas.setPageSize((width * 2 + DIVIDER_WIDTH, height))
            rendered_page = pdf_doc[page_index - 1]
            rendered_image = rendered_page.render(scale=1.7).to_pil()
            image_buffer = BytesIO()
            rendered_image.save(image_buffer, format="PNG")
            image_buffer.seek(0)
            output_canvas.drawImage(ImageReader(image_buffer), 0, 0, width=width, height=height)
            draw_divider(output_canvas, width, height)
            image_buffer.seek(0)
            right_offset = width + DIVIDER_WIDTH
            output_canvas.drawImage(ImageReader(image_buffer), right_offset, 0, width=width, height=height)
            for segment in segments_by_page.get(page_index, []):
                draw_translated_line(output_canvas, right_offset, height, segment.translation, segment.bbox)
            output_canvas.showPage()
    finally:
        pdf_doc.close()
    output_canvas.save()
    return segments


def write_pdf_translation_text(segments: Iterable[Segment], output_txt: Path) -> None:
    output_txt.parent.mkdir(parents=True, exist_ok=True)
    current_key: tuple[str, int] | None = None
    lines = []
    for segment in segments:
        key = (segment.source_file, segment.page)
        if key != current_key:
            if lines:
                lines.append("")
            lines.append(f"【{segment.source_file} - 第 {segment.page} 页】")
            current_key = key
        lines.append(f"原文：{segment.text}")
        lines.append(f"中文：{segment.translation}")
    output_txt.write_text("\n".join(lines), encoding="utf-8")


def docx_paragraph_records(document) -> list[tuple[object, str]]:
    records: list[tuple[object, str]] = []
    seen_paragraphs: set[int] = set()

    def add_paragraph(paragraph, location: str) -> None:
        paragraph_key = id(paragraph._element)
        if paragraph_key in seen_paragraphs:
            return
        seen_paragraphs.add(paragraph_key)
        records.append((paragraph, location))

    def add_table(table, location: str) -> None:
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                cell_location = f"{location}/row:{row_index}/cell:{cell_index}"
                for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                    add_paragraph(paragraph, f"{cell_location}/p:{paragraph_index}")
                for nested_index, nested_table in enumerate(cell.tables, start=1):
                    add_table(nested_table, f"{cell_location}/table:{nested_index}")

    def add_part(part, location: str) -> None:
        for paragraph_index, paragraph in enumerate(part.paragraphs, start=1):
            add_paragraph(paragraph, f"{location}/p:{paragraph_index}")
        for table_index, table in enumerate(part.tables, start=1):
            add_table(table, f"{location}/table:{table_index}")

    for index, paragraph in enumerate(document.paragraphs, start=1):
        add_paragraph(paragraph, f"paragraph:{index}")
    for table_index, table in enumerate(document.tables, start=1):
        add_table(table, f"table:{table_index}")

    seen_parts = set()
    section_part_names = (
        "header",
        "first_page_header",
        "even_page_header",
        "footer",
        "first_page_footer",
        "even_page_footer",
    )
    for section_index, section in enumerate(document.sections, start=1):
        for part_name in section_part_names:
            part = getattr(section, part_name)
            part_key = id(part._element)
            if part_key in seen_parts:
                continue
            seen_parts.add(part_key)
            add_part(part, f"section:{section_index}/{part_name}")
    return records


def collect_docx_texts(input_docx: Path) -> list[str]:
    document = Document(str(input_docx))
    return [normalize_text(paragraph.text) for paragraph, _ in docx_paragraph_records(document) if normalize_text(paragraph.text)]


def translate_docx(
    input_docx: Path,
    output_docx: Path,
    output_txt: Path,
    translation_cache: dict[str, str] | None = None,
) -> list[dict]:
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    output_txt.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(input_docx))
    segments = []

    def set_paragraph_text(paragraph, translated: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = translated
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(translated)

    def replace_paragraph(paragraph, location: str) -> None:
        original = normalize_text(paragraph.text)
        if not original:
            return
        translated, translation_source = translate_office_text(original, translation_cache)
        set_paragraph_text(paragraph, translated)
        segments.append(
            {
                "location": location,
                "original": original,
                "translation": translated,
                "translation_source": translation_source,
            }
        )

    for paragraph, location in docx_paragraph_records(document):
        replace_paragraph(paragraph, location)

    document.save(str(output_docx))
    txt_lines = []
    for segment in segments:
        txt_lines.append(f"【{segment['location']}】")
        txt_lines.append(f"原文：{segment['original']}")
        txt_lines.append(f"中文：{segment['translation']}")
    output_txt.write_text("\n".join(txt_lines), encoding="utf-8")
    return segments


def detect_docx_high_risk_objects(input_docx: Path) -> list[str]:
    try:
        with zipfile.ZipFile(input_docx, "r") as archive:
            xml_parts = {
                name: archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.startswith("word/") and (name.endswith(".xml") or name.endswith(".rels"))
            }
    except (OSError, zipfile.BadZipFile):
        return ["DOCX 结构检测失败：文件无法作为标准 DOCX ZIP 读取，需人工复核"]

    combined_xml = "\n".join(xml_parts.values())
    header_footer_xml = "\n".join(
        content
        for name, content in xml_parts.items()
        if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
    )
    warnings = []
    if any(marker in combined_xml for marker in ("w:txbxContent", "v:textbox", "wps:txbx", "txbxContent")):
        warnings.append("DOCX 高风险对象：检测到文本框，python-docx 不能可靠翻译其中内容，需人工复核")
    if "word/comments.xml" in xml_parts or "commentRangeStart" in combined_xml or "commentReference" in combined_xml:
        warnings.append("DOCX 高风险对象：检测到批注，批注内容和批注锚点需人工复核")
    if re.search(r"<w:(?:ins|del|moveFrom|moveTo)\b", combined_xml):
        warnings.append("DOCX 高风险对象：检测到修订痕迹，插入/删除/移动内容需人工复核")
    if any(marker in combined_xml for marker in ("<dgm:", "smartArt", "SmartArt", "/diagramData", "/diagramDrawing")):
        warnings.append("DOCX 高风险对象：检测到 SmartArt/Diagram，图形内文字需人工复核")
    if "wp:anchor" in header_footer_xml:
        warnings.append("DOCX 高风险对象：检测到页眉页脚浮动对象，浮动对象内文字需人工复核")
    return warnings


def inspect_xlsx_warnings(workbook, input_xlsx: Path) -> list[str]:
    warnings: list[str] = []
    if input_xlsx.suffix.lower() == ".xlsm":
        warnings.append("宏工作簿: 已保留 .xlsm 宏容器，宏内容需人工复核")
    security = getattr(workbook, "security", None)
    if security and (getattr(security, "lockStructure", False) or getattr(security, "lockWindows", False)):
        warnings.append("工作簿结构保护: 已保留，翻译后需人工确认保护状态")
    if getattr(workbook, "_external_links", None):
        warnings.append("外部链接: 工作簿包含外部链接，翻译未改写链接目标")
    for sheet in workbook.worksheets:
        if sheet.sheet_state != "visible":
            warnings.append(f"隐藏 Sheet: {sheet.title}")
        if getattr(sheet.auto_filter, "ref", None):
            warnings.append(f"筛选: {sheet.title}!{sheet.auto_filter.ref}")
        if getattr(sheet.protection, "sheet", False):
            warnings.append(f"Sheet 保护: {sheet.title}")
        for row in sheet.iter_rows():
            for cell in row:
                if cell.comment is not None:
                    warnings.append(f"批注: {sheet.title}!{cell.coordinate}")
        images = getattr(sheet, "_images", [])
        if images:
            warnings.append(f"图片: {sheet.title} 包含 {len(images)} 个图片对象，图片文字未翻译")
        charts = getattr(sheet, "_charts", [])
        if charts:
            warnings.append(f"图表: {sheet.title} 包含 {len(charts)} 个图表对象，图表内部文字需人工复核")
        tables = getattr(sheet, "tables", {})
        if tables:
            warnings.append(f"表格对象: {sheet.title} 包含 {len(tables)} 个结构化表格，筛选/样式需人工复核")
    return warnings


def copy_xlsx_without_custom_properties(input_xlsx: Path, output_xlsx: Path) -> None:
    output_xlsx.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(input_xlsx, "r") as source_archive, zipfile.ZipFile(output_xlsx, "w", zipfile.ZIP_DEFLATED) as target_archive:
        for item in source_archive.infolist():
            if item.filename == "docProps/custom.xml":
                continue
            target_archive.writestr(item, source_archive.read(item.filename))


def load_xlsx_for_translation(input_xlsx: Path):
    keep_vba = input_xlsx.suffix.lower() == ".xlsm"
    try:
        return load_workbook(str(input_xlsx), keep_vba=keep_vba), []
    except TypeError as exc:
        message = str(exc)
        if "StringProperty" not in message or ".name should be" not in message:
            raise
        with tempfile.TemporaryDirectory(prefix="rfq_xlsx_clean_") as tmp:
            cleaned_xlsx = Path(tmp) / input_xlsx.name
            copy_xlsx_without_custom_properties(input_xlsx, cleaned_xlsx)
            workbook = load_workbook(str(cleaned_xlsx), keep_vba=keep_vba)
        return workbook, ["XLSX 自定义属性损坏，已在临时副本中移除 docProps/custom.xml 后继续翻译；原文件未修改"]


def close_xlsx_workbook(workbook) -> None:
    workbook.close()
    vba_archive = getattr(workbook, "vba_archive", None)
    if vba_archive is not None:
        try:
            vba_archive.close()
        finally:
            workbook.vba_archive = None


def collect_xlsx_texts(input_xlsx: Path) -> list[str]:
    workbook, _ = load_xlsx_for_translation(input_xlsx)
    try:
        return [
            normalize_text(cell.value)
            for sheet in workbook.worksheets
            for row in sheet.iter_rows()
            for cell in row
            if isinstance(cell.value, str)
            and normalize_text(cell.value)
            and not normalize_text(cell.value).startswith("=")
        ]
    finally:
        close_xlsx_workbook(workbook)


def translate_xlsx(
    input_xlsx: Path,
    output_xlsx: Path,
    output_txt: Path,
    translation_cache: dict[str, str] | None = None,
) -> tuple[list[dict], list[str]]:
    output_xlsx.parent.mkdir(parents=True, exist_ok=True)
    output_txt.parent.mkdir(parents=True, exist_ok=True)
    workbook, load_warnings = load_xlsx_for_translation(input_xlsx)
    warnings = load_warnings + inspect_xlsx_warnings(workbook, input_xlsx)
    segments = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                value = cell.value
                if not isinstance(value, str):
                    continue
                original = normalize_text(value)
                if not original or original.startswith("="):
                    continue
                translated, translation_source = translate_office_text(original, translation_cache)
                if translated == value:
                    if not office_requires_model_translation(original):
                        continue
                else:
                    cell.value = translated
                segments.append(
                    {
                        "location": f"sheet:{sheet.title}/cell:{cell.coordinate}",
                        "sheet": sheet.title,
                        "cell": cell.coordinate,
                        "original": original,
                        "translation": translated,
                        "translation_source": translation_source,
                    }
                )
    workbook.save(str(output_xlsx))
    close_xlsx_workbook(workbook)
    txt_lines = []
    for segment in segments:
        txt_lines.append(f"【{segment['location']}】")
        txt_lines.append(f"原文：{segment['original']}")
        txt_lines.append(f"中文：{segment['translation']}")
    output_txt.write_text("\n".join(txt_lines), encoding="utf-8")
    return segments, warnings


def build_manifest_entry(
    source_path: Path,
    output_pdf: Path | None,
    output_txt: Path | None,
    status: str,
    page_count: int | None,
    method: str,
    risks: list[str],
    output_docx: Path | None = None,
) -> dict:
    outputs = {}
    if output_pdf:
        outputs["pdf"] = str(output_pdf)
    if output_txt:
        outputs["txt"] = str(output_txt)
    if output_docx:
        outputs["docx"] = str(output_docx)
    return {
        "source_file": source_path.name,
        "source_path": str(source_path),
        "status": status,
        "page_count": page_count,
        "method": method,
        "outputs": outputs,
        "risks": risks,
    }


def detect_pdf_language(segments: list[Segment]) -> str:
    joined = "\n".join(segment.text for segment in segments[:80])
    cyrillic = len(re.findall(r"[\u0400-\u04ff]", joined))
    latin = len(re.findall(r"[A-Za-z]", joined))
    if cyrillic and latin:
        return "English/Russian mixed"
    if cyrillic:
        return "Russian"
    if latin:
        return "English"
    return "Unknown"


def normalize_project_mode(mode: str) -> str:
    normalized = normalize_text(mode or "平衡")
    lookup_key = normalized.lower()
    if lookup_key not in SUPPORTED_PROJECT_MODES:
        supported = "、".join(sorted(SUPPORTED_PROJECT_MODES))
        raise ValueError(f"不支持的处理模式：{mode!r}；当前支持：{supported}")
    return SUPPORTED_PROJECT_MODES[lookup_key]


def normalize_pdf_engine(pdf_engine: str | None = None) -> str:
    requested = pdf_engine or PDF_ENGINE_LEGACY
    key = str(requested).strip().lower().replace("-", "_")
    if key in {"legacy", "old", "b_legacy"}:
        return PDF_ENGINE_LEGACY
    if key in {"pdfmathtranslate_next", "pdf2zh_next"}:
        return PDF_ENGINE_PDFMATHTRANSLATE_NEXT
    supported = "、".join(sorted(SUPPORTED_PDF_ENGINES))
    raise ValueError(f"不支持的 PDF 翻译引擎：{pdf_engine!r}；当前支持：{supported}")


def user_environment_value(name: str, default: str | None = None) -> str | None:
    value = os.environ.get(name)
    if value is not None and value.strip():
        return value
    if os.name == "nt":
        try:
            import winreg

            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, "Environment") as key:
                registry_value, _ = winreg.QueryValueEx(key, name)
            if registry_value is not None and str(registry_value).strip():
                return str(registry_value)
        except (FileNotFoundError, OSError):
            pass
    return default


def split_private_glossary_paths(raw: str | None) -> list[str]:
    if not raw:
        return []
    return [
        value.strip().strip("\"'")
        for value in re.split(r"[,;\r\n]+", raw)
        if value.strip().strip("\"'")
    ]


def load_private_glossary_terms() -> tuple[dict[str, str], dict]:
    raw_paths = split_private_glossary_paths(
        user_environment_value(PRIVATE_GLOSSARY_ENV)
    )
    if not raw_paths:
        return {}, {
            "configured": False,
            "file_count": 0,
            "entry_count": 0,
            "signature": None,
        }

    digest = hashlib.sha256()
    terms: dict[str, str] = {}
    seen_paths: set[str] = set()
    file_count = 0
    for raw_path in raw_paths:
        try:
            path = Path(os.path.expandvars(raw_path)).expanduser().resolve()
        except OSError:
            raise FileNotFoundError("私有术语文件不存在或不可读取") from None
        path_key = str(path).casefold()
        if path_key in seen_paths:
            continue
        seen_paths.add(path_key)
        if not path.is_file():
            raise FileNotFoundError("私有术语文件不存在或不可读取")
        if path.suffix.lower() not in {".csv", ".json"}:
            raise ValueError("私有术语文件仅支持 CSV 或 JSON")
        try:
            content = path.read_bytes()
            if path.suffix.lower() == ".json":
                rows = json.loads(content.decode("utf-8-sig"))
                if not isinstance(rows, list):
                    raise ValueError
            else:
                decoded = content.decode("utf-8-sig")
                rows = list(csv.DictReader(decoded.splitlines()))
        except (OSError, UnicodeError, json.JSONDecodeError, ValueError, csv.Error):
            raise ValueError("私有术语文件无法解析，请检查 CSV/JSON 格式") from None
        digest.update(len(content).to_bytes(8, "big"))
        digest.update(content)
        file_count += 1
        for row in rows:
            if not isinstance(row, dict):
                raise ValueError("私有术语文件的词条必须是对象或表格行")
            source = normalize_text(str(row.get("source", "")))
            target = normalize_text(str(row.get("target", "")))
            if not source or not target:
                raise ValueError("私有术语文件的 source 和 target 不能为空")
            terms[source.casefold()] = target
    return terms, {
        "configured": True,
        "file_count": file_count,
        "entry_count": len(terms),
        "signature": digest.hexdigest()[:16],
    }


def positive_env_int(name: str, default: int, maximum: int | None = None) -> int:
    raw = user_environment_value(name)
    if raw is None or not raw.strip():
        return default
    try:
        value = int(raw)
    except ValueError as exc:
        raise ValueError(f"环境变量 {name} 必须是正整数") from exc
    if value <= 0:
        raise ValueError(f"环境变量 {name} 必须是正整数")
    if maximum is not None and value > maximum:
        return maximum
    return value


def default_pdf_runtime_root() -> Path:
    local_app_data = os.environ.get("LOCALAPPDATA") or str(
        Path.home() / "AppData" / "Local"
    )
    return Path(local_app_data) / "RFQTranslationTool" / "BRuntime"


def resolve_pdf_runtime_python() -> Path:
    configured = user_environment_value("B_PDF_TRANSLATION_PYTHON")
    if configured:
        return Path(configured).expanduser().resolve()
    runtime_root = Path(
        user_environment_value(
            "B_PDF_TRANSLATION_RUNTIME_DIR",
            str(default_pdf_runtime_root()),
        )
    ).expanduser()
    return runtime_root / ".venv" / "Scripts" / "python.exe"


def validated_model_base_url(raw_value: str) -> str:
    value = str(raw_value or "").strip().rstrip("/")
    parsed = urllib.parse.urlsplit(value)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise ValueError("模型服务地址必须是完整的 http/https URL")
    if parsed.username or parsed.password or parsed.query or parsed.fragment:
        raise ValueError("模型服务地址不得包含账号、密码、查询参数或片段")
    return value


def office_model_provider_config() -> dict:
    base_url = validated_model_base_url(
        user_environment_value(
            "VECTOR_ENGINE_BASE_URL",
            DEFAULT_VECTOR_ENGINE_BASE_URL,
        )
    )
    model = user_environment_value(
        "VECTOR_ENGINE_MODEL",
        DEFAULT_VECTOR_ENGINE_MODEL,
    ).strip()
    if not model:
        raise ValueError("VECTOR_ENGINE_MODEL 不能为空")
    return {
        "service": "openaicompatible",
        "base_url": base_url,
        "model": model,
        "api_key_env": VECTOR_ENGINE_API_KEY_ENV,
        "api_key_configured": bool(user_environment_value(VECTOR_ENGINE_API_KEY_ENV)),
    }


def office_config_signature() -> str:
    provider = office_model_provider_config()
    _private_terms, private_glossary_summary = load_private_glossary_terms()
    payload = {
        "service": provider["service"],
        "base_url": provider["base_url"],
        "model": provider["model"],
        "target_language": OFFICE_TRANSLATION_TARGET_LANGUAGE,
        "prompt_contract_version": OFFICE_PROMPT_CONTRACT_VERSION,
        "protection_contract_version": OFFICE_PROTECTION_CONTRACT_VERSION,
        "office_batch_isolation_version": (
            OFFICE_BATCH_FAILURE_ISOLATION_CONTRACT_VERSION
        ),
        "batch_size": positive_env_int(
            "B_OFFICE_TRANSLATION_BATCH_SIZE", 16, maximum=32
        ),
        "batch_max_chars": positive_env_int(
            "B_OFFICE_TRANSLATION_BATCH_MAX_CHARS", 6000, maximum=12000
        ),
        "timeout": positive_env_int(
            "B_OFFICE_TRANSLATION_TIMEOUT", 90, maximum=300
        ),
        "private_glossary_signature": private_glossary_summary["signature"],
    }
    digest = hashlib.sha256(
        json.dumps(payload, sort_keys=True, ensure_ascii=False).encode("utf-8")
    ).hexdigest()[:16]
    return (
        f"office:{provider['model']}:"
        f"{OFFICE_BATCH_FAILURE_ISOLATION_CONTRACT_VERSION}:{digest}"
    )


def pdfmathtranslate_provider_config() -> dict:
    service = user_environment_value(
        "B_PDF_TRANSLATION_SERVICE",
        DEFAULT_PDF_TRANSLATION_SERVICE,
    ).strip().lower()
    if service not in {
        "openaicompatible",
        "openaicompatbatch",
        "siliconflowfree",
    }:
        raise ValueError(
            "B_PDF_TRANSLATION_SERVICE 当前仅支持 "
            "openaicompatible、openaicompatbatch 或 siliconflowfree"
        )
    batched_service = service == "openaicompatbatch"
    config = {
        "service": service,
        "profile": user_environment_value("B_PDF_TRANSLATION_PROFILE", "fast").strip().lower(),
        "qps": positive_env_int(
            "B_PDF_TRANSLATION_QPS", 12 if batched_service else 4, maximum=20
        ),
        "pool_max_workers": positive_env_int(
            "B_PDF_TRANSLATION_WORKERS", 12 if batched_service else 4, maximum=16
        ),
        "thinking": (
            "disabled"
            if user_environment_value("B_PDF_TRANSLATION_DISABLE_THINKING", "1").strip().lower()
            not in {"0", "false", "no", "off"}
            else "enabled"
        ),
    }
    if config["profile"] not in {"quality", "balanced", "fast"}:
        raise ValueError("B_PDF_TRANSLATION_PROFILE 当前仅支持 quality、balanced 或 fast")
    if service in {"openaicompatible", "openaicompatbatch"}:
        config.update(
            {
                "base_url": validated_model_base_url(
                    user_environment_value(
                        "VECTOR_ENGINE_BASE_URL",
                        DEFAULT_VECTOR_ENGINE_BASE_URL,
                    )
                ),
                "model": user_environment_value(
                    "VECTOR_ENGINE_MODEL",
                    DEFAULT_VECTOR_ENGINE_MODEL,
                ).strip(),
                "api_key_env": VECTOR_ENGINE_API_KEY_ENV,
                "api_key_configured": bool(user_environment_value(VECTOR_ENGINE_API_KEY_ENV)),
            }
        )
    else:
        config.update(
            {
                "base_url": None,
                "model": "SiliconFlowFree",
                "api_key_env": None,
                "api_key_configured": True,
            }
        )
    if batched_service:
        config.update(
            {
                "batch_size": positive_env_int(
                    "B_PDF_TRANSLATION_BATCH_SIZE", 20, maximum=32
                ),
                "batch_max_chars": positive_env_int(
                    "B_PDF_TRANSLATION_BATCH_MAX_CHARS", 6000, maximum=12000
                ),
                "batch_flush_ms": positive_env_int(
                    "B_PDF_TRANSLATION_BATCH_FLUSH_MS", 100, maximum=2000
                ),
                "batch_timeout": positive_env_int(
                    "B_PDF_TRANSLATION_BATCH_TIMEOUT", 90, maximum=300
                ),
                "batch_max_retries": positive_env_int(
                    "B_PDF_TRANSLATION_BATCH_RETRIES", 4, maximum=8
                ),
                "batch_request_workers": positive_env_int(
                    "B_PDF_TRANSLATION_BATCH_REQUEST_WORKERS", 4, maximum=8
                ),
                "repair_model": user_environment_value(
                    "B_PDF_TRANSLATION_REPAIR_MODEL",
                    DEFAULT_VECTOR_ENGINE_REPAIR_MODEL,
                ).strip(),
                "babeldoc_batch_token_limit": positive_env_int(
                    "B_PDF_TRANSLATION_BABELDOC_BATCH_TOKEN_LIMIT", 1600, maximum=8000
                ),
                "babeldoc_batch_count_limit": positive_env_int(
                    "B_PDF_TRANSLATION_BABELDOC_BATCH_COUNT_LIMIT", 40, maximum=100
                ),
                "doclayout_image_size": positive_env_int(
                    "B_PDF_TRANSLATION_DOCLAYOUT_IMAGE_SIZE", 800, maximum=1280
                ),
                "translate_cyrillic_formula_text": (
                    user_environment_value(
                        "B_PDF_TRANSLATION_TRANSLATE_CYRILLIC_FORMULA_TEXT",
                        "1",
                    ).strip().lower()
                    not in {"0", "false", "no", "off"}
                ),
            }
        )
    else:
        config.update(
            {
                "batch_size": None,
                "batch_max_chars": None,
                "batch_flush_ms": None,
                "batch_timeout": None,
                "batch_max_retries": None,
                "batch_request_workers": None,
                "repair_model": None,
                "babeldoc_batch_token_limit": None,
                "babeldoc_batch_count_limit": None,
                "doclayout_image_size": None,
                "translate_cyrillic_formula_text": None,
            }
        )
    return config


def project_config_signature(pdf_engine: str) -> str:
    office_signature = office_config_signature()
    if pdf_engine != PDF_ENGINE_PDFMATHTRANSLATE_NEXT:
        return f"{PROJECT_TRANSLATION_CONFIG_SIGNATURE}:{pdf_engine}:{office_signature}"
    provider = pdfmathtranslate_provider_config()
    component_versions = runtime_component_versions()
    glossary_path = PDF_RUNTIME_SOURCE_DIR / "rfq_default_glossary.json"
    glossary_digest = (
        hashlib.sha256(glossary_path.read_bytes()).hexdigest()[:12]
        if glossary_path.is_file()
        else "missing"
    )
    provider_payload = "|".join(
        [
            str(provider["service"]),
            str(provider.get("base_url") or ""),
            str(provider.get("model") or ""),
            str(provider["profile"]),
            str(provider["qps"]),
            str(provider["pool_max_workers"]),
            str(provider["thinking"]),
            str(provider.get("batch_size") or ""),
            str(provider.get("batch_max_chars") or ""),
            str(provider.get("batch_flush_ms") or ""),
            str(provider.get("batch_timeout") or ""),
            str(provider.get("batch_max_retries") or ""),
            str(provider.get("batch_request_workers") or ""),
            str(provider.get("repair_model") or ""),
            str(provider.get("babeldoc_batch_token_limit") or ""),
            str(provider.get("babeldoc_batch_count_limit") or ""),
            str(provider.get("doclayout_image_size") or ""),
            str(provider.get("translate_cyrillic_formula_text") or ""),
            glossary_digest,
            PDF_PREFLIGHT_CONTRACT_VERSION,
            PDF_PAGE_RANGE_CONTRACT_VERSION,
            PDF_FALLBACK_CONTRACT_VERSION,
            json.dumps(configured_preflight_thresholds(), sort_keys=True),
            json.dumps(component_versions, sort_keys=True),
            PDF_PROTECTION_CONTRACT_VERSION,
            OCR_CONTRACT_VERSION,
            str(PDF_OCR_DPI),
            str(PDF_OCR_MIN_PAGE_CHARS),
            str(PDF_OCR_MIN_AVERAGE_CONFIDENCE),
            str(PDF_OCR_LOW_CONFIDENCE_THRESHOLD),
        ]
    )
    provider_digest = hashlib.sha256(provider_payload.encode("utf-8")).hexdigest()[:12]
    return (
        f"{PROJECT_TRANSLATION_CONFIG_SIGNATURE}:{pdf_engine}:"
        f"{provider['service']}:{provider['model']}:{provider_digest}:{office_signature}"
    )


def translation_build_metadata() -> dict:
    return {
        "module": "B",
        "module_version": B_TRANSLATION_COMPONENT_VERSION,
        "build_commit": (
            user_environment_value("RFQ_BUILD_COMMIT")
            or user_environment_value("GIT_COMMIT")
            or "source-checkout"
        ),
        "preflight_version": PDF_PREFLIGHT_CONTRACT_VERSION,
        "page_range_contract_version": PDF_PAGE_RANGE_CONTRACT_VERSION,
        "fallback_contract_version": PDF_FALLBACK_CONTRACT_VERSION,
        "ocr_contract_version": OCR_CONTRACT_VERSION,
        "office_batch_isolation_version": (
            OFFICE_BATCH_FAILURE_ISOLATION_CONTRACT_VERSION
        ),
        "output_naming_contract_version": OUTPUT_NAMING_CONTRACT_VERSION,
        "component_versions": runtime_component_versions(),
    }


def short_digest(path: Path) -> str:
    return file_sha256(path)[:10]


def pdfmathtranslate_output_name(source_path: Path) -> str:
    return f"中文翻译_{short_digest(source_path)}.pdf"


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file_obj:
        for chunk in iter(lambda: file_obj.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_json_file(path: Path, default):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return default


def write_json_file(path: Path, payload) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{time.perf_counter_ns():x}.tmp")
    try:
        temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        for attempt in range(6):
            try:
                temporary.replace(path)
                return
            except PermissionError:
                if attempt == 5:
                    raise
                time.sleep(0.05)
    finally:
        try:
            temporary.unlink(missing_ok=True)
        except OSError:
            pass


def merge_translation_cache_files(shared_cache_path: Path, worker_cache_path: Path) -> None:
    shared_cache = load_json_file(shared_cache_path, {})
    worker_cache_payload = load_json_file(worker_cache_path, {})
    if isinstance(shared_cache, dict) and isinstance(worker_cache_payload, dict):
        shared_cache.update(worker_cache_payload)
        write_json_file(shared_cache_path, shared_cache)


def remove_temporary_tree(path: Path, attempts: int = 6) -> bool:
    for attempt in range(attempts):
        if not path.exists():
            return True
        try:
            shutil.rmtree(path)
        except OSError:
            if attempt == attempts - 1:
                return False
            time.sleep(0.1)
    return not path.exists()


def append_jsonl(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as file_obj:
        file_obj.write(json.dumps(payload, ensure_ascii=False) + "\n")


def read_text_safe(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return ""


def redact_sensitive_text(text: str) -> str:
    redacted = str(text or "")
    for env_name in (VECTOR_ENGINE_API_KEY_ENV,):
        secret = user_environment_value(env_name)
        if secret:
            redacted = redacted.replace(secret, "***")
    redacted = re.sub(r"(?i)Bearer\s+[A-Za-z0-9._~+/=-]+", "Bearer ***", redacted)
    redacted = re.sub(r"(?i)(api[_-]?key[=:\s]+)[^\s,;]+", r"\1***", redacted)
    return redacted


def latest_pdf_runtime_manifest(output_root: Path) -> dict:
    manifests = sorted(
        output_root.rglob("b_pdfmathtranslate_next_manifest.json"),
        key=lambda item: item.stat().st_mtime,
        reverse=True,
    )
    if not manifests:
        return {}
    payload = load_json_file(manifests[0], {})
    if isinstance(payload, dict):
        payload["_manifest_path"] = str(manifests[0])
        return payload
    return {"_manifest_path": str(manifests[0])}


def copy_if_file(source: Path | None, target: Path) -> str | None:
    if source is None or not source.is_file():
        return None
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)
    return str(target)


def copy_redacted_text_file(source: Path | None, target: Path) -> str | None:
    if source is None or not source.is_file():
        return None
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(
        redact_sensitive_text(source.read_text(encoding="utf-8", errors="replace")),
        encoding="utf-8",
    )
    return str(target)


def pdf_page_count(path: Path) -> int | None:
    try:
        return len(PdfReader(str(path)).pages)
    except Exception:
        return None


def pdf_translation_preflight(path: Path) -> dict:
    return inspect_pdf_preflight(path)


def extract_pdf_text_for_qa(path: Path, max_pages: int = 20) -> str:
    try:
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages[:max_pages]:
            texts.append(page.extract_text() or "")
        return "\n".join(texts)
    except Exception:
        return ""


def render_pdf_sample_pages(pdf_path: Path, evidence_dir: Path, page_numbers: list[int]) -> list[str]:
    renders: list[str] = []
    try:
        document = pdfium.PdfDocument(str(pdf_path))
        page_count = len(document)
        for page_number in page_numbers:
            if page_number < 1 or page_number > page_count:
                continue
            output = evidence_dir / f"p{page_number}.png"
            page = document[page_number - 1]
            bitmap = page.render(scale=1.5)
            image = bitmap.to_pil()
            image.save(output)
            renders.append(str(output))
        document.close()
    except Exception:
        return renders
    return renders


def project_system_dir_from_cache(cache_path: Path) -> Path:
    pointer_path = cache_path.with_suffix(cache_path.suffix + ".system_dir")
    if pointer_path.is_file():
        try:
            pointed_dir = Path(pointer_path.read_text(encoding="utf-8").strip())
            if pointed_dir.name == PROJECT_SYSTEM_DIRNAME and pointed_dir.is_dir():
                return pointed_dir
        except OSError:
            pass
    if cache_path.parent.name == "translation_cache_workers":
        return cache_path.parent.parent
    return cache_path.parent


def utc_now_text() -> str:
    return datetime.now().isoformat(timespec="microseconds")


def relative_posix(path: Path, root: Path) -> str:
    return path.resolve().relative_to(root.resolve()).as_posix()


def resolve_project_output_dir(project_package: Path, output_dir: Path | None) -> Path:
    candidate = output_dir or Path(PROJECT_TRANSLATED_DIRNAME)
    resolved = candidate.resolve() if candidate.is_absolute() else (project_package / candidate).resolve()
    try:
        resolved.relative_to(project_package.resolve())
    except ValueError as exc:
        raise ValueError("输出目录必须位于项目资料包内") from exc
    return resolved


def project_file_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in PDF_SUFFIXES:
        return "pdf"
    if suffix in DOCX_SUFFIXES:
        return "docx"
    if suffix in XLSX_SUFFIXES:
        return "xlsx"
    if suffix in LEGACY_DOC_SUFFIXES:
        return "doc_legacy"
    if suffix in LEGACY_XLS_SUFFIXES:
        return "xls_legacy"
    return "unsupported"


def is_supported_translation_file(path: Path) -> bool:
    return project_file_type(path) != "unsupported"


def resolve_requested_file(
    project_package: Path,
    source_dir: Path,
    requested_path: str | Path,
) -> tuple[Path, str, str | None]:
    requested = Path(requested_path)
    if requested.is_absolute():
        return requested, requested.as_posix(), "待翻译文件必须使用项目资料包内的相对路径"
    if requested.parts and requested.parts[0].casefold() == PROJECT_SOURCE_DIRNAME.casefold():
        candidate = (project_package / requested).resolve()
    else:
        candidate = (source_dir / requested).resolve()
    try:
        inside_source = candidate.relative_to(source_dir.resolve())
    except ValueError:
        return candidate, requested.as_posix(), "待翻译文件必须位于 01_原始询价文件 内"
    source_relative = (Path(PROJECT_SOURCE_DIRNAME) / inside_source).as_posix()
    if not is_supported_translation_file(candidate):
        return candidate, source_relative, f"暂不支持的文件类型：{candidate.suffix or '无扩展名'}"
    return candidate, source_relative, None


def resolve_requested_pdf(
    project_package: Path,
    source_dir: Path,
    requested_path: str | Path,
) -> tuple[Path, str, str | None]:
    return resolve_requested_file(project_package, source_dir, requested_path)


def discover_project_files(source_dir: Path) -> list[tuple[Path, str, str | None]]:
    discovered = []
    for path in source_dir.rglob("*"):
        if not path.is_file() or path.name.startswith("~$") or not is_supported_translation_file(path):
            continue
        inside_source = path.resolve().relative_to(source_dir.resolve())
        source_relative = (Path(PROJECT_SOURCE_DIRNAME) / inside_source).as_posix()
        discovered.append((path.resolve(), source_relative, None))
    return sorted(discovered, key=lambda item: item[1].casefold())


def discover_project_pdfs(source_dir: Path) -> list[tuple[Path, str, str | None]]:
    return [
        item for item in discover_project_files(source_dir)
        if project_file_type(item[0]) == "pdf"
    ]


def selected_upload_relative_files(project_package: Path, manifest_path: Path | None = None) -> list[str] | None:
    selected_manifest = manifest_path or project_package / PROJECT_SYSTEM_DIRNAME / SELECTED_UPLOAD_MANIFEST_NAME
    if not selected_manifest.exists():
        return None
    payload = load_json_file(selected_manifest, {})
    files = payload.get("files", []) if isinstance(payload, dict) else []
    selected_files = []
    for item in files:
        if not isinstance(item, dict) or not item.get("selected"):
            continue
        relative_path = (
            item.get("stored_relative_path")
            or item.get("package_relative_path")
            or item.get("relative_path")
            or item.get("browser_relative_path")
        )
        if relative_path:
            selected_files.append(str(relative_path))
    return selected_files


def primary_output_path(entry: dict) -> str | None:
    outputs = entry.get("outputs", {}) if isinstance(entry, dict) else {}
    return (
        entry.get("output_pdf")
        or entry.get("output_docx")
        or entry.get("output_xlsx")
        or outputs.get("pdf")
        or outputs.get("docx")
        or outputs.get("xlsx")
        or entry.get("output_txt")
        or outputs.get("txt")
    )


OUTPUT_NAMING_ENTRY_FIELDS = (
    "output_naming_source_relative_path",
    "display_file_name",
    "download_file_name",
    "display_relative_path",
    "physical_output_file",
    "physical_output_relative_path",
    "physical_output_path",
    "output_extension",
    "output_name_conflict_index",
    "output_name_conflict_resolved",
    "physical_name_sanitized",
    "output_path_shortened",
    "output_relative_parent_preserved",
    "output_naming_contract_version",
)


def attach_output_naming_fields(entry: dict, plan: OutputNamingPlan | None) -> None:
    if plan is None:
        return
    fields = plan.to_manifest_fields()
    entry["output_naming_source_relative_path"] = fields.pop("source_relative_path")
    entry.update(fields)
    entry["output_naming_contract_version"] = OUTPUT_NAMING_CONTRACT_VERSION


def output_naming_relative_path(source_relative_path: str) -> str:
    source = PurePosixPath(source_relative_path.replace("\\", "/"))
    if source.parts and source.parts[0].casefold() == PROJECT_SOURCE_DIRNAME.casefold():
        source = PurePosixPath(*source.parts[1:])
    return source.as_posix()


def reusable_output_naming_fields(
    previous_entry: dict | None,
    source_relative_path: str,
    translated_dir: Path,
) -> dict | None:
    """Recover a stable B13 plan, including compatible pre-B13 original names."""

    if not isinstance(previous_entry, dict):
        return None
    source_posix = output_naming_relative_path(source_relative_path)
    required = {
        "output_naming_source_relative_path",
        "display_file_name",
        "physical_output_relative_path",
        "output_extension",
    }
    if required.issubset(previous_entry):
        fields = {key: previous_entry.get(key) for key in OUTPUT_NAMING_ENTRY_FIELDS}
        fields["source_relative_path"] = fields.pop("output_naming_source_relative_path")
        return fields

    output_text = primary_output_path(previous_entry)
    if not output_text:
        return None
    output_path = Path(output_text)
    try:
        physical_relative = output_path.resolve().relative_to(translated_dir.resolve())
        output_extension = translated_output_extension(PurePosixPath(source_posix).suffix)
    except (OSError, ValueError):
        return None
    source = PurePosixPath(source_posix)
    match = re.fullmatch(
        rf"{re.escape(source.stem)}-译(?: \((\d+)\))?{re.escape(output_extension)}",
        output_path.name,
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    conflict_index = int(match.group(1) or 1)
    display_relative = PurePosixPath(*source.parent.parts, output_path.name)
    if physical_relative.as_posix().casefold() != display_relative.as_posix().casefold():
        return None
    return {
        "source_relative_path": source_posix,
        "display_file_name": output_path.name,
        "download_file_name": output_path.name,
        "display_relative_path": display_relative.as_posix(),
        "physical_output_file": output_path.name,
        "physical_output_relative_path": physical_relative.as_posix(),
        "physical_output_path": str(output_path),
        "output_extension": output_extension,
        "output_name_conflict_index": conflict_index,
        "output_name_conflict_resolved": conflict_index > 1,
        "physical_name_sanitized": False,
        "output_path_shortened": False,
        "output_relative_parent_preserved": True,
        "output_naming_contract_version": OUTPUT_NAMING_CONTRACT_VERSION,
    }


def materialize_planned_project_outputs(
    entry: dict,
    plan: OutputNamingPlan | None,
    *,
    allow_replace: bool,
) -> None:
    """Move staged artifacts to the planned project path and update outputs."""

    attach_output_naming_fields(entry, plan)
    if plan is None:
        return
    outputs = entry.setdefault("outputs", {})
    primary_key = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".xlsx": "xlsx",
        ".xlsm": "xlsx",
    }[plan.output_extension]
    staged_text = outputs.get(primary_key)
    if not staged_text:
        return
    staged = Path(staged_text)
    if not staged.is_file():
        return
    target = Path(plan.physical_output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    if staged.resolve() != target.resolve():
        if target.exists() and not allow_replace:
            raise FileExistsError(f"译文目标已存在，拒绝覆盖：{plan.physical_relative_path}")
        os.replace(staged, target)
    outputs[primary_key] = str(target)

    staged_txt_text = outputs.get("txt")
    if staged_txt_text:
        staged_txt = Path(staged_txt_text)
        if staged_txt.is_file():
            target_txt = target.with_name(f"{target.name}.txt")
            if staged_txt.resolve() != target_txt.resolve():
                if target_txt.exists() and not allow_replace:
                    raise FileExistsError(
                        f"译文文本目标已存在，拒绝覆盖：{target_txt.name}"
                    )
                os.replace(staged_txt, target_txt)
            outputs["txt"] = str(target_txt)


def project_manifest_summary(files: list[dict]) -> dict:
    statuses = ("pending", "running", "success", "partial", "failed", "blocked", "skipped")
    summary = {
        "input_files": len(files),
        "delivered_files": sum(
            item.get("status") in {"success", "partial"}
            or (item.get("status") == "skipped" and bool(primary_output_path(item)))
            for item in files
        ),
        "ocr_required_files": sum(bool(item.get("ocr_required")) for item in files),
    }
    for status in statuses:
        summary[status] = sum(item.get("status") == status for item in files)
    return summary


def failed_project_entry(
    source_path: Path,
    source_relative_path: str,
    error: str,
    mode: str,
) -> dict:
    return {
        "source_file": source_path.name,
        "source_path": str(source_path),
        "source_relative_path": source_relative_path,
        "file_type": project_file_type(source_path),
        "output_file": None,
        "output_path": None,
        "output_relative_path": None,
        "output_pdf": None,
        "output_docx": None,
        "output_xlsx": None,
        "output_txt": None,
        "status": "failed",
        "page_count": None,
        "language": "Unknown",
        "segment_count": 0,
        "method": "not_processed",
        "translation_method": "not_processed",
        "mode": mode,
        "outputs": {},
        "warnings": [],
        "errors": [error],
        "risks": [error],
        "ocr_required": False,
        "timing_breakdown": {},
        "source_sha256_before": file_sha256(source_path) if source_path.exists() and source_path.is_file() else None,
        "source_sha256_after": file_sha256(source_path) if source_path.exists() and source_path.is_file() else None,
        "source_unchanged": True if source_path.exists() and source_path.is_file() else None,
    }


def base_progress_entry(source_path: Path, source_relative_path: str, mode: str) -> dict:
    return {
        "source_file": source_path.name,
        "source_path": str(source_path),
        "source_relative_path": source_relative_path,
        "file_type": project_file_type(source_path),
        "source_sha256": file_sha256(source_path) if source_path.exists() and source_path.is_file() else None,
        "page_count": None,
        "status": "pending",
        "started_at": None,
        "completed_at": None,
        "elapsed_seconds": None,
        "output_pdf": None,
        "output_docx": None,
        "output_xlsx": None,
        "output_txt": None,
        "output_file": None,
        "output_path": None,
        "output_relative_path": None,
        "display_file_name": None,
        "download_file_name": None,
        "display_relative_path": None,
        "physical_output_file": None,
        "physical_output_relative_path": None,
        "physical_output_path": None,
        "output_naming_contract_version": OUTPUT_NAMING_CONTRACT_VERSION,
        "segment_count": 0,
        "translation_method": None,
        "cache_hit": False,
        "skipped_reason": None,
        "mode": mode,
        "warnings": [],
        "errors": [],
        "ocr_required": False,
        "timing_breakdown": {},
        "config_signature": PROJECT_TRANSLATION_CONFIG_SIGNATURE,
    }


def skipped_project_entry(
    source_path: Path,
    source_relative_path: str,
    reason: str,
    mode: str,
    previous_entry: dict | None = None,
) -> dict:
    entry = base_progress_entry(source_path, source_relative_path, mode)
    reused_result_status = None
    if previous_entry:
        reused_result_status = previous_entry.get("reused_result_status") or previous_entry.get("status")
        for key in (
            "page_count",
            "language",
            "segment_count",
            "method",
            "translation_method",
            "translation_engine",
            "model_configured",
            "pdf_engine",
            "outputs",
            "output_file",
            "output_path",
            "output_relative_path",
            "output_pdf",
            "output_docx",
            "output_xlsx",
            "output_txt",
            "warnings",
            "errors",
            "risks",
            "error_summary",
            "ocr_required",
            "pdf_preflight",
            "qa",
            "pdfmathtranslate",
            *OUTPUT_NAMING_ENTRY_FIELDS,
        ):
            if key in previous_entry:
                entry[key] = previous_entry[key]
        previous_timing = previous_entry.get("timing_breakdown")
        if isinstance(previous_timing, dict) and previous_timing:
            entry["previous_timing_breakdown"] = previous_timing
    now = utc_now_text()
    previous_risks = list(entry.get("risks") or [])
    is_reused_partial = reused_result_status == "partial"
    if is_reused_partial:
        reason = "已存在且输入未变化，复用上次 partial 结果（未重新调用模型）"
    if reason not in previous_risks:
        previous_risks.append(reason)
    previous_errors = list(entry.get("errors") or []) if is_reused_partial else []
    entry.update(
        {
            "status": "partial" if is_reused_partial else "skipped",
            "completed_at": now,
            "elapsed_seconds": 0.0,
            "cache_hit": bool(previous_entry),
            "skipped_reason": reason,
            "reused_result_status": reused_result_status,
            "errors": previous_errors,
            "risks": previous_risks,
            "source_sha256_before": entry["source_sha256"],
            "source_sha256_after": entry["source_sha256"],
            "source_unchanged": True if entry["source_sha256"] else None,
            "config_signature": PROJECT_TRANSLATION_CONFIG_SIGNATURE,
            "timing_breakdown": {
                "total_seconds": 0.0,
                "file_processing_seconds": 0.0,
            },
        }
    )
    entry.setdefault("outputs", {})
    return entry


def normalize_project_entry_paths(entry: dict, package: Path) -> None:
    outputs = entry.setdefault("outputs", {})
    entry["output_pdf"] = outputs.get("pdf")
    entry["output_docx"] = outputs.get("docx")
    entry["output_xlsx"] = outputs.get("xlsx")
    entry["output_txt"] = outputs.get("txt")
    output_text = primary_output_path(entry)
    if output_text:
        output_path = Path(output_text)
        entry["output_file"] = output_path.name
        entry["output_path"] = str(output_path)
        entry["output_relative_path"] = relative_posix(output_path, package)
    else:
        entry["output_file"] = None
        entry["output_path"] = None
        entry["output_relative_path"] = None


def can_skip_project_file(
    previous_entry: dict | None,
    source_sha256: str,
    mode: str,
    file_type: str,
    config_signature: str = PROJECT_TRANSLATION_CONFIG_SIGNATURE,
) -> bool:
    if not previous_entry:
        return False
    if previous_entry.get("status") not in {"success", "partial", "skipped"}:
        return False
    previous_hash = previous_entry.get("source_sha256") or previous_entry.get("source_sha256_before")
    if previous_hash != source_sha256:
        return False
    if previous_entry.get("mode") != mode:
        return False
    if previous_entry.get("config_signature") != config_signature:
        return False
    outputs = previous_entry.get("outputs", {})
    if file_type == "pdf":
        required_outputs = [previous_entry.get("output_pdf") or outputs.get("pdf"), previous_entry.get("output_txt") or outputs.get("txt")]
    elif file_type in {"docx", "doc_legacy"}:
        required_outputs = [previous_entry.get("output_docx") or outputs.get("docx"), previous_entry.get("output_txt") or outputs.get("txt")]
    elif file_type in {"xlsx", "xls_legacy"}:
        required_outputs = [previous_entry.get("output_xlsx") or outputs.get("xlsx"), previous_entry.get("output_txt") or outputs.get("txt")]
    else:
        return False
    return all(output and Path(output).exists() for output in required_outputs)


def can_skip_project_pdf(previous_entry: dict | None, source_sha256: str, mode: str) -> bool:
    return can_skip_project_file(previous_entry, source_sha256, mode, "pdf")


def is_non_pdf_selection_skip(selection_error: str | None) -> bool:
    return bool(selection_error and selection_error.startswith("暂不支持的文件类型"))


def build_progress_manifest(
    package: Path,
    source_dir: Path,
    translated_dir: Path,
    system_dir: Path,
    mode: str,
    progress_files: list[dict],
    artifacts: dict,
    generated_at: str,
    pdf_engine: str = PDF_ENGINE_LEGACY,
    config_signature: str = PROJECT_TRANSLATION_CONFIG_SIGNATURE,
) -> dict:
    return {
        "module": "B",
        "module_name": "PDF/文件中文翻译引擎",
        "contract_version": "4.0",
        "scope": "b_translation_progress",
        "generated_at": generated_at,
        "updated_at": utc_now_text(),
        "project_package": str(package),
        "source_dir": str(source_dir),
        "translated_dir": str(translated_dir),
        "system_data_dir": str(system_dir),
        "mode": mode,
        "pdf_engine": pdf_engine,
        "config_signature": config_signature,
        "output_naming_contract_version": OUTPUT_NAMING_CONTRACT_VERSION,
        "build": translation_build_metadata(),
        "summary": project_manifest_summary(progress_files),
        "files": progress_files,
        "artifacts": artifacts,
    }


def write_progress_manifest(
    progress_path: Path,
    package: Path,
    source_dir: Path,
    translated_dir: Path,
    system_dir: Path,
    mode: str,
    progress_files: list[dict],
    artifacts: dict,
    generated_at: str,
    pdf_engine: str = PDF_ENGINE_LEGACY,
    config_signature: str = PROJECT_TRANSLATION_CONFIG_SIGNATURE,
) -> None:
    write_json_file(
        progress_path,
        build_progress_manifest(
            package,
            source_dir,
            translated_dir,
            system_dir,
            mode,
            progress_files,
            artifacts,
            generated_at,
            pdf_engine,
            config_signature,
        ),
    )


def process_project_pdf(
    source_path: Path,
    output_dir: Path,
    cache_path: Path,
    mode: str,
    output_file_name: str | None = None,
) -> tuple[dict, list[dict]]:
    total_started_perf = time.perf_counter()
    preflight_started_perf = time.perf_counter()
    reader = PdfReader(str(source_path))
    page_count = len(reader.pages)
    page_indices = list(range(page_count))
    source_regions = collect_template_layout_source_regions(source_path, page_indices)
    language_segments = [
        Segment(source_path.name, region.page, region.text, "", region.bbox)
        for region in source_regions
    ]
    language = detect_pdf_language(language_segments)
    pages_with_text = {region.page for region in source_regions}
    missing_text_pages = [page for page in range(1, page_count + 1) if page not in pages_with_text]
    pdf_preflight_seconds = round(time.perf_counter() - preflight_started_perf, 3)

    if not source_regions:
        process_pdf_seconds = round(time.perf_counter() - total_started_perf, 3)
        return (
            {
                "status": "blocked",
                "page_count": page_count,
                "language": language,
                "segment_count": 0,
                "method": "text_layer_preflight",
                "translation_method": "not_started",
                "model_configured": False,
                "mode": mode,
                "outputs": {},
                "warnings": [],
                "errors": ["未检测到可翻译文本层，需要 OCR 后再处理"],
                "risks": ["未检测到可翻译文本层，需要 OCR 后再处理"],
                "ocr_required": True,
                "timing_breakdown": {
                    "pdf_preflight_seconds": pdf_preflight_seconds,
                    "translation_cache_seconds": 0.0,
                    "pdf_render_seconds": 0.0,
                    "text_export_seconds": 0.0,
                    "process_project_pdf_seconds": process_pdf_seconds,
                },
            },
            [],
        )

    translation_cache_started_perf = time.perf_counter()
    translation_cache = build_layout_translation_cache(source_regions, source_path.name, cache_path)
    translation_cache_seconds = round(time.perf_counter() - translation_cache_started_perf, 3)

    pdf_render_started_perf = time.perf_counter()
    output_dir.mkdir(parents=True, exist_ok=True)
    delivered_name = output_file_name or active_output_file_name(translated_pdf_name(source_path))
    output_pdf = output_dir / delivered_name
    output_txt = output_dir / f"{delivered_name}.txt"
    regions = generate_auto_layout_sample_pdf(
        source_path,
        output_pdf,
        page_indices,
        translation_cache=translation_cache,
    )
    pdf_render_seconds = round(time.perf_counter() - pdf_render_started_perf, 3)

    text_export_started_perf = time.perf_counter()
    write_layout_regions_text(regions, output_txt, source_path.name)
    text_export_seconds = round(time.perf_counter() - text_export_started_perf, 3)
    process_pdf_seconds = round(time.perf_counter() - total_started_perf, 3)

    warnings = [
        "翻译使用共享缓存、Google HTTP 翻译通道和本地机械术语规则回退，未配置专用大模型",
        "技术编号、Tag、单位、标准号和型号已按现有保护规则处理，仍需技术人员复核",
        "自动按页面类型选择模板覆盖或表格重绘回填；极窄单元格可能出现小字号",
    ]
    status = "success"
    ocr_required = False
    if missing_text_pages:
        status = "partial"
        ocr_required = True
        warnings.append(
            "以下页面未检测到可翻译文本层，需要 OCR 复核：" + "、".join(str(page) for page in missing_text_pages)
        )
    if language in {"Russian", "English/Russian mixed"}:
        status = "partial"
        warnings.append("检测到俄文内容；当前英语翻译通道无法保证俄文译文质量，已标记低置信度")

    segments = [
        {
            "source_file": source_path.name,
            "page": region.page,
            "text": region.text,
            "translation": region.translation,
            "bbox": list(region.bbox),
            "align": region.align,
            "is_table": region.is_table,
        }
        for region in regions
    ]
    entry = {
        "status": status,
        "page_count": page_count,
        "language": language,
        "segment_count": len(segments),
        "method": "auto_layout_template_or_table_repaint_overlay",
        "translation_method": "cached_google_http_with_local_mechanical_rules_fallback",
        "model_configured": False,
        "mode": mode,
        "outputs": {"pdf": str(output_pdf), "txt": str(output_txt)},
        "warnings": warnings,
        "errors": [],
        "risks": list(warnings),
        "ocr_required": ocr_required,
        "timing_breakdown": {
            "pdf_preflight_seconds": pdf_preflight_seconds,
            "translation_cache_seconds": translation_cache_seconds,
            "pdf_render_seconds": pdf_render_seconds,
            "text_export_seconds": text_export_seconds,
            "process_project_pdf_seconds": process_pdf_seconds,
        },
    }
    return entry, segments


def blocked_pdf_preflight_entry(
    source_path: Path,
    mode: str,
    preflight: dict,
    error_summary: str,
) -> tuple[dict, list[dict]]:
    ocr_required = preflight.get("classification") in {
        "scanned_pdf",
        "mixed_pdf",
        "vector_or_image_only_pdf",
    }
    error_code = str(
        preflight.get("error_code")
        or ("pdf_requires_ocr" if ocr_required else "pdf_unreadable_or_encrypted")
    )
    return (
        {
            "status": "blocked",
            "page_count": preflight.get("page_count"),
            "language": "Unknown",
            "segment_count": 0,
            "method": "pdf_preflight_blocked",
            "translation_method": "not_processed",
            "translation_engine": PDF_ENGINE_PDFMATHTRANSLATE_NEXT,
            "model_configured": False,
            "mode": mode,
            "outputs": {},
            "warnings": [error_summary],
            "errors": [],
            "risks": [error_summary],
            "error_summary": error_summary,
            "error_code": error_code,
            "ocr_required": ocr_required,
            "pdf_preflight": preflight,
            "pdf_route": preflight.get("route"),
            "fallback_attempted": False,
            "fallback_reason": None,
            "qa": {
                "source_page_count": preflight.get("page_count"),
                "output_page_count": None,
                "page_count_matches": False,
                "output_pdf_openable": False,
                "pdf_preflight": preflight,
            },
            "timing_breakdown": {
                "pdf_preflight_seconds": preflight.get("elapsed_seconds"),
            },
        },
        [],
    )


def _process_project_pdf_pdfmathtranslate_next(
    source_path: Path,
    output_dir: Path,
    cache_path: Path,
    mode: str,
    temp_output_root: Path,
    preflight: dict | None = None,
    output_file_name: str | None = None,
) -> tuple[dict, list[dict]]:
    total_started_perf = time.perf_counter()
    system_dir = project_system_dir_from_cache(cache_path)
    source_digest = short_digest(source_path)
    evidence_dir = system_dir / f"pdf证据_{source_digest}"
    evidence_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    delivered_name = output_file_name or active_output_file_name(
        pdfmathtranslate_output_name(source_path)
    )
    output_pdf = output_dir / delivered_name
    output_txt = output_dir / f"{delivered_name}.txt"
    output_pdf.unlink(missing_ok=True)
    output_txt.unlink(missing_ok=True)
    pdf_preflight = preflight or pdf_translation_preflight(source_path)
    source_page_count = pdf_preflight.get("page_count")
    pdf_classification = str(pdf_preflight.get("classification") or "unreadable_pdf")
    ocr_pages = [int(page) for page in pdf_preflight.get("ocr_pages", [])]
    ocr_required = pdf_classification in {
        "scanned_pdf",
        "mixed_pdf",
        "vector_or_image_only_pdf",
    } and bool(ocr_pages)
    if pdf_classification == "unreadable_pdf" or pdf_preflight.get("route") == "blocked":
        return blocked_pdf_preflight_entry(
            source_path,
            mode,
            pdf_preflight,
            "PDF 已损坏、加密或没有可读取内容，当前未完成翻译",
        )
    provider = pdfmathtranslate_provider_config()
    runtime_python = resolve_pdf_runtime_python()

    if not runtime_python.is_file():
        if ocr_required:
            return blocked_pdf_preflight_entry(
                source_path,
                mode,
                pdf_preflight,
                "这是扫描版 PDF，需要 OCR；当前 PDF/OCR 运行时未安装，未完成翻译",
            )
        raise FileNotFoundError(
            "PDF 翻译运行时尚未安装或当前账号无权访问默认目录。"
            f"请使用当前普通 Windows 账号运行 {PDF_RUNTIME_DEPLOY_DIR / 'install_windows.ps1'}；"
            "如默认 LOCALAPPDATA 目录不可写，请将 B_PDF_TRANSLATION_RUNTIME_DIR "
            "设为该账号可写的短路径后重试。"
        )
    if not PDF_RUNTIME_WRAPPER.is_file():
        if ocr_required:
            return blocked_pdf_preflight_entry(
                source_path,
                mode,
                pdf_preflight,
                "这是扫描版 PDF，需要 OCR；当前 OCR 处理组件缺失，未完成翻译",
            )
        raise FileNotFoundError(f"PDF 翻译正式 wrapper 不存在: {PDF_RUNTIME_WRAPPER}")
    if provider["service"] in {
        "openaicompatible",
        "openaicompatbatch",
    } and not provider["api_key_configured"]:
        raise RuntimeError(
            f"缺少 Windows 用户环境变量 {provider['api_key_env']}，"
            "无法调用 VectorEngine PDF 翻译模型"
        )

    command = [
        str(runtime_python),
        str(PDF_RUNTIME_WRAPPER),
        "--pdf",
        str(source_path),
        "--pages",
        f"1-{int(source_page_count)}",
        "--service",
        str(provider["service"]),
        "--profile",
        str(provider["profile"]),
        "--lang-in",
        "auto",
        "--lang-out",
        "zh",
        "--output-root",
        str(temp_output_root),
        "--render-limit",
        "0",
        "--qps",
        str(provider["qps"]),
        "--pool-max-workers",
        str(provider["pool_max_workers"]),
        "--no-dual",
        "--include-full-output-pages",
        "--translate-table-text",
        "--no-auto-extract-glossary",
        "--primary-font-family",
        "sans-serif",
        "--use-rfq-prompt",
        "--use-default-glossary",
        "--use-runtime-protection-glossary",
        "--babeldoc-batch-token-limit",
        str(provider.get("babeldoc_batch_token_limit") or 800),
        "--babeldoc-batch-count-limit",
        str(provider.get("babeldoc_batch_count_limit") or 20),
    ]
    if ocr_required:
        command.extend(
            [
                "--ocr-mode",
                "rapidocr",
                "--ocr-pages",
                ",".join(str(page) for page in ocr_pages),
                "--ocr-dpi",
                str(PDF_OCR_DPI),
                "--ocr-min-page-chars",
                str(PDF_OCR_MIN_PAGE_CHARS),
                "--ocr-min-average-confidence",
                str(PDF_OCR_MIN_AVERAGE_CONFIDENCE),
                "--ocr-low-confidence-threshold",
                str(PDF_OCR_LOW_CONFIDENCE_THRESHOLD),
            ]
        )
    else:
        command.append("--skip-scanned-detection")
    if provider["service"] == "openaicompatible":
        command.extend(
            [
                "--openai-compatible-base-url",
                str(provider["base_url"]),
                "--openai-compatible-model",
                str(provider["model"]),
            ]
        )
    elif provider["service"] == "openaicompatbatch":
        command.extend(
            [
                "--openaicompatbatch-base-url",
                str(provider["base_url"]),
                "--openaicompatbatch-model",
                str(provider["model"]),
                "--openaicompatbatch-repair-model",
                str(provider["repair_model"]),
                "--openaicompatbatch-batch-size",
                str(provider["batch_size"]),
                "--openaicompatbatch-max-chars",
                str(provider["batch_max_chars"]),
                "--openaicompatbatch-flush-ms",
                str(provider["batch_flush_ms"]),
                "--openaicompatbatch-timeout",
                str(provider["batch_timeout"]),
                "--openaicompatbatch-max-retries",
                str(provider["batch_max_retries"]),
                "--openaicompatbatch-request-workers",
                str(provider["batch_request_workers"]),
                "--doclayout-image-size",
                str(provider["doclayout_image_size"]),
            ]
        )
        command.append(
            "--translate-cyrillic-formula-text"
            if provider["translate_cyrillic_formula_text"]
            else "--protect-cyrillic-as-formula"
        )
    child_env = os.environ.copy()
    api_key_env = provider.get("api_key_env")
    if api_key_env:
        api_key = user_environment_value(str(api_key_env))
        if api_key:
            child_env[str(api_key_env)] = api_key

    def execute_attempt(
        attempt_number: int,
        attempt_command: list[str],
        route: str,
    ) -> tuple[subprocess.CompletedProcess[str], str, str, dict, dict]:
        attempt_root = temp_output_root / f"a{attempt_number}"
        current_command = list(attempt_command)
        output_index = current_command.index("--output-root") + 1
        current_command[output_index] = str(attempt_root)
        result = subprocess.run(
            current_command,
            cwd=str(PDF_RUNTIME_SOURCE_DIR),
            env=child_env,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
        )
        stdout = redact_sensitive_text(result.stdout)
        stderr = redact_sensitive_text(result.stderr)
        manifest = latest_pdf_runtime_manifest(attempt_root)
        outputs_payload = (
            manifest.get("outputs", {})
            if isinstance(manifest.get("outputs"), dict)
            else {}
        )
        wrapper_error_code = str(manifest.get("error_code") or "")
        combined_log = "\n".join([stdout, stderr])
        if not wrapper_error_code and re.search(
            r"The document contains no paragraphs\.?", combined_log, flags=re.I
        ):
            wrapper_error_code = "pdf_no_paragraphs_detected"
        if not wrapper_error_code and not outputs_payload.get("mono_pdf"):
            wrapper_error_code = "pdf_engine_no_output"
        attempt_payload = {
            "attempt": attempt_number,
            "route": route,
            "returncode": result.returncode,
            "status": manifest.get("status") or ("success" if result.returncode == 0 else "failed"),
            "error_code": wrapper_error_code or None,
            "ocr_mode": "rapidocr" if "--ocr-mode" in current_command else "off",
        }
        (evidence_dir / f"a{attempt_number}_out.log").write_text(
            stdout, encoding="utf-8", errors="replace"
        )
        (evidence_dir / f"a{attempt_number}_err.log").write_text(
            stderr, encoding="utf-8", errors="replace"
        )
        return result, stdout, stderr, manifest, attempt_payload

    attempts: list[dict] = []
    completed, safe_stdout, safe_stderr, wrapper_manifest, attempt_payload = execute_attempt(
        1,
        command,
        "ocr" if ocr_required else "text",
    )
    attempts.append(attempt_payload)
    fallback_attempted = False
    fallback_reason = None
    first_error_code = attempt_payload.get("error_code")
    if not ocr_required and first_error_code in {
        "pdf_no_paragraphs_detected",
        "pdf_engine_no_output",
        "pdf_page_range_invalid",
    }:
        fallback_attempted = True
        fallback_reason = str(first_error_code)
        ocr_required = True
        ocr_pages = list(range(1, int(source_page_count) + 1))
        fallback_command = [item for item in command if item != "--skip-scanned-detection"]
        fallback_command.extend(
            [
                "--ocr-mode",
                "rapidocr",
                "--ocr-pages",
                ",".join(str(page) for page in ocr_pages),
                "--ocr-dpi",
                str(PDF_OCR_DPI),
                "--ocr-min-page-chars",
                str(PDF_OCR_MIN_PAGE_CHARS),
                "--ocr-min-average-confidence",
                str(PDF_OCR_MIN_AVERAGE_CONFIDENCE),
                "--ocr-low-confidence-threshold",
                str(PDF_OCR_LOW_CONFIDENCE_THRESHOLD),
            ]
        )
        completed, safe_stdout, safe_stderr, wrapper_manifest, attempt_payload = execute_attempt(
            2,
            fallback_command,
            "ocr_fallback",
        )
        attempts.append(attempt_payload)

    (evidence_dir / "out.log").write_text(safe_stdout, encoding="utf-8", errors="replace")
    (evidence_dir / "err.log").write_text(safe_stderr, encoding="utf-8", errors="replace")

    wrapper_manifest_path = Path(str(wrapper_manifest.get("_manifest_path"))) if wrapper_manifest.get("_manifest_path") else None
    evidence_manifest_path = evidence_dir / "m.json"
    write_json_file(evidence_manifest_path, wrapper_manifest)
    evidence_report_path = copy_if_file(
        wrapper_manifest_path.with_name("b_pdfmathtranslate_next_report.md") if wrapper_manifest_path else None,
        evidence_dir / "r.md",
    )
    evidence_log_path = copy_redacted_text_file(
        Path(str(wrapper_manifest.get("log"))) if wrapper_manifest.get("log") else None,
        evidence_dir / "l.log",
    )

    outputs = wrapper_manifest.get("outputs", {}) if isinstance(wrapper_manifest.get("outputs"), dict) else {}
    mono_pdf = Path(str(outputs.get("mono_pdf"))) if outputs.get("mono_pdf") else None
    ocr_payload = (
        wrapper_manifest.get("ocr", {})
        if isinstance(wrapper_manifest.get("ocr"), dict)
        else {}
    )
    warnings: list[str] = []
    errors: list[str] = []
    if completed.returncode != 0:
        errors.append((safe_stderr.strip() or safe_stdout.strip() or f"PDFMathTranslate 退出码 {completed.returncode}")[:1500])
    if wrapper_manifest.get("status") and wrapper_manifest.get("status") != "success":
        errors.append(f"PDFMathTranslate wrapper status={wrapper_manifest.get('status')}")
    if mono_pdf is None or not mono_pdf.is_file():
        errors.append("PDFMathTranslate 未生成 mono translated PDF")
    else:
        shutil.copy2(mono_pdf, output_pdf)
    if ocr_required:
        warnings.extend(str(item) for item in ocr_payload.get("warnings", []))
        if ocr_payload.get("status") == "failed":
            errors.append(
                str(
                    ocr_payload.get("error_summary")
                    or "这是扫描版 PDF，需要 OCR；当前未完成翻译"
                )
            )

    output_page_count = pdf_page_count(output_pdf) if output_pdf.is_file() else None
    page_count_matches = source_page_count is not None and output_page_count == source_page_count
    output_pdf_openable = output_page_count is not None
    output_text = extract_pdf_text_for_qa(output_pdf) if output_pdf.is_file() else ""
    residual_cyrillic_chars = len(re.findall(r"[\u0400-\u04FF]", output_text))
    residual_latin_chars = len(re.findall(r"[A-Za-z]", output_text))
    qa = wrapper_manifest.get("qa", {}) if isinstance(wrapper_manifest.get("qa"), dict) else {}
    protected_tokens = qa.get("protected_tokens", {}) if isinstance(qa.get("protected_tokens"), dict) else {}
    protected_missing_count = int(protected_tokens.get("missing_count") or 0)
    protected_missing = protected_tokens.get("missing_sample", []) if isinstance(protected_tokens.get("missing_sample"), list) else []
    qa_text_metrics = qa.get("text_metrics", {}) if isinstance(qa.get("text_metrics"), dict) else {}
    page_bounds = qa.get("page_bounds", {}) if isinstance(qa.get("page_bounds"), dict) else {}
    output_boundary_violation_count = int(page_bounds.get("output_violation_count") or 0)
    qa_output_metrics = qa_text_metrics.get("output", {}) if isinstance(qa_text_metrics.get("output"), dict) else {}
    actionable_cyrillic_chars = int(
        qa_output_metrics.get("actionable_cyrillic_chars")
        if qa_output_metrics.get("actionable_cyrillic_chars") is not None
        else residual_cyrillic_chars
    )
    actionable_cyrillic_pages = (
        qa.get("actionable_cyrillic_pages", [])
        if isinstance(qa.get("actionable_cyrillic_pages"), list)
        else []
    )
    wrapper_request = wrapper_manifest.get("request", {}) if isinstance(wrapper_manifest.get("request"), dict) else {}
    language_policy = str(wrapper_request.get("language_policy") or "")
    render_pages = [1]
    if source_page_count and source_page_count >= 10:
        render_pages.append(10)
    sample_renders = render_pdf_sample_pages(output_pdf, evidence_dir, render_pages) if output_pdf.is_file() else []

    final_error_code = str(
        wrapper_manifest.get("error_code")
        or attempts[-1].get("error_code")
        or ""
    ) or None
    if not output_pdf_openable:
        errors.append("输出 PDF 无法打开或页数不可读取")
        final_error_code = final_error_code or "pdf_engine_no_output"
    if output_pdf_openable and source_page_count is not None and not page_count_matches:
        errors.append(f"输出页数不一致：source={source_page_count}, output={output_page_count}")
        final_error_code = "pdf_page_range_invalid"
    if protected_missing_count:
        warnings.append(f"保护编号缺失：{protected_missing_count}")
    if output_boundary_violation_count:
        warnings.append(f"输出 PDF 检测到页边界外文本：{output_boundary_violation_count}")
    if actionable_cyrillic_chars:
        page_text = "、".join(
            str(item.get("source_page"))
            for item in actionable_cyrillic_pages
            if isinstance(item, dict) and item.get("source_page") is not None
        )
        warnings.append(
            f"输出 PDF 仍检测到需复核的俄文自然语言：{actionable_cyrillic_chars}"
            + (f"（页：{page_text}）" if page_text else "")
        )
    if residual_latin_chars and language_policy != "preserve_english_translate_cyrillic":
        warnings.append(f"输出 PDF 仍检测到英文字母：{residual_latin_chars}")
    wrapper_warnings = qa.get("warnings", []) if isinstance(qa.get("warnings"), list) else []
    actionable_latin_fragment_count = int(qa.get("actionable_latin_fragment_count") or 0)
    actionable_latin_fragments = (
        qa.get("actionable_latin_fragments", [])
        if isinstance(qa.get("actionable_latin_fragments"), list)
        else []
    )
    if ocr_required and actionable_latin_fragment_count:
        warnings.append(
            f"扫描翻译仍有可翻译英文残留：{actionable_latin_fragment_count}"
        )
    warnings.extend(str(item) for item in wrapper_warnings)
    if not sample_renders:
        warnings.append("未能生成抽样渲染图，请人工打开 PDF 复核")

    if errors:
        status = "blocked" if ocr_required else "failed"
    elif (
        protected_missing_count
        or output_boundary_violation_count
        or (source_page_count is not None and not page_count_matches)
        or (ocr_required and ocr_payload.get("status") == "partial")
        or (ocr_required and actionable_latin_fragment_count)
    ):
        status = "partial"
    else:
        status = "success"
        final_error_code = None

    log_stats = wrapper_manifest.get("log_stats", {}) if isinstance(wrapper_manifest.get("log_stats"), dict) else {}
    wrapper_elapsed = wrapper_manifest.get("elapsed_seconds")
    process_pdf_seconds = round(time.perf_counter() - total_started_perf, 3)
    qa_payload = {
        "source_page_count": source_page_count,
        "output_page_count": output_page_count,
        "page_count_matches": page_count_matches,
        "output_pdf_openable": output_pdf_openable,
        "residual_cyrillic_chars": residual_cyrillic_chars,
        "actionable_cyrillic_chars": actionable_cyrillic_chars,
        "actionable_cyrillic_pages": actionable_cyrillic_pages,
        "residual_latin_chars": residual_latin_chars,
        "actionable_latin_fragment_count": actionable_latin_fragment_count,
        "actionable_latin_fragments": actionable_latin_fragments[:30],
        "language_policy": language_policy,
        "protected_missing_count": protected_missing_count,
        "protected_missing": protected_missing[:50],
        "page_bounds": page_bounds,
        "sample_renders": sample_renders,
        "wrapper_qa": qa,
        "pdf_preflight": pdf_preflight,
        "ocr": ocr_payload,
    }
    error_summary = None
    if status in {"blocked", "failed"}:
        error_summary = str(
            wrapper_manifest.get("error_summary")
            or ocr_payload.get("error_summary")
            or (errors[0] if errors else "PDF 翻译未完成")
        )
    output_note_lines = [
                f"source_pdf: {source_path}",
                f"output_pdf: {output_pdf}",
                f"engine: {PDF_ENGINE_PDFMATHTRANSLATE_NEXT}",
                f"service: {provider['service']}",
                f"model: {provider['model']}",
                f"base_url: {provider.get('base_url')}",
                f"api_key_env: {provider.get('api_key_env')}",
                f"status: {status}",
                f"pdf_classification: {pdf_classification}",
                f"ocr_required: {ocr_required}",
                f"ocr_status: {ocr_payload.get('status') if ocr_required else 'not_required'}",
                f"error_summary: {error_summary or ''}",
                f"source_page_count: {source_page_count}",
                f"output_page_count: {output_page_count}",
                f"residual_cyrillic_chars: {residual_cyrillic_chars}",
                f"actionable_cyrillic_chars: {actionable_cyrillic_chars}",
                f"residual_latin_chars: {residual_latin_chars}",
                f"language_policy: {language_policy}",
                f"protected_missing_count: {protected_missing_count}",
                f"warnings: {'; '.join(warnings)}",
                f"errors: {'; '.join(errors)}",
    ]
    if status == "blocked":
        output_txt.unlink(missing_ok=True)
    else:
        output_txt.write_text("\n".join(output_note_lines), encoding="utf-8")
    timing_breakdown = {
        "process_project_pdf_seconds": process_pdf_seconds,
        "pdf_preflight_seconds": pdf_preflight.get("elapsed_seconds"),
        "ocr_seconds": ocr_payload.get("elapsed_seconds") if ocr_required else None,
        "pdfmathtranslate_wrapper_seconds": wrapper_elapsed,
        "pdfmathtranslate_engine_seconds": log_stats.get("engine_time_seconds"),
        "pdfmathtranslate_finish_cost_seconds": log_stats.get("finish_cost_seconds"),
        "pdfmathtranslate_peak_memory_mb": log_stats.get("peak_memory_mb"),
        "pdfmathtranslate_fallback_ratio": log_stats.get("fallback_ratio"),
        "pdfmathtranslate_translation_total": log_stats.get("translation_total"),
        "pdfmathtranslate_translation_successful": log_stats.get("translation_successful"),
        "pdfmathtranslate_translation_fallback": log_stats.get("translation_fallback"),
        "provider_llm_request_count": log_stats.get("provider_llm_request_count"),
        "provider_llm_elapsed_sum": log_stats.get("provider_llm_elapsed_sum"),
        "provider_repair_request_count": log_stats.get("provider_repair_request_count"),
        "provider_repair_item_count": log_stats.get("provider_repair_item_count"),
        "provider_repair_elapsed_sum": log_stats.get("provider_repair_elapsed_sum"),
    }
    entry = {
        "status": status,
        "page_count": source_page_count,
        "language": "auto",
        "segment_count": 0,
        "method": "pdfmathtranslate_next_babeldoc_layout",
        "translation_method": (
            f"pdfmathtranslate_next_{provider['service']}_{provider['model']}"
        ),
        "translation_engine": PDF_ENGINE_PDFMATHTRANSLATE_NEXT,
        "model_configured": True,
        "mode": mode,
        "outputs": (
            {}
            if status == "blocked"
            else {"pdf": str(output_pdf), "txt": str(output_txt)}
        ),
        "warnings": warnings,
        "errors": errors,
        "risks": list(warnings) + list(errors),
        "error_summary": error_summary,
        "error_code": final_error_code,
        "ocr_required": ocr_required,
        "pdf_preflight": pdf_preflight,
        "pdf_route": "ocr_fallback" if fallback_attempted else pdf_preflight.get("route"),
        "fallback_attempted": fallback_attempted,
        "fallback_reason": fallback_reason,
        "attempts": attempts,
        "engine_version": runtime_component_versions().get("pdf2zh_next"),
        "preflight_version": PDF_PREFLIGHT_CONTRACT_VERSION,
        "component_versions": runtime_component_versions(),
        "qa": qa_payload,
        "pdfmathtranslate": {
            "service": provider["service"],
            "model": provider["model"],
            "base_url": provider.get("base_url"),
            "api_key_env": provider.get("api_key_env"),
            "api_key_configured": provider["api_key_configured"],
            "profile": provider["profile"],
            "qps": provider["qps"],
            "pool_max_workers": provider["pool_max_workers"],
            "thinking": provider["thinking"],
            "batch_size": provider.get("batch_size"),
            "batch_max_chars": provider.get("batch_max_chars"),
            "batch_flush_ms": provider.get("batch_flush_ms"),
            "batch_timeout": provider.get("batch_timeout"),
            "batch_max_retries": provider.get("batch_max_retries"),
            "batch_request_workers": provider.get("batch_request_workers"),
            "repair_model": provider.get("repair_model"),
            "babeldoc_batch_token_limit": provider.get("babeldoc_batch_token_limit"),
            "babeldoc_batch_count_limit": provider.get("babeldoc_batch_count_limit"),
            "doclayout_image_size": provider.get("doclayout_image_size"),
            "translate_cyrillic_formula_text": provider.get(
                "translate_cyrillic_formula_text"
            ),
            "wrapper_manifest": str(evidence_manifest_path),
            "wrapper_report": evidence_report_path,
            "wrapper_log": evidence_log_path,
            "evidence_dir": str(evidence_dir),
            "command_redacted": wrapper_manifest.get("command_redacted", []),
            "log_stats": log_stats,
            "ocr": ocr_payload,
            "attempts": attempts,
        },
        "timing_breakdown": timing_breakdown,
    }
    return entry, []


def process_project_pdf_pdfmathtranslate_next(
    source_path: Path,
    output_dir: Path,
    cache_path: Path,
    mode: str,
    output_file_name: str | None = None,
) -> tuple[dict, list[dict]]:
    preflight = getattr(PDF_PREFLIGHT_STATE, "value", None)
    temp_output_root = Path(tempfile.mkdtemp(prefix="rfq_b_pdfmathtranslate_"))
    entry: dict | None = None
    try:
        entry, segments = _process_project_pdf_pdfmathtranslate_next(
            source_path,
            output_dir,
            cache_path,
            mode,
            temp_output_root,
            preflight,
            output_file_name,
        )
        return entry, segments
    except Exception:
        failure_preflight = preflight or pdf_translation_preflight(source_path)
        if failure_preflight.get("classification") in {
            "scanned_pdf",
            "mixed_pdf",
            "vector_or_image_only_pdf",
        }:
            entry, segments = blocked_pdf_preflight_entry(
                source_path,
                mode,
                failure_preflight,
                "这是扫描或无文本层 PDF，需要 OCR；当前 OCR 处理异常，未完成翻译",
            )
            return entry, segments
        raise
    finally:
        cleaned = remove_temporary_tree(temp_output_root)
        if entry is not None:
            pdfmathtranslate = entry.setdefault("pdfmathtranslate", {})
            pdfmathtranslate.pop("temp_output_root", None)
            pdfmathtranslate["temporary_workdir_cleaned"] = cleaned
            if not cleaned:
                warning = "PDF 临时工作目录清理失败，需管理员检查当前用户 TEMP 目录"
                entry.setdefault("warnings", []).append(warning)
                entry.setdefault("risks", []).append(warning)


def process_project_docx(
    source_path: Path,
    output_dir: Path,
    cache_path: Path,
    mode: str,
    output_file_name: str | None = None,
) -> tuple[dict, list[dict]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    delivered_name = output_file_name or active_output_file_name(translated_office_name(source_path))
    output_docx = output_dir / delivered_name
    output_txt = output_dir / f"{delivered_name}.txt"
    warnings = [
        "DOCX 第一版翻译段落、表格、页眉和页脚文本；图片文字和复杂对象需人工复核",
    ]
    warnings.extend(detect_docx_high_risk_objects(source_path))
    translation_cache, diagnostics = build_office_translation_cache(
        collect_docx_texts(source_path),
        cache_path,
    )
    segments = translate_docx(source_path, output_docx, output_txt, translation_cache)
    warnings.extend(diagnostics["warnings"])
    unresolved = [
        segment
        for segment in segments
        if segment.get("translation_source") == "untranslated"
        and office_requires_model_translation(segment["original"])
    ]
    unresolved_count = max(len(unresolved), diagnostics["unresolved_count"])
    status = "partial" if unresolved_count else "success"
    used_model_cache = any(segment.get("translation_source") == "model_cache" for segment in segments)
    entry = {
        "status": status,
        "page_count": None,
        "language": "Office text",
        "segment_count": len(segments),
        "method": "docx_paragraph_table_text_translation",
        "translation_method": (
            "openai_compatible_batch_cache_with_local_mechanical_rules"
            if used_model_cache
            else "local_mechanical_rules_partial" if status == "partial" else "local_mechanical_rules"
        ),
        "model_configured": diagnostics["model_configured"],
        "model_translation": diagnostics,
        "mode": mode,
        "outputs": {"docx": str(output_docx), "txt": str(output_txt)},
        "warnings": warnings,
        "errors": list(diagnostics["errors"]),
        "risks": list(warnings) + list(diagnostics["errors"]),
        "ocr_required": False,
    }
    return entry, segments


def process_project_xlsx(
    source_path: Path,
    output_dir: Path,
    cache_path: Path,
    mode: str,
    output_file_name: str | None = None,
) -> tuple[dict, list[dict]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    delivered_name = output_file_name or active_output_file_name(translated_office_name(source_path))
    output_xlsx = output_dir / delivered_name
    output_txt = output_dir / f"{delivered_name}.txt"
    translation_cache, diagnostics = build_office_translation_cache(
        collect_xlsx_texts(source_path),
        cache_path,
    )
    segments, xlsx_warnings = translate_xlsx(
        source_path,
        output_xlsx,
        output_txt,
        translation_cache,
    )
    base_warning = "XLSX 第一版只翻译普通文本单元格，公式、数字、日期、合并单元格和基础样式保持原样"
    warnings = [base_warning] + xlsx_warnings
    warnings.extend(diagnostics["warnings"])
    unresolved = [
        segment
        for segment in segments
        if segment.get("translation_source") == "untranslated"
        and office_requires_model_translation(segment["original"])
    ]
    unresolved_count = max(len(unresolved), diagnostics["unresolved_count"])
    status = "partial" if unresolved_count else "success"
    used_model_cache = any(segment.get("translation_source") == "model_cache" for segment in segments)
    entry = {
        "status": status,
        "page_count": None,
        "language": "Workbook text",
        "segment_count": len(segments),
        "method": "xlsx_text_cell_translation_preserve_formulas",
        "translation_method": (
            "openai_compatible_batch_cache_with_local_mechanical_rules"
            if used_model_cache
            else "local_mechanical_rules_partial" if status == "partial" else "local_mechanical_rules"
        ),
        "model_configured": diagnostics["model_configured"],
        "model_translation": diagnostics,
        "mode": mode,
        "outputs": {"xlsx": str(output_xlsx), "txt": str(output_txt)},
        "warnings": warnings,
        "errors": list(diagnostics["errors"]),
        "risks": warnings + list(diagnostics["errors"]),
        "ocr_required": False,
    }
    return entry, segments


def resolve_libreoffice_executable() -> Path | None:
    configured = user_environment_value("B_OFFICE_CONVERTER_PATH")
    candidates = [
        Path(configured).expanduser() if configured else None,
        Path(shutil.which("soffice.exe") or shutil.which("soffice") or ""),
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    ]
    for candidate in candidates:
        if candidate and str(candidate) != "." and candidate.is_file():
            return candidate.resolve()
    return None


def legacy_office_blocked_entry(source_path: Path, mode: str) -> dict:
    message = (
        f"旧格式 {source_path.suffix.lower()} 需要 LibreOffice 只读转换；"
        "当前未找到 soffice，请先转换为 .docx/.xlsx 或安装 LibreOffice"
    )
    return {
        "status": "blocked",
        "page_count": None,
        "language": "Legacy Office",
        "segment_count": 0,
        "method": "legacy_office_requires_libreoffice",
        "translation_method": None,
        "model_configured": False,
        "mode": mode,
        "outputs": {},
        "warnings": [message],
        "errors": [],
        "risks": [message],
        "ocr_required": False,
        "legacy_conversion": {
            "status": "blocked",
            "converter": "LibreOffice headless",
            "source_format": source_path.suffix.lower(),
        },
    }


def process_project_legacy_office(
    source_path: Path,
    output_dir: Path,
    cache_path: Path,
    mode: str,
    output_file_name: str | None = None,
) -> tuple[dict, list[dict]]:
    converter = resolve_libreoffice_executable()
    if converter is None:
        return legacy_office_blocked_entry(source_path, mode), []

    target_suffix = ".docx" if source_path.suffix.lower() == ".doc" else ".xlsx"
    target_format = target_suffix.lstrip(".")
    with tempfile.TemporaryDirectory(prefix="rfq_b_legacy_office_") as tmp:
        temp_root = Path(tmp)
        profile_dir = temp_root / "lo_profile"
        profile_dir.mkdir(parents=True, exist_ok=True)
        command = [
            str(converter),
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
            "--convert-to",
            target_format,
            "--outdir",
            str(temp_root),
            str(source_path),
        ]
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=180,
            check=False,
        )
        converted = temp_root / f"{source_path.stem}{target_suffix}"
        if completed.returncode != 0 or not converted.is_file():
            detail = (completed.stderr or completed.stdout or "未生成转换文件").strip()
            message = (
                f"旧格式 {source_path.suffix.lower()} 转换失败，请人工转换为 "
                f"{target_suffix} 后重试：{detail[:500]}"
            )
            entry = legacy_office_blocked_entry(source_path, mode)
            entry.update(
                {
                    "status": "failed",
                    "method": "legacy_office_conversion_failed",
                    "warnings": [],
                    "errors": [message],
                    "risks": [message],
                }
            )
            entry["legacy_conversion"]["status"] = "failed"
            return entry, []

        if target_suffix == ".docx":
            entry, segments = process_project_docx(
                converted,
                output_dir,
                cache_path,
                mode,
                output_file_name,
            )
        else:
            entry, segments = process_project_xlsx(
                converted,
                output_dir,
                cache_path,
                mode,
                output_file_name,
            )
        warning = (
            f"旧格式 {source_path.suffix.lower()} 已通过 LibreOffice headless 只读副本"
            f"转换为 {target_suffix} 后翻译；转换前后版式和复杂对象需人工复核"
        )
        entry["warnings"] = [warning] + list(entry.get("warnings", []))
        entry["risks"] = [warning] + list(entry.get("risks", []))
        entry["method"] = f"legacy_office_libreoffice_to_{target_format}+{entry['method']}"
        entry["legacy_conversion"] = {
            "status": "success",
            "converter": "LibreOffice headless",
            "source_format": source_path.suffix.lower(),
            "converted_format": target_suffix,
        }
        return entry, segments


def process_project_file(
    source_path: Path,
    output_dir: Path,
    cache_path: Path,
    mode: str,
    pdf_engine: str = PDF_ENGINE_LEGACY,
    pdf_preflight: dict | None = None,
) -> tuple[dict, list[dict]]:
    file_type = project_file_type(source_path)
    if file_type == "pdf":
        if pdf_engine == PDF_ENGINE_PDFMATHTRANSLATE_NEXT:
            previous_preflight = getattr(PDF_PREFLIGHT_STATE, "value", None)
            PDF_PREFLIGHT_STATE.value = pdf_preflight
            try:
                return process_project_pdf_pdfmathtranslate_next(
                    source_path,
                    output_dir,
                    cache_path,
                    mode,
                )
            finally:
                if previous_preflight is None:
                    try:
                        del PDF_PREFLIGHT_STATE.value
                    except AttributeError:
                        pass
                else:
                    PDF_PREFLIGHT_STATE.value = previous_preflight
        return process_project_pdf(source_path, output_dir, cache_path, mode)
    if file_type == "docx":
        return process_project_docx(source_path, output_dir, cache_path, mode)
    if file_type == "xlsx":
        return process_project_xlsx(source_path, output_dir, cache_path, mode)
    if file_type in {"doc_legacy", "xls_legacy"}:
        return process_project_legacy_office(source_path, output_dir, cache_path, mode)
    raise ValueError(f"暂不支持的文件类型：{source_path.suffix or '无扩展名'}")


def process_project_package(
    project_package: str | Path,
    relative_files: list[str | Path] | None = None,
    mode: str = "平衡",
    output_dir: str | Path | None = None,
    selected_files_manifest: str | Path | None = None,
    pdf_concurrency: int = 2,
    pdf_engine: str | None = None,
) -> dict:
    try:
        normalized_pdf_concurrency = int(pdf_concurrency)
    except (TypeError, ValueError) as exc:
        raise ValueError("pdf_concurrency 必须是 1 到 4 的整数") from exc
    if not 1 <= normalized_pdf_concurrency <= 4:
        raise ValueError("pdf_concurrency 必须是 1 到 4 的整数")
    package = Path(project_package).resolve()
    if not package.is_dir():
        raise FileNotFoundError(f"项目资料包不存在：{package}")
    source_dir = package / PROJECT_SOURCE_DIRNAME
    if not source_dir.is_dir():
        raise FileNotFoundError(f"项目资料包缺少 {PROJECT_SOURCE_DIRNAME}：{source_dir}")

    normalized_mode = normalize_project_mode(mode)
    normalized_pdf_engine = normalize_pdf_engine(pdf_engine)
    config_signature = project_config_signature(normalized_pdf_engine)
    translated_dir = resolve_project_output_dir(package, Path(output_dir) if output_dir is not None else None)
    source_dir_resolved = source_dir.resolve()
    if translated_dir == source_dir_resolved or source_dir_resolved in translated_dir.parents:
        raise ValueError(f"输出目录不能位于 {PROJECT_SOURCE_DIRNAME} 内")
    system_dir = package / PROJECT_SYSTEM_DIRNAME
    translated_dir.mkdir(parents=True, exist_ok=True)
    system_dir.mkdir(parents=True, exist_ok=True)

    translation_manifest_path = system_dir / "translation_manifest.json"
    selected_manifest_path = system_dir / "selected_translation_manifest.json"
    translation_segments_path = system_dir / "translation_segments.json"
    cache_path = system_dir / "translation_cache.json"
    progress_path = system_dir / "b_translation_progress.json"
    events_path = system_dir / "b_translation_events.jsonl"
    if not cache_path.exists():
        write_json_file(cache_path, {})

    artifacts = {
        "translation_manifest": str(translation_manifest_path),
        "selected_translation_manifest": str(selected_manifest_path),
        "translation_segments": str(translation_segments_path),
        "translation_cache": str(cache_path),
        "b_translation_progress": str(progress_path),
        "b_translation_events": str(events_path),
    }
    selection_manifest_path = Path(selected_files_manifest) if selected_files_manifest is not None else None
    selected_manifest_files = None if relative_files is not None else selected_upload_relative_files(package, selection_manifest_path)
    selection_scope = {
        "status": "not_applied",
        "source": "explicit_relative_files" if relative_files is not None else "source_scan",
        "selected_manifest_path": str(selection_manifest_path or package / PROJECT_SYSTEM_DIRNAME / SELECTED_UPLOAD_MANIFEST_NAME),
    }
    if relative_files is not None:
        requested_files = [
            resolve_requested_file(package, source_dir, requested_path)
            for requested_path in relative_files
        ]
    elif selected_manifest_files is not None:
        requested_files = [
            resolve_requested_file(package, source_dir, requested_path)
            for requested_path in selected_manifest_files
        ]
        selection_scope.update(
            {
                "status": "applied",
                "source": SELECTED_UPLOAD_MANIFEST_NAME,
                "selected_file_count": len(selected_manifest_files),
            }
        )
    else:
        requested_files = discover_project_files(source_dir)
    deduplicated_files = []
    seen_source_paths = set()
    for item in requested_files:
        key = item[1].casefold()
        if key in seen_source_paths:
            continue
        seen_source_paths.add(key)
        deduplicated_files.append(item)
    requested_files = deduplicated_files

    previous_manifest = load_json_file(translation_manifest_path, {})
    previous_files = previous_manifest.get("files", []) if isinstance(previous_manifest, dict) else []
    previous_files_by_source = {
        str(item.get("source_relative_path", item.get("source_file", ""))).replace("\\", "/"): item
        for item in previous_files
        if isinstance(item, dict)
    }
    previous_output_owners: dict[str, set[str]] = {}
    for source_key, previous_entry in previous_files_by_source.items():
        previous_outputs = previous_entry.get("outputs", {})
        if not isinstance(previous_outputs, dict):
            continue
        for output_text in previous_outputs.values():
            if not output_text:
                continue
            output_key = os.path.normcase(os.path.abspath(os.fspath(output_text)))
            previous_output_owners.setdefault(output_key, set()).add(source_key.casefold())
    naming_source_items = [
        (
            source_relative_path.replace("\\", "/"),
            output_naming_relative_path(source_relative_path),
        )
        for source_path, source_relative_path, selection_error in requested_files
        if not selection_error and source_path.suffix.lower() in SUPPORTED_TRANSLATION_SUFFIXES
    ]
    reusable_naming_plans: dict[str, dict] = {}
    reusable_naming_sources: set[str] = set()
    for source_relative_path, naming_relative_path in naming_source_items:
        reusable = reusable_output_naming_fields(
            previous_files_by_source.get(source_relative_path),
            source_relative_path,
            translated_dir,
        )
        if reusable is not None:
            reusable_naming_plans[naming_relative_path] = reusable
            reusable_naming_sources.add(source_relative_path.casefold())
    reusable_physical_paths = {
        str(fields.get("physical_output_relative_path") or "").replace("\\", "/").casefold()
        for fields in reusable_naming_plans.values()
    }
    occupied_physical_paths = [
        path.relative_to(translated_dir).as_posix()
        for path in translated_dir.rglob("*")
        if path.is_file()
        and path.relative_to(translated_dir).as_posix().casefold() not in reusable_physical_paths
    ]
    current_naming_sources = {
        source_relative_path.casefold()
        for source_relative_path, _naming_relative_path in naming_source_items
    }
    occupied_display_paths = [
        str(item.get("display_relative_path"))
        for item in previous_files
        if isinstance(item, dict)
        and item.get("display_relative_path")
        and str(item.get("source_relative_path") or "").replace("\\", "/").casefold()
        not in current_naming_sources
    ]
    naming_plans = plan_translated_outputs(
        [item[1] for item in naming_source_items],
        translated_dir,
        occupied_physical_relative_paths=occupied_physical_paths,
        occupied_display_relative_paths=occupied_display_paths,
        reusable_plans=reusable_naming_plans,
    )
    naming_plans_by_source = {
        source_relative_path: plan
        for (source_relative_path, _naming_relative_path), plan in zip(
            naming_source_items,
            naming_plans,
        )
    }

    run_timestamp = datetime.now().isoformat(timespec="microseconds")
    result_entries: list[dict | None] = [None for _item in requested_files]
    selected_segments_by_index: dict[int, list[dict]] = {}
    replaced_segment_source_paths: set[str] = set()
    progress_files = [
        base_progress_entry(source_path, source_relative_path, normalized_mode)
        for source_path, source_relative_path, _selection_error in requested_files
    ]
    for progress_entry in progress_files:
        plan = naming_plans_by_source.get(
            str(progress_entry["source_relative_path"]).replace("\\", "/")
        )
        attach_output_naming_fields(progress_entry, plan)
    write_progress_manifest(
        progress_path,
        package,
        source_dir,
        translated_dir,
        system_dir,
        normalized_mode,
        progress_files,
        artifacts,
        run_timestamp,
        normalized_pdf_engine,
        config_signature,
    )
    state_lock = Lock()
    worker_cache_root: Path | None = None

    def write_progress_locked() -> None:
        write_progress_manifest(
            progress_path,
            package,
            source_dir,
            translated_dir,
            system_dir,
            normalized_mode,
            progress_files,
            artifacts,
            run_timestamp,
            normalized_pdf_engine,
            config_signature,
        )

    def append_event_locked(payload: dict) -> None:
        append_jsonl(events_path, payload)

    def store_result(file_index: int, entry: dict, segments: list[dict] | None = None) -> None:
        result_entries[file_index] = entry
        progress_files[file_index] = entry
        if segments is not None:
            selected_segments_by_index[file_index] = segments
            replaced_segment_source_paths.add(entry["source_relative_path"])
        write_progress_locked()

    def naming_plan_for(source_relative_path: str) -> OutputNamingPlan | None:
        return naming_plans_by_source.get(source_relative_path.replace("\\", "/"))

    def remove_superseded_outputs(
        source_relative_path: str,
        previous_entry: dict | None,
        current_entry: dict,
    ) -> int:
        if not isinstance(previous_entry, dict) or current_entry.get("status") not in {"success", "partial"}:
            return 0
        current_outputs = current_entry.get("outputs", {})
        if not isinstance(current_outputs, dict):
            return 0
        current_keys = {
            os.path.normcase(os.path.abspath(os.fspath(output_text)))
            for output_text in current_outputs.values()
            if output_text
        }
        previous_outputs = previous_entry.get("outputs", {})
        if not isinstance(previous_outputs, dict):
            return 0
        source_owner = source_relative_path.replace("\\", "/").casefold()
        removed = 0
        for output_text in previous_outputs.values():
            if not output_text:
                continue
            output_path = Path(output_text)
            output_key = os.path.normcase(os.path.abspath(os.fspath(output_path)))
            if output_key in current_keys:
                continue
            if previous_output_owners.get(output_key, {source_owner}) != {source_owner}:
                continue
            try:
                output_path.resolve().relative_to(translated_dir.resolve())
            except (OSError, ValueError):
                continue
            if output_path.is_file():
                output_path.unlink()
                removed += 1
        return removed

    def record_selection_error(file_index: int, source_path: Path, source_relative_path: str, selection_error: str) -> None:
        if is_non_pdf_selection_skip(selection_error):
            entry = skipped_project_entry(source_path, source_relative_path, selection_error, normalized_mode)
            entry["config_signature"] = config_signature
            event_name = "file_skipped"
            event_payload = {
                "timestamp": utc_now_text(),
                "event": event_name,
                "source_file": source_path.name,
                "source_relative_path": source_relative_path,
                "status": "skipped",
                "skipped_reason": selection_error,
            }
        else:
            entry = failed_project_entry(source_path, source_relative_path, selection_error, normalized_mode)
            entry["config_signature"] = config_signature
            event_name = "file_failed"
            event_payload = {
                "timestamp": utc_now_text(),
                "event": event_name,
                "source_file": source_path.name,
                "source_relative_path": source_relative_path,
                "status": "failed",
                "errors": entry["errors"],
            }
        attach_output_naming_fields(entry, naming_plan_for(source_relative_path))
        with state_lock:
            store_result(file_index, entry)
            append_event_locked(event_payload)

    def record_missing_file(file_index: int, source_path: Path, source_relative_path: str) -> None:
        entry = failed_project_entry(
            source_path,
            source_relative_path,
            f"源文件不存在：{source_relative_path}",
            normalized_mode,
        )
        entry["config_signature"] = config_signature
        attach_output_naming_fields(entry, naming_plan_for(source_relative_path))
        with state_lock:
            store_result(file_index, entry)
            append_event_locked(
                {
                    "timestamp": utc_now_text(),
                    "event": "file_failed",
                    "source_file": source_path.name,
                    "source_relative_path": source_relative_path,
                    "status": "failed",
                    "errors": entry["errors"],
                }
            )

    def worker_cache_for(file_index: int, source_relative_path: str) -> Path:
        nonlocal worker_cache_root
        cache_key = hashlib.sha256(source_relative_path.encode("utf-8")).hexdigest()[:16]
        with state_lock:
            if worker_cache_root is None:
                worker_cache_root = Path(tempfile.mkdtemp(prefix="rfq_b_cache_"))
            worker_cache = worker_cache_root / f"{file_index:04d}_{cache_key}.json"
            write_json_file(worker_cache, load_json_file(cache_path, {}))
            worker_cache.with_suffix(worker_cache.suffix + ".system_dir").write_text(
                str(system_dir),
                encoding="utf-8",
            )
        return worker_cache

    def merge_worker_cache(worker_cache: Path) -> None:
        if not worker_cache.exists():
            return
        with state_lock:
            merge_translation_cache_files(cache_path, worker_cache)

    def run_processable_file(file_index: int, source_path: Path, source_relative_path: str) -> None:
        file_type = project_file_type(source_path)
        source_key = source_relative_path.replace("\\", "/")
        naming_plan = naming_plan_for(source_relative_path)
        source_sha256_before = file_sha256(source_path)
        pdf_preflight: dict | None = None
        file_config_signature = config_signature
        if file_type == "pdf" and normalized_pdf_engine == PDF_ENGINE_PDFMATHTRANSLATE_NEXT:
            pdf_preflight = pdf_translation_preflight(source_path)
            preflight_signature = str(pdf_preflight.get("result_signature") or "missing")
            payload = (
                f"{config_signature}|{source_sha256_before}|{preflight_signature}|"
                f"{PDF_PAGE_RANGE_CONTRACT_VERSION}|{PDF_FALLBACK_CONTRACT_VERSION}"
            )
            file_config_signature = (
                f"{config_signature}:pdf-file:"
                f"{hashlib.sha256(payload.encode('utf-8')).hexdigest()[:16]}"
            )
        previous_entry = previous_files_by_source.get(source_key)
        if can_skip_project_file(
            previous_entry,
            source_sha256_before,
            normalized_mode,
            file_type,
            file_config_signature,
        ):
            entry = skipped_project_entry(
                source_path,
                source_relative_path,
                "已存在且输入未变化，按续跑规则跳过",
                normalized_mode,
                previous_entry=previous_entry,
            )
            entry["config_signature"] = file_config_signature
            if file_type == "pdf":
                entry["pdf_engine"] = normalized_pdf_engine
                entry.setdefault("pdf_preflight", pdf_preflight)
                entry.setdefault("pdf_route", (pdf_preflight or {}).get("route"))
            attach_output_naming_fields(entry, naming_plan)
            normalize_project_entry_paths(entry, package)
            with state_lock:
                store_result(file_index, entry)
                append_event_locked(
                    {
                        "timestamp": utc_now_text(),
                        "event": "file_skipped",
                        "source_file": source_path.name,
                        "source_relative_path": source_relative_path,
                        "status": entry["status"],
                        "cache_hit": True,
                        "skipped_reason": entry["skipped_reason"],
                        "reused_result_status": entry.get("reused_result_status"),
                        "display_file_name": entry.get("display_file_name"),
                        "download_file_name": entry.get("download_file_name"),
                        "timing_breakdown": entry.get("timing_breakdown", {}),
                    }
                )
            return

        staging_output_dir = Path(tempfile.mkdtemp(prefix="rfq_b_output_"))
        started_at = utc_now_text()
        started_perf = time.perf_counter()
        with state_lock:
            progress_files[file_index].update(
                {
                    "status": "running",
                    "source_sha256": source_sha256_before,
                    "started_at": started_at,
                    "completed_at": None,
                    "elapsed_seconds": None,
                    "cache_hit": False,
                    "skipped_reason": None,
                }
            )
            write_progress_locked()
            append_event_locked(
                {
                    "timestamp": started_at,
                    "event": "file_started",
                    "source_file": source_path.name,
                    "source_relative_path": source_relative_path,
                    "status": "running",
                    "source_sha256": source_sha256_before,
                    "pdf_engine": normalized_pdf_engine if file_type == "pdf" else None,
                    "pdf_route": (pdf_preflight or {}).get("route") if file_type == "pdf" else None,
                }
            )

        worker_cache = cache_path
        worker_cache_prepare_seconds = 0.0
        worker_cache_merge_seconds = 0.0
        file_processing_seconds = 0.0
        process_started_perf = None
        source_sha256_after: str | None = None
        try:
            if file_type == "pdf":
                worker_cache_started_perf = time.perf_counter()
                worker_cache = worker_cache_for(file_index, source_relative_path)
                worker_cache_prepare_seconds = round(time.perf_counter() - worker_cache_started_perf, 3)
            process_started_perf = time.perf_counter()
            previous_output_file_name = getattr(OUTPUT_NAMING_STATE, "file_name", None)
            if naming_plan is not None:
                OUTPUT_NAMING_STATE.file_name = naming_plan.physical_file_name
            try:
                entry, segments = process_project_file(
                    source_path=source_path,
                    output_dir=staging_output_dir,
                    cache_path=worker_cache,
                    mode=normalized_mode,
                    pdf_engine=normalized_pdf_engine,
                    pdf_preflight=pdf_preflight,
                )
            finally:
                if previous_output_file_name is None:
                    try:
                        del OUTPUT_NAMING_STATE.file_name
                    except AttributeError:
                        pass
                else:
                    OUTPUT_NAMING_STATE.file_name = previous_output_file_name
            file_processing_seconds = round(time.perf_counter() - process_started_perf, 3)
            if file_type == "pdf":
                worker_cache_merge_started_perf = time.perf_counter()
                merge_worker_cache(worker_cache)
                worker_cache_merge_seconds = round(time.perf_counter() - worker_cache_merge_started_perf, 3)
            source_sha256_after = file_sha256(source_path)
            if (
                source_sha256_after == source_sha256_before
                and entry.get("status") in {"success", "partial"}
            ):
                materialize_planned_project_outputs(
                    entry,
                    naming_plan,
                    allow_replace=source_key.casefold() in reusable_naming_sources,
                )
        except Exception as exc:
            if process_started_perf is not None:
                file_processing_seconds = round(time.perf_counter() - process_started_perf, 3)
            entry = failed_project_entry(
                source_path,
                source_relative_path,
                f"处理失败：{exc!r}",
                normalized_mode,
            )
            segments = []
            attach_output_naming_fields(entry, naming_plan)
        finally:
            shutil.rmtree(staging_output_dir, ignore_errors=True)

        if source_sha256_after is None:
            source_sha256_after = file_sha256(source_path)
        completed_at = utc_now_text()
        elapsed_seconds = round(time.perf_counter() - started_perf, 3)
        timing_breakdown = entry.get("timing_breakdown", {})
        if not isinstance(timing_breakdown, dict):
            timing_breakdown = {}
        timing_breakdown = dict(timing_breakdown)
        timing_breakdown.update(
            {
                "worker_cache_prepare_seconds": worker_cache_prepare_seconds,
                "file_processing_seconds": file_processing_seconds,
                "worker_cache_merge_seconds": worker_cache_merge_seconds,
                "total_seconds": elapsed_seconds,
            }
        )
        entry.update(
            {
                "source_file": source_path.name,
                "source_path": str(source_path),
                "source_relative_path": source_relative_path,
                "file_type": file_type,
                "source_sha256_before": source_sha256_before,
                "source_sha256_after": source_sha256_after,
                "source_sha256": source_sha256_after,
                "source_unchanged": source_sha256_before == source_sha256_after,
                "started_at": started_at,
                "completed_at": completed_at,
                "elapsed_seconds": elapsed_seconds,
                "cache_hit": False,
                "skipped_reason": None,
                "config_signature": file_config_signature,
                "pdf_engine": normalized_pdf_engine if file_type == "pdf" else None,
                "timing_breakdown": timing_breakdown,
            }
        )
        entry.setdefault("outputs", {})
        entry.setdefault("warnings", [])
        entry.setdefault("errors", [])
        entry.setdefault("risks", list(entry["warnings"]) + list(entry["errors"]))
        if not entry["source_unchanged"]:
            entry["status"] = "failed"
            entry["errors"].append("原文件哈希在处理前后发生变化，已阻断交付")
            entry["risks"].append("原文件哈希在处理前后发生变化，已阻断交付")
        if entry.get("status") not in {"success", "partial"}:
            entry["outputs"] = {}
        attach_output_naming_fields(entry, naming_plan)
        normalize_project_entry_paths(entry, package)
        entry["superseded_output_files_removed"] = remove_superseded_outputs(
            source_relative_path,
            previous_entry,
            entry,
        )

        segment_payloads = []
        for segment in segments:
            segment_payload = dict(segment)
            segment_payload["source_file"] = source_path.name
            segment_payload["source_path"] = str(source_path)
            segment_payload["source_relative_path"] = source_relative_path
            segment_payload["file_type"] = file_type
            segment_payloads.append(segment_payload)
        entry["segment_count"] = len(segment_payloads)
        with state_lock:
            store_result(file_index, entry, segment_payloads)
            append_event_locked(
                {
                    "timestamp": completed_at,
                    "event": "file_finished" if entry["status"] in {"success", "partial", "blocked"} else "file_failed",
                    "source_file": source_path.name,
                    "source_relative_path": source_relative_path,
                    "file_type": file_type,
                    "status": entry["status"],
                    "elapsed_seconds": elapsed_seconds,
                    "output_pdf": entry.get("output_pdf"),
                    "output_docx": entry.get("output_docx"),
                    "output_xlsx": entry.get("output_xlsx"),
                    "output_txt": entry.get("output_txt"),
                    "display_file_name": entry.get("display_file_name"),
                    "download_file_name": entry.get("download_file_name"),
                    "physical_output_relative_path": entry.get("physical_output_relative_path"),
                    "pdf_engine": entry.get("pdf_engine"),
                    "timing_breakdown": entry.get("timing_breakdown", {}),
                    "warnings": entry.get("warnings", []),
                    "errors": entry.get("errors", []),
                }
            )

    def flush_pdf_batch(pdf_batch: list[tuple[int, Path, str]]) -> None:
        if not pdf_batch:
            return
        max_workers = normalized_pdf_concurrency
        if max_workers == 1 or len(pdf_batch) == 1:
            for item_index, item_source, item_relative in pdf_batch:
                run_processable_file(item_index, item_source, item_relative)
            return
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [
                executor.submit(run_processable_file, item_index, item_source, item_relative)
                for item_index, item_source, item_relative in pdf_batch
            ]
            for future in as_completed(futures):
                future.result()

    pdf_batch: list[tuple[int, Path, str]] = []
    deferred_non_pdf: list[tuple[int, Path, str, str | None, str]] = []
    for file_index, (source_path, source_relative_path, selection_error) in enumerate(requested_files):
        if selection_error:
            deferred_non_pdf.append(
                (file_index, source_path, source_relative_path, selection_error, "selection_error")
            )
            continue
        if not source_path.is_file():
            deferred_non_pdf.append(
                (file_index, source_path, source_relative_path, None, "missing")
            )
            continue
        if project_file_type(source_path) == "pdf":
            pdf_batch.append((file_index, source_path, source_relative_path))
            continue
        deferred_non_pdf.append(
            (file_index, source_path, source_relative_path, None, "office")
        )
    try:
        flush_pdf_batch(pdf_batch)
        for file_index, source_path, source_relative_path, selection_error, item_kind in deferred_non_pdf:
            if item_kind == "selection_error":
                record_selection_error(
                    file_index,
                    source_path,
                    source_relative_path,
                    selection_error or "文件选择失败",
                )
            elif item_kind == "missing":
                record_missing_file(file_index, source_path, source_relative_path)
            else:
                run_processable_file(file_index, source_path, source_relative_path)
    finally:
        if worker_cache_root is not None:
            shutil.rmtree(worker_cache_root, ignore_errors=True)

    files = [entry for entry in result_entries if entry is not None]
    selected_segments = []
    for file_index in sorted(selected_segments_by_index):
        selected_segments.extend(selected_segments_by_index[file_index])

    selected_manifest = {
        "module": "B",
        "module_name": "PDF/文件中文翻译引擎",
        "contract_version": "2.0",
        "scope": "selected_project_file_translation",
        "generated_at": run_timestamp,
        "project_package": str(package),
        "source_dir": str(source_dir),
        "translated_dir": str(translated_dir),
        "system_data_dir": str(system_dir),
        "mode": normalized_mode,
        "pdf_engine": normalized_pdf_engine,
        "config_signature": config_signature,
        "output_naming_contract_version": OUTPUT_NAMING_CONTRACT_VERSION,
        "build": translation_build_metadata(),
        "translation_method": (
            "mixed_pdfmathtranslate_next_pdf_and_docx_xlsx_local_rules"
            if normalized_pdf_engine == PDF_ENGINE_PDFMATHTRANSLATE_NEXT
            else "mixed_pdf_docx_xlsx_translation_with_local_rules_and_pdf_cache"
        ),
        "model_configured": False,
        "quality_limitations": [
            "PDF 使用 PDFMathTranslate-next 时保持原版式能力较强，但仍需复核残留源文、术语和保护编号",
            "旧 PDF 引擎仍可通过 pdf_engine=legacy 回退",
            "扫描页或无文本层页面会执行本地 OCR；低置信度识别、复杂版式和图像噪点仍需人工复核",
            "DOCX/XLSX 第一版优先保留可读结构，不承诺复杂对象和图片文字完全翻译",
        ],
        "selection_scope": selection_scope,
        "selected_relative_paths": [item[1] for item in requested_files],
        "files": files,
        "summary": project_manifest_summary(files),
        "artifacts": artifacts,
    }

    previous_segments = load_json_file(translation_segments_path, [])
    if not isinstance(previous_segments, list):
        previous_segments = []
    if selection_scope["status"] == "applied":
        current_source_paths = {item[1] for item in requested_files}
        previous_segments = [
            segment for segment in previous_segments
            if segment.get("source_relative_path") in current_source_paths
        ]
    merged_segments = [
        segment
        for segment in previous_segments
        if segment.get("source_relative_path") not in replaced_segment_source_paths
    ]
    merged_segments.extend(selected_segments)

    merged_files = {} if selection_scope["status"] == "applied" else {
        item.get("source_relative_path", item.get("source_file", "")): item
        for item in previous_files
        if isinstance(item, dict)
    }
    for entry in files:
        merged_files[entry["source_relative_path"]] = entry
    project_files = sorted(merged_files.values(), key=lambda item: item.get("source_relative_path", "").casefold())
    project_manifest = {
        "module": "B",
        "module_name": "PDF/文件中文翻译引擎",
        "contract_version": "2.0",
        "scope": "project_pdf_translation_index",
        "generated_at": run_timestamp,
        "project_package": str(package),
        "source_dir": str(source_dir),
        "translated_dir": str(translated_dir),
        "system_data_dir": str(system_dir),
        "mode": normalized_mode,
        "pdf_engine": normalized_pdf_engine,
        "config_signature": config_signature,
        "output_naming_contract_version": OUTPUT_NAMING_CONTRACT_VERSION,
        "build": translation_build_metadata(),
        "translation_method": selected_manifest["translation_method"],
        "model_configured": selected_manifest["model_configured"],
        "quality_limitations": selected_manifest["quality_limitations"],
        "selection_scope": selection_scope,
        "last_selected_relative_paths": selected_manifest["selected_relative_paths"],
        "files": project_files,
        "summary": project_manifest_summary(project_files),
        "artifacts": artifacts,
    }

    write_json_file(translation_segments_path, merged_segments)
    write_json_file(translation_manifest_path, project_manifest)
    write_json_file(selected_manifest_path, selected_manifest)
    return selected_manifest


def write_layout_regions_text(regions: list[TextRegion], output_txt: Path, source_file: str) -> None:
    output_txt.parent.mkdir(parents=True, exist_ok=True)
    lines = [f"样张翻译区域：{source_file}", ""]
    for region in regions:
        lines.append(f"【第 {region.page} 页 / {'表格内' if region.is_table else '正文'} / {region.align}】")
        lines.append(f"原文：{region.text}")
        lines.append(f"中文：{region.translation}")
    output_txt.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate Chinese comparison PDFs for an RFQ project package.")
    parser.add_argument(
        "--project-package",
        type=Path,
        help="Project package containing 01_原始询价文件.",
    )
    parser.add_argument(
        "--file",
        action="append",
        dest="relative_files",
        help="Relative PDF/DOCX/XLSX path under the project package or 01_原始询价文件; repeat for multiple files.",
    )
    parser.add_argument(
        "--mode",
        default="平衡",
        help="Processing mode. Current supported value: 平衡 (or balanced).",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        help="Optional output directory inside the project package. Defaults to 02_中文翻译文件.",
    )
    parser.add_argument(
        "--selected-files-manifest",
        type=Path,
        help="Optional selected_upload_files_manifest.json from A/J. Used when --file is not provided.",
    )
    parser.add_argument(
        "--pdf-concurrency",
        type=int,
        default=2,
        help="Maximum concurrent PDF translations. Defaults to 2.",
    )
    args = parser.parse_args()
    if not args.project_package:
        error = {
            "status": "blocked",
            "error_code": "project_package_required",
            "error_summary": "必须显式提供 --project-package 项目资料包路径，不会回退到任何内置样例。",
        }
        # ASCII-safe JSON keeps the structured Chinese message intact when a
        # Windows service captures stderr with a legacy console code page.
        print(json.dumps(error, ensure_ascii=True), file=sys.stderr)
        return 2

    result = process_project_package(
        project_package=args.project_package,
        relative_files=args.relative_files,
        mode=args.mode,
        output_dir=args.output_dir,
        selected_files_manifest=args.selected_files_manifest,
        pdf_concurrency=args.pdf_concurrency,
    )
    print(json.dumps(result["summary"], ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
