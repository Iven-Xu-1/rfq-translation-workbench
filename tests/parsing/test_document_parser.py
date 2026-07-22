import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document


THREAD_DIR = Path(__file__).resolve().parents[1]
PROCESS_DIR = THREAD_DIR / "02_过程文件"
sys.path.insert(0, str(PROCESS_DIR))

from rfq_text_parser.parser import parse_project_package


def write_minimal_xlsx(path: Path) -> None:
    files = {
        "[Content_Types].xml": """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
  <Override PartName="/xl/sharedStrings.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sharedStrings+xml"/>
</Types>""",
        "_rels/.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>""",
        "xl/_rels/workbook.xml.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/sharedStrings" Target="sharedStrings.xml"/>
</Relationships>""",
        "xl/workbook.xml": """<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="Pump List" sheetId="1" r:id="rId1"/></sheets>
</workbook>""",
        "xl/sharedStrings.xml": """<?xml version="1.0" encoding="UTF-8"?>
<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" count="4" uniqueCount="4">
  <si><t>Tag</t></si><si><t>Flow</t></si><si><t>P-101</t></si><si><t>10 m3/h</t></si>
</sst>""",
        "xl/worksheets/sheet1.xml": """<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1"><c r="A1" t="s"><v>0</v></c><c r="B1" t="s"><v>1</v></c></row>
    <row r="2"><c r="A2" t="s"><v>2</v></c><c r="B2" t="s"><v>3</v></c></row>
  </sheetData>
</worksheet>""",
    }
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, content in files.items():
            zf.writestr(name, content)


def write_simple_pdf(path: Path) -> None:
    path.write_bytes(
        b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >> endobj
4 0 obj << /Length 56 >> stream
BT /F1 12 Tf 72 720 Td (Pump flow 10 m3/h) Tj ET
endstream endobj
trailer << /Root 1 0 R >>
%%EOF"""
    )


class DocumentParserTests(unittest.TestCase):
    def make_project(self, tmp: Path) -> Path:
        package = tmp / "SYN-TEST-RFQ-001"
        (package / "01_原始询价文件").mkdir(parents=True)
        (package / "系统数据").mkdir()
        return package

    def test_txt_blocks_include_line_source_locations(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            package = self.make_project(Path(td))
            (package / "01_原始询价文件" / "说明.txt").write_text("Line one\n第二行", encoding="utf-8")

            result = parse_project_package(package)

            doc = result["documents"][0]
            self.assertEqual(doc["parse_status"], "success")
            self.assertEqual(doc["extracted_blocks"][1]["source_location"]["line"], 2)
            self.assertEqual(doc["extracted_blocks"][1]["original_text"], "第二行")

    def test_xlsx_cells_and_tables_keep_sheet_row_column_sources(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            package = self.make_project(Path(td))
            write_minimal_xlsx(package / "01_原始询价文件" / "泵清单.xlsx")

            result = parse_project_package(package)

            doc = result["documents"][0]
            self.assertEqual(doc["parse_status"], "success")
            self.assertEqual(doc["sheet_count"], 1)
            cell_blocks = [block for block in doc["extracted_blocks"] if block["block_type"] == "sheet_cell"]
            self.assertEqual(cell_blocks[2]["source_location"]["cell"], "A2")
            self.assertEqual(doc["tables"][0]["rows"][1]["cells"][1]["text"], "10 m3/h")

    def test_docx_paragraphs_and_table_cells_keep_word_sources(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            package = self.make_project(Path(td))
            docx_path = package / "01_原始询价文件" / "技术要求.docx"
            doc = Document()
            doc.add_paragraph("Pump shall be API 610.")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Tag"
            table.cell(0, 1).text = "Head"
            table.cell(1, 0).text = "P-101"
            table.cell(1, 1).text = "50 m"
            doc.save(docx_path)

            result = parse_project_package(package)

            parsed = result["documents"][0]
            self.assertEqual(parsed["parse_status"], "success")
            self.assertEqual(parsed["extracted_blocks"][0]["source_location"]["paragraph_index"], 1)
            self.assertEqual(parsed["tables"][0]["rows"][1]["cells"][1]["source_location"]["column"], 2)

    def test_simple_pdf_text_is_extracted_with_page_source(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            package = self.make_project(Path(td))
            write_simple_pdf(package / "01_原始询价文件" / "datasheet.pdf")

            result = parse_project_package(package)

            doc = result["documents"][0]
            self.assertEqual(doc["file_type"], "pdf")
            self.assertEqual(doc["parse_status"], "success")
            self.assertIn("Pump flow 10 m3/h", doc["extracted_blocks"][0]["original_text"])
            self.assertEqual(doc["extracted_blocks"][0]["source_location"]["page"], 1)

    def test_failed_file_does_not_block_other_files_and_writes_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            package = self.make_project(Path(td))
            original_dir = package / "01_原始询价文件"
            (original_dir / "ok.txt").write_text("usable text", encoding="utf-8")
            (original_dir / "broken.pdf").write_bytes(b"not a pdf")

            result = parse_project_package(package)

            statuses = {doc["file_name"]: doc["parse_status"] for doc in result["documents"]}
            self.assertEqual(statuses["ok.txt"], "success")
            self.assertEqual(statuses["broken.pdf"], "failed")
            output_dir = package / "系统数据" / "文本解析结果"
            self.assertTrue((output_dir / "parsed_documents.json").exists())
            report = (output_dir / "extraction_report.txt").read_text(encoding="utf-8")
            self.assertIn("成功文件", report)
            self.assertIn("失败文件", report)
            parsed = json.loads((output_dir / "parsed_documents.json").read_text(encoding="utf-8"))
            self.assertIn("documents", parsed)


if __name__ == "__main__":
    unittest.main()
