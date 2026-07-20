from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document


THREAD_ROOT = Path(__file__).resolve().parents[1]
PROCESS_DIR = THREAD_ROOT / "02_过程文件"

import sys

sys.path.insert(0, str(PROCESS_DIR))

from d3_pump_cards.docx_renderer import render_parameter_cards_docx  # noqa: E402
from d3_pump_cards.public_template import create_public_pump_card_template  # noqa: E402


def value(value_cn: str, original: str, unit: str = "", *, status: str = "已提取", source: str = "合成数据 P1") -> dict:
    return {
        "value_cn": value_cn,
        "original_value": original,
        "unit": unit,
        "status": status,
        "source_short": source,
    }


def field(*values: dict, status: str = "已提取") -> dict:
    return {"status": status, "values": list(values)}


def synthetic_card(tag: str, *, long: bool = False, mpa: bool = False) -> dict:
    pressure_unit = "MPa" if mpa else "bar.g"
    fields = {
        "pump_type": field(value("卧式泵", "Horizontal pump")),
        "fluid_name": field(value("合成介质", "Synthetic fluid")),
        "density": field(value("980", "980", "kg/m3")),
        "viscosity": field(value("2.0–3.5", "2.0-3.5", "cP")),
        "fluid_temperature": field(value("35 / 55", "35 / 55", "℃")),
        "environment_temperature": field(status="原文件未提供"),
        "others": field(value("连续运行", "Continuous duty")),
        "process_capacity": field(value("18–22", "18-22", "m3/h")),
        "rated_capacity": field(value("24", "24", "m3/h")),
        "suction_pressure": field(value("1.20", "1.20", pressure_unit)),
        "discharge_pressure": field(
            value("6.50", "6.50", pressure_unit),
            value("6.80", "6.80", pressure_unit, status="冲突待确认", source="合成数据 P2"),
        ),
        "wetted_parts_material": field(value("316L 不锈钢", "316L stainless steel")),
    }
    if long:
        fields["others"] = field(
            value(
                "\n".join(
                    (
                        "长卡合成验证文本",
                        "保留多值、区间和冲突，不对不确定结论做选型判断",
                        "来源列独立展示文件、页码和定位片段",
                        "本行用于验证中文换行、行高自动扩展和卡片防跨页",
                    )
                ),
                "Long synthetic validation text for wrapping and page isolation.",
            )
        )
        fields["wetted_parts_material"] = field(
            value("316L 不锈钢", "316L stainless steel"),
            value("双相不锈钢（待确认）", "Duplex stainless steel", status="冲突待确认", source="合成数据 P3"),
        )
    return {"tag_no": tag, "fields": fields, "page_mode": "single" if long else "double"}


class PublicTemplateTests(unittest.TestCase):
    def test_template_is_six_columns_and_contains_two_card_blocks(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "public_template.docx"
            create_public_pump_card_template(path)
            document = Document(path)
            self.assertEqual(document.core_properties.title, "通用泵参数卡片模板")
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(len(document.tables[0]._tbl.tblGrid.gridCol_lst), 6)
            self.assertEqual(len(document.tables[0].rows), 22)
            text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
            for label in ("泵类型", "Tag No.", "介质名称", "密度", "粘度", "工艺流量", "过流部分材质", "来源"):
                self.assertIn(label, text)

    def test_two_short_cards_share_one_planned_page(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            template = Path(temporary) / "template.docx"
            output = Path(temporary) / "two_cards.docx"
            create_public_pump_card_template(template)
            layout = render_parameter_cards_docx(
                template_path=template,
                output_path=output,
                project_title="合成泵参数卡片验证",
                cards=[synthetic_card("SYN-P-101A/B"), synthetic_card("SYN-P-102", mpa=True)],
            )
            self.assertEqual(layout["planned_page_count"], 1)
            self.assertEqual(layout["two_card_page_count"], 1)
            document = Document(output)
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(len(document.tables[0]._tbl.tblGrid.gridCol_lst), 6)
            text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
            for expected in ("SYN-P-101A/B", "SYN-P-102", "2.0–3.5", "35 / 55", "原文件未提供", "bar.g", "MPa", "6.80（待确认）"):
                self.assertIn(expected, text)

    def test_long_card_isolated_on_its_own_planned_page(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            template = Path(temporary) / "template.docx"
            output = Path(temporary) / "mixed_cards.docx"
            create_public_pump_card_template(template)
            layout = render_parameter_cards_docx(
                template_path=template,
                output_path=output,
                project_title="合成泵参数卡片验证",
                cards=[synthetic_card("SYN-P-101A/B"), synthetic_card("SYN-P-102", mpa=True), synthetic_card("SYN-P-103", long=True)],
            )
            self.assertEqual(layout["planned_page_count"], 2)
            self.assertEqual(layout["page_groups"], [["SYN-P-101A/B", "SYN-P-102"], ["SYN-P-103"]])
            self.assertEqual(layout["two_card_page_count"], 1)
            self.assertEqual(layout["single_card_page_count"], 1)


if __name__ == "__main__":
    unittest.main()
