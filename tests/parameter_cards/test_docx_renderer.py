from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml.etree import QName


THREAD_ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = THREAD_ROOT.parent
PROCESS_DIR = THREAD_ROOT / "02_过程文件"
sys.path.insert(0, str(PROCESS_DIR))

from d3_pump_cards.docx_renderer import (  # noqa: E402
    group_cards_for_pages,
    render_parameter_cards_docx,
)
from d3_pump_cards.public_template import create_public_pump_card_template  # noqa: E402


def field(value_cn: str, value_original: str | None = None, unit: str = "") -> dict:
    return {
        "status": "已提取",
        "values": [
            {
                "value_cn": value_cn,
                "original_value": value_original if value_original is not None else value_cn,
                "normalized_value": value_cn,
                "unit": unit,
                "source_short": "工艺表 P3",
                "source_ref_id": "src-1",
                "status": "已提取",
            }
        ],
    }


def card(tag: str, *, long: bool = False) -> dict:
    others = "\n".join(f"要求{i}" for i in range(1, 9)) if long else "数量：1"
    return {
        "card_id": f"card-{tag}",
        "tag_no": tag,
        "page_mode": "single" if long else "double",
        "fields": {
            "pump_type": field("容积式泵（卧式）", "Positive Displacement (Horizontal)"),
            "tag_no": field(tag, tag),
            "fluid_name": field("冷凝液", "Cold Condensates"),
            "density": field("981.7", "981.7", "kg/m3"),
            "viscosity": field("0.46", "0.46", "cP"),
            "fluid_temperature": field("60", "60", "℃"),
            "environment_temperature": {"status": "原文件未提供", "values": []},
            "others": field(others, others),
            "process_capacity": field("1820", "1820", "L/h"),
            "rated_capacity": field("2000", "2000", "L/h"),
            "suction_pressure": field("3.4（绝压，待确认表压）", "3.4 (absolute; gauge pressure to be confirmed)", "MPa"),
            "discharge_pressure": field("7（绝压，待确认表压）", "7 (absolute; gauge pressure to be confirmed)", "MPa"),
            "wetted_parts_material": field("泵壳：碳钢；柱塞：11/13 Cr", "Casing CS; Plunger 11/13 Cr"),
        },
    }


class TestD3DocxRenderer(unittest.TestCase):
    def test_page_grouping_pairs_short_cards_and_isolates_long_cards(self) -> None:
        cards = [card("P-1"), card("P-2"), card("P-3", long=True), card("P-4")]
        groups = group_cards_for_pages(cards)
        self.assertEqual([[c["tag_no"] for c in group] for group in groups], [["P-1", "P-2"], ["P-3"], ["P-4"]])

    def test_renderer_clones_template_blocks_and_removes_example_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "public_template.docx"
            create_public_pump_card_template(template)
            output = Path(tmp) / "cards.docx"
            render_parameter_cards_docx(
                template_path=template,
                output_path=output,
                project_title="测试项目参数卡片",
                cards=[card("P-1"), card("P-2"), card("P-3", long=True)],
            )

            generated = Document(output)
            self.assertEqual(len(generated.tables), 2)
            self.assertEqual([len(t.rows) for t in generated.tables], [22, 11])
            all_text = "\n".join(p.text for p in generated.paragraphs)
            all_text += "\n" + "\n".join(cell.text for t in generated.tables for row in t.rows for cell in row.cells)
            self.assertIn("P-1", all_text)
            self.assertIn("P-2", all_text)
            self.assertIn("P-3", all_text)
            self.assertNotIn("以此类推", all_text)
            self.assertNotIn("Value", all_text)
            self.assertTrue(all(t._tbl.tblPr.find(qn("w:tblpPr")) is None for t in generated.tables))

            first_table = generated.tables[0]
            self.assertEqual(len(first_table._tbl.tblGrid.gridCol_lst), 6)
            grid_widths = [int(column.get(qn("w:w"))) for column in first_table._tbl.tblGrid.gridCol_lst]
            self.assertGreaterEqual(grid_widths[2], 800)
            self.assertGreaterEqual(grid_widths[4], 800)
            self.assertNotIn("【", first_table.rows[0].cells[0].text)
            for row in first_table.rows[:11]:
                self.assertNotIn("【", row.cells[1].text)
                self.assertNotIn("【", row.cells[3].text)
            self.assertIn("来源", first_table.rows[0].cells[5].text)
            self.assertIn("工艺表 P3", first_table.rows[0].cells[5].text)
            self.assertIn("工艺表 P3", first_table.rows[1].cells[5].text)

            body_children = list(generated._element.body)
            names = [QName(child).localname for child in body_children]
            self.assertEqual(names[-1], "sectPr")
            break_indexes = [
                index
                for index, child in enumerate(body_children)
                if QName(child).localname == "p" and child.find(".//" + qn("w:br")) is not None
            ]
            table_indexes = [index for index, name in enumerate(names) if name == "tbl"]
            self.assertEqual(len(break_indexes), 1)
            self.assertLess(table_indexes[0], break_indexes[0])
            self.assertLess(break_indexes[0], table_indexes[1])

            bookmark_starts = list(generated._element.body.iter(qn("w:bookmarkStart")))
            bookmark_ids = [item.get(qn("w:id")) for item in bookmark_starts]
            self.assertEqual(len(bookmark_ids), len(set(bookmark_ids)))

    def test_specific_gravity_without_unit_never_defaults_to_density_unit(self) -> None:
        gravity_card = card("P-SG")
        gravity_card["fields"]["density"] = field("比重 1.03", "1.03", "")
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "public_template.docx"
            create_public_pump_card_template(template)
            output = Path(tmp) / "specific_gravity.docx"
            render_parameter_cards_docx(
                template_path=template,
                output_path=output,
                project_title="比重单位验证",
                cards=[gravity_card],
            )
            generated = Document(output)
            density_row = generated.tables[0].rows[1]
            self.assertEqual(density_row.cells[2].text, "无量纲")
            self.assertEqual(density_row.cells[4].text, "无量纲")
            self.assertNotIn("kg/m3", density_row.cells[2].text + density_row.cells[4].text)


if __name__ == "__main__":
    unittest.main()
