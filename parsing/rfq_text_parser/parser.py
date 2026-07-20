from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import re
import sys
import traceback
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET


PARSER_VERSION = "0.1.0"
SUPPORTED_SUFFIXES = {".txt", ".pdf", ".docx", ".xlsx", ".xls"}
WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
SHEET_NS = {"s": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
REL_NS = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def parse_project_package(package_path: str | Path, output_dir: str | Path | None = None) -> dict[str, Any]:
    package = Path(package_path)
    original_dir = package / "01_原始询价文件"
    if output_dir is None:
        output = package / "系统数据" / "文本解析结果"
    else:
        output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)

    documents: list[dict[str, Any]] = []
    if not original_dir.exists():
        raise FileNotFoundError(f"原始询价文件目录不存在: {original_dir}")

    for file_path in sorted([p for p in original_dir.iterdir() if p.is_file()], key=lambda p: p.name.lower()):
        document = parse_file(file_path, package)
        documents.append(document)
        write_json(output / f"document_{document['document_id']}.json", document)

    summary = {
        "parser_version": PARSER_VERSION,
        "project_package": str(package),
        "source_dir": str(original_dir),
        "output_dir": str(output),
        "processed_at": utc_now(),
        "documents": documents,
        "summary": build_summary(documents),
    }
    manifest = {
        "parser_version": PARSER_VERSION,
        "processed_at": summary["processed_at"],
        "project_package": str(package),
        "source_dir": str(original_dir),
        "output_dir": str(output),
        "file_count": len(documents),
        "status_counts": summary["summary"]["status_counts"],
    }

    write_json(output / "parsed_documents.json", summary)
    write_json(output / "parser_manifest.json", manifest)
    (output / "extraction_report.txt").write_text(build_report(summary), encoding="utf-8")
    return summary


def parse_file(file_path: Path, package_path: Path | None = None) -> dict[str, Any]:
    document_id = stable_document_id(file_path, package_path)
    document = new_document(document_id, file_path)
    try:
        suffix = file_path.suffix.lower()
        if suffix == ".txt":
            parse_txt(file_path, document)
        elif suffix == ".docx":
            parse_docx(file_path, document)
        elif suffix == ".xlsx":
            parse_xlsx(file_path, document)
        elif suffix == ".xls":
            parse_xls(file_path, document)
        elif suffix == ".pdf":
            parse_pdf(file_path, document)
        else:
            document["parse_status"] = "skipped"
            document["warnings"].append({"code": "unsupported_file_type", "message": f"暂不支持文件类型: {suffix}"})
    except Exception as exc:
        document["parse_status"] = "failed"
        document["errors"].append(
            {
                "code": "parse_exception",
                "message": str(exc),
                "traceback": traceback.format_exc(limit=4),
            }
        )
    if not document["extracted_blocks"] and document["parse_status"] == "success":
        document["parse_status"] = "low_quality"
        document["warnings"].append({"code": "no_text_extracted", "message": "未提取到可用文本，可能需要 OCR 或人工复核。"})
    return document


def new_document(document_id: str, file_path: Path) -> dict[str, Any]:
    return {
        "document_id": document_id,
        "file_name": file_path.name,
        "file_type": file_path.suffix.lower().lstrip(".") or "unknown",
        "original_path": str(file_path),
        "parse_status": "success",
        "language_hint": "unknown",
        "page_count": None,
        "sheet_count": None,
        "extracted_blocks": [],
        "tables": [],
        "warnings": [],
        "errors": [],
    }


def parse_txt(file_path: Path, document: dict[str, Any]) -> None:
    text = read_text_fallback(file_path)
    for index, line in enumerate(text.splitlines(), start=1):
        if line.strip():
            add_block(
                document,
                "paragraph",
                line,
                {"type": "txt", "line": index},
                confidence=1.0,
            )


def parse_docx(file_path: Path, document: dict[str, Any]) -> None:
    from docx import Document

    doc = Document(str(file_path))
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            add_block(
                document,
                "paragraph",
                text,
                {"type": "word", "paragraph_index": index},
                confidence=0.98,
            )

    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row_index, row in enumerate(table.rows, start=1):
            cells = []
            row_texts = []
            for column_index, cell in enumerate(row.cells, start=1):
                text = normalize_space(cell.text)
                source = {
                    "type": "word",
                    "table_index": table_index,
                    "row": row_index,
                    "column": column_index,
                }
                cells.append({"text": text, "source_location": source})
                if text:
                    row_texts.append(text)
                    add_block(document, "table_cell", text, source, confidence=0.98)
            rows.append({"row_index": row_index, "cells": cells})
            if row_texts:
                add_block(
                    document,
                    "table_row",
                    " | ".join(row_texts),
                    {"type": "word", "table_index": table_index, "row": row_index},
                    confidence=0.98,
                )
        document["tables"].append(
            {
                "table_id": f"{document['document_id']}_t{table_index:03d}",
                "source_location": {"type": "word", "table_index": table_index},
                "rows": rows,
            }
        )


def parse_xlsx(file_path: Path, document: dict[str, Any]) -> None:
    workbook = read_xlsx(file_path)
    document["sheet_count"] = len(workbook)
    for sheet_index, sheet in enumerate(workbook, start=1):
        sheet_name = sheet["name"]
        rows_for_table = []
        for row in sheet["rows"]:
            row_cells = []
            row_texts = []
            row_number = row["row"]
            for cell in row["cells"]:
                value = str(cell["value"]).strip()
                if not value:
                    continue
                source = {
                    "type": "excel",
                    "sheet": sheet_name,
                    "sheet_index": sheet_index,
                    "row": row_number,
                    "column": cell["column"],
                    "cell": cell["ref"],
                }
                row_cells.append({"text": value, "source_location": source})
                row_texts.append(value)
                add_block(document, "sheet_cell", value, source, confidence=0.98)
            if row_cells:
                rows_for_table.append({"row_index": row_number, "cells": row_cells})
                add_block(
                    document,
                    "table_row",
                    " | ".join(row_texts),
                    {"type": "excel", "sheet": sheet_name, "sheet_index": sheet_index, "row": row_number},
                    confidence=0.96,
                )
        if rows_for_table:
            document["tables"].append(
                {
                    "table_id": f"{document['document_id']}_s{sheet_index:03d}",
                    "source_location": {"type": "excel", "sheet": sheet_name, "sheet_index": sheet_index},
                    "rows": rows_for_table,
                }
            )


def parse_xls(file_path: Path, document: dict[str, Any]) -> None:
    try:
        import xlrd
    except ModuleNotFoundError:
        document["parse_status"] = "skipped"
        document["warnings"].append(
            {
                "code": "xls_dependency_missing",
                "message": "旧版 .xls 需要 xlrd 依赖；当前环境未安装，已跳过且未修改原始文件。",
            }
        )
        return

    book = xlrd.open_workbook(str(file_path))
    document["sheet_count"] = book.nsheets
    for sheet_index in range(book.nsheets):
        sheet = book.sheet_by_index(sheet_index)
        rows_for_table = []
        for row_index in range(sheet.nrows):
            row_cells = []
            row_texts = []
            for column_index in range(sheet.ncols):
                value = normalize_space(str(sheet.cell_value(row_index, column_index)))
                if not value:
                    continue
                source = {
                    "type": "excel",
                    "sheet": sheet.name,
                    "sheet_index": sheet_index + 1,
                    "row": row_index + 1,
                    "column": column_index + 1,
                    "cell": f"{column_letter(column_index + 1)}{row_index + 1}",
                }
                row_cells.append({"text": value, "source_location": source})
                row_texts.append(value)
                add_block(document, "sheet_cell", value, source, confidence=0.92)
            if row_cells:
                rows_for_table.append({"row_index": row_index + 1, "cells": row_cells})
                add_block(
                    document,
                    "table_row",
                    " | ".join(row_texts),
                    {"type": "excel", "sheet": sheet.name, "sheet_index": sheet_index + 1, "row": row_index + 1},
                    confidence=0.92,
                )
        if rows_for_table:
            document["tables"].append(
                {
                    "table_id": f"{document['document_id']}_s{sheet_index + 1:03d}",
                    "source_location": {"type": "excel", "sheet": sheet.name, "sheet_index": sheet_index + 1},
                    "rows": rows_for_table,
                }
            )


def parse_pdf(file_path: Path, document: dict[str, Any]) -> None:
    raw = file_path.read_bytes()
    if not raw.startswith(b"%PDF"):
        document["parse_status"] = "failed"
        document["errors"].append({"code": "not_pdf", "message": "文件扩展名为 PDF，但文件头不是 %PDF。"})
        return

    extracted = extract_pdf_with_pypdf(file_path, document)
    if extracted:
        return

    text = extract_simple_pdf_text(raw)
    document["page_count"] = 1 if text else None
    if text:
        add_block(
            document,
            "page_text",
            text,
            {"type": "pdf", "page": 1, "block_index": 1, "method": "simple_stream_text"},
            confidence=0.72,
        )
        document["warnings"].append(
            {"code": "pdf_simple_extractor", "message": "使用基础 PDF 文本流提取，复杂 PDF 可能不完整。"}
        )
    else:
        document["parse_status"] = "low_quality"
        document["warnings"].append({"code": "ocr_needed", "message": "未提取到文本，可能是扫描 PDF 或需要 OCR。"})


def extract_pdf_with_pypdf(file_path: Path, document: dict[str, Any]) -> bool:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError:
        return False

    try:
        reader = PdfReader(str(file_path))
        if reader.is_encrypted:
            document["parse_status"] = "failed"
            document["errors"].append({"code": "encrypted_pdf", "message": "PDF 已加密，未解析。"})
            return True
        document["page_count"] = len(reader.pages)
        found = False
        for page_index, page in enumerate(reader.pages, start=1):
            text = normalize_space(page.extract_text() or "")
            if text:
                found = True
                add_block(
                    document,
                    "page_text",
                    text,
                    {"type": "pdf", "page": page_index, "block_index": 1, "method": "pypdf"},
                    confidence=0.9,
                )
        if not found:
            document["parse_status"] = "low_quality"
            document["warnings"].append({"code": "ocr_needed", "message": "PDF 没有可提取文本，可能需要 OCR。"})
        return True
    except Exception as exc:
        document["warnings"].append({"code": "pypdf_failed", "message": f"pypdf 提取失败，尝试基础文本流: {exc}"})
        return False


def read_xlsx(file_path: Path) -> list[dict[str, Any]]:
    with zipfile.ZipFile(file_path) as zf:
        shared_strings = read_shared_strings(zf)
        workbook_xml = ET.fromstring(zf.read("xl/workbook.xml"))
        rels_xml = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
        rels = {rel.attrib["Id"]: rel.attrib["Target"] for rel in rels_xml.findall("r:Relationship", REL_NS)}
        sheets = []
        for sheet in workbook_xml.findall(".//s:sheet", SHEET_NS):
            name = sheet.attrib.get("name", "Sheet")
            rel_id = sheet.attrib.get(f"{{{OFFICE_REL_NS}}}id")
            target = rels.get(rel_id or "", "")
            if not target:
                continue
            sheet_path = "xl/" + target.lstrip("/")
            if sheet_path not in zf.namelist():
                sheet_path = "xl/" + target.replace("../", "").lstrip("/")
            rows = read_sheet_rows(zf.read(sheet_path), shared_strings)
            sheets.append({"name": name, "rows": rows})
        return sheets


def read_shared_strings(zf: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in zf.namelist():
        return []
    root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
    values = []
    for item in root.findall("s:si", SHEET_NS):
        parts = [node.text or "" for node in item.findall(".//s:t", SHEET_NS)]
        values.append("".join(parts))
    return values


def read_sheet_rows(xml_bytes: bytes, shared_strings: list[str]) -> list[dict[str, Any]]:
    root = ET.fromstring(xml_bytes)
    rows = []
    for row in root.findall(".//s:row", SHEET_NS):
        row_number = int(row.attrib.get("r", "0") or "0")
        cells = []
        for cell in row.findall("s:c", SHEET_NS):
            ref = cell.attrib.get("r", "")
            cell_type = cell.attrib.get("t", "")
            value_node = cell.find("s:v", SHEET_NS)
            inline_node = cell.find("s:is/s:t", SHEET_NS)
            value = ""
            if cell_type == "s" and value_node is not None and value_node.text is not None:
                index = int(value_node.text)
                value = shared_strings[index] if 0 <= index < len(shared_strings) else ""
            elif inline_node is not None and inline_node.text is not None:
                value = inline_node.text
            elif value_node is not None and value_node.text is not None:
                value = value_node.text
            cells.append({"ref": ref, "column": column_index_from_ref(ref), "value": value})
        rows.append({"row": row_number, "cells": cells})
    return rows


def extract_simple_pdf_text(raw: bytes) -> str:
    matches: list[str] = []
    for stream in re.findall(rb"stream\s*(.*?)\s*endstream", raw, flags=re.S):
        for item in re.findall(rb"\((.*?)\)\s*Tj", stream, flags=re.S):
            matches.append(decode_pdf_literal(item))
        for array in re.findall(rb"\[(.*?)\]\s*TJ", stream, flags=re.S):
            for item in re.findall(rb"\((.*?)\)", array, flags=re.S):
                matches.append(decode_pdf_literal(item))
    return normalize_space(" ".join(matches))


def decode_pdf_literal(value: bytes) -> str:
    value = value.replace(rb"\(", b"(").replace(rb"\)", b")").replace(rb"\\", b"\\")
    return value.decode("latin-1", errors="ignore")


def add_block(
    document: dict[str, Any],
    block_type: str,
    text: str,
    source_location: dict[str, Any],
    confidence: float,
) -> None:
    cleaned = normalize_space(text)
    if not cleaned:
        return
    order = len(document["extracted_blocks"]) + 1
    document["extracted_blocks"].append(
        {
            "block_id": f"{document['document_id']}_b{order:05d}",
            "block_type": block_type,
            "source_location": source_location,
            "original_text": cleaned,
            "text_order": order,
            "confidence": confidence,
        }
    )


def build_summary(documents: list[dict[str, Any]]) -> dict[str, Any]:
    counts: dict[str, int] = {}
    for document in documents:
        counts[document["parse_status"]] = counts.get(document["parse_status"], 0) + 1
    return {
        "file_count": len(documents),
        "status_counts": counts,
        "success_files": [doc["file_name"] for doc in documents if doc["parse_status"] == "success"],
        "failed_files": [doc["file_name"] for doc in documents if doc["parse_status"] == "failed"],
        "low_quality_files": [doc["file_name"] for doc in documents if doc["parse_status"] == "low_quality"],
        "skipped_files": [doc["file_name"] for doc in documents if doc["parse_status"] == "skipped"],
    }


def build_report(summary: dict[str, Any]) -> str:
    lines = [
        "线程 C 文本解析报告",
        f"处理时间: {summary['processed_at']}",
        f"项目资料包: {summary['project_package']}",
        f"原始文件目录: {summary['source_dir']}",
        f"解析输出目录: {summary['output_dir']}",
        "",
        f"文件总数: {summary['summary']['file_count']}",
        f"成功文件: {', '.join(summary['summary']['success_files']) or '无'}",
        f"失败文件: {', '.join(summary['summary']['failed_files']) or '无'}",
        f"低质量文件: {', '.join(summary['summary']['low_quality_files']) or '无'}",
        f"跳过文件: {', '.join(summary['summary']['skipped_files']) or '无'}",
        "",
        "逐文件结果:",
    ]
    for document in summary["documents"]:
        lines.append(f"- {document['file_name']} [{document['parse_status']}] blocks={len(document['extracted_blocks'])} tables={len(document['tables'])}")
        for warning in document["warnings"]:
            lines.append(f"  警告: {warning.get('code')} - {warning.get('message')}")
        for error in document["errors"]:
            lines.append(f"  错误: {error.get('code')} - {error.get('message')}")
    return "\n".join(lines) + "\n"


def stable_document_id(file_path: Path, package_path: Path | None = None) -> str:
    try:
        if package_path:
            key = file_path.relative_to(package_path).as_posix()
        else:
            key = file_path.name
    except ValueError:
        key = str(file_path)
    digest = hashlib.sha1(key.lower().encode("utf-8")).hexdigest()[:12]
    return f"doc_{digest}"


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def read_text_fallback(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def column_index_from_ref(ref: str) -> int | None:
    match = re.match(r"([A-Z]+)", ref.upper())
    if not match:
        return None
    result = 0
    for char in match.group(1):
        result = result * 26 + ord(char) - ord("A") + 1
    return result


def column_letter(index: int) -> str:
    letters = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        letters = chr(ord("A") + remainder) + letters
    return letters


def utc_now() -> str:
    return dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Parse RFQ package source documents into thread C JSON outputs.")
    parser.add_argument("project_package", help="项目资料包路径，例如 项目_测试客户_RFQ_001")
    parser.add_argument("--output-dir", help="可选输出目录；默认写入 项目资料包/系统数据/文本解析结果")
    args = parser.parse_args(argv)
    result = parse_project_package(args.project_package, args.output_dir)
    print(json.dumps({"output_dir": result["output_dir"], "summary": result["summary"]}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
