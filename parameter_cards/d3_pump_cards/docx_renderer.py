from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


CARD_ROW_COUNT = 11
TEMPLATE_CARD_COUNT = 2
SOURCE_COLUMN_GRID_WIDTHS = (1450, 2220, 800, 2250, 800, 999)

ROW_FIELD_MAP = {
    "fluid_name": 0,
    "density": 1,
    "viscosity": 2,
    "fluid_temperature": 3,
    "environment_temperature": 4,
    "others": 5,
    "process_capacity": 6,
    "rated_capacity": 7,
    "suction_pressure": 8,
    "discharge_pressure": 9,
    "wetted_parts_material": 10,
}

LABELS = {
    "fluid_name": ("介质名称：", "Fluid Name:"),
    "density": ("密度：", "Density:"),
    "viscosity": ("粘度：", "Viscosity:"),
    "fluid_temperature": ("介质温度：", "Fluid Temperature:"),
    "environment_temperature": ("环境温度：", "Environment Temperature:"),
    "others": ("其他：", "Others:"),
    "process_capacity": ("工艺流量：", "Process Capacity:"),
    "rated_capacity": ("额定流量：", "Rated Capacity:"),
    "suction_pressure": ("入口压力：", "Suction Pressure:"),
    "discharge_pressure": ("出口压力：", "Discharge Pressure:"),
    "wetted_parts_material": ("过流部分材质：", "Wetted Parts Material:"),
}

DEFAULT_UNITS = {
    "density": "kg/m3",
    "viscosity": "cP",
    "fluid_temperature": "℃",
    "environment_temperature": "℃",
    "process_capacity": "L/h",
    "rated_capacity": "L/h",
    "suction_pressure": "bar.g",
    "discharge_pressure": "bar.g",
}


def _card_is_long(card: dict[str, Any]) -> bool:
    if card.get("page_mode") == "single":
        return True
    if card.get("page_mode") == "double":
        return False
    line_count = 0
    total_chars = 0
    for field in (card.get("fields") or {}).values():
        for value in field.get("values", []):
            cn = str(value.get("value_cn") or "")
            original = str(value.get("original_value") or "")
            line_count += max(cn.count("\n"), original.count("\n")) + 1
            total_chars += len(cn) + len(original)
    return line_count > 18 or total_chars > 650


def group_cards_for_pages(cards: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    groups: list[list[dict[str, Any]]] = []
    pending: list[dict[str, Any]] = []
    for card in cards:
        if _card_is_long(card):
            if pending:
                groups.append(pending)
                pending = []
            groups.append([card])
            continue
        pending.append(card)
        if len(pending) == 2:
            groups.append(pending)
            pending = []
    if pending:
        groups.append(pending)
    return groups


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _clear_cell(cell: _Cell) -> Paragraph:
    paragraphs = list(cell.paragraphs)
    first = paragraphs[0]
    for paragraph in paragraphs[1:]:
        cell._tc.remove(paragraph._p)
    _clear_paragraph_runs(first)
    first.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first.paragraph_format.space_before = Pt(0)
    first.paragraph_format.space_after = Pt(0)
    first.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return first


def _format_run(run: Any, *, size: float, red: bool = False, bold: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) if red else RGBColor(0x00, 0x00, 0x00)


def _add_multiline_value(paragraph: Paragraph, text: str, *, size: float) -> None:
    lines = text.splitlines() or [""]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        _format_run(run, size=size)


def _set_labeled_cell(cell: _Cell, label: str, value: str, *, label_red: bool = False, value_size: float = 8.5) -> None:
    paragraph = _clear_cell(cell)
    label_run = paragraph.add_run(label)
    _format_run(label_run, size=9.5, red=label_red)
    if value:
        paragraph.add_run(" ")
        _add_multiline_value(paragraph, value, size=value_size)


def _set_unit_cell(cell: _Cell, unit: str) -> None:
    paragraph = _clear_cell(cell)
    run = paragraph.add_run(unit)
    _format_run(run, size=8.0, red=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_identity_paragraph(cell: _Cell, label: str, value: str) -> None:
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(label)
    _format_run(label_run, size=8.5, red=True)
    value_run = paragraph.add_run(value)
    _format_run(value_run, size=8.5)


def _field_text(field: dict[str, Any] | None, *, language: str) -> str:
    if not field or field.get("status") == "原文件未提供" or not field.get("values"):
        return "原文件未提供" if language == "cn" else "Not provided in source"
    lines: list[str] = []
    for value in field.get("values", []):
        text = str(value.get("value_cn") if language == "cn" else value.get("original_value") or "")
        if value.get("status") == "冲突待确认":
            text += "（待确认）" if language == "cn" else " (to be confirmed)"
        elif value.get("status") == "低置信度待复核":
            text += "（低置信度）" if language == "cn" else " (low confidence)"
        lines.append(text)
    return "\n".join(lines)


def _field_sources(field: dict[str, Any] | None) -> str:
    if not field or field.get("status") == "原文件未提供" or not field.get("values"):
        return "—"
    ordered: list[str] = []
    for value in field.get("values", []):
        source = str(value.get("source_short") or "来源待复核")
        if value.get("status") == "冲突待确认":
            source += "（待确认）"
        elif value.get("status") == "低置信度待复核":
            source += "（低置信度）"
        if source not in ordered:
            ordered.append(source)
    return "\n".join(ordered)


def _field_unit(field: dict[str, Any] | None, default: str, *, field_name: str = "") -> str:
    units = []
    for value in (field or {}).get("values", []):
        unit = str(value.get("unit") or "").strip()
        if unit and unit not in units:
            units.append(unit)
    if units:
        return "/".join(units)
    if field_name == "density":
        for value in (field or {}).get("values", []):
            value_cn = str(value.get("value_cn") or "")
            original = str(value.get("original_value") or "")
            if "比重" in value_cn or "specific gravity" in original.casefold() or re.search(r"\bSG\s*=", original, flags=re.IGNORECASE):
                return "无量纲"
    return default


def _set_identity_cell(cell: _Cell, card: dict[str, Any]) -> None:
    base = _clear_cell(cell)
    cell._tc.remove(base._p)
    fields = card.get("fields") or {}
    pump_cn = _field_text(fields.get("pump_type"), language="cn")
    pump_original = _field_text(fields.get("pump_type"), language="en")
    tag = card.get("tag_no") or _field_text(fields.get("tag_no"), language="en")
    _add_identity_paragraph(cell, "泵类型：", pump_cn)
    _add_identity_paragraph(cell, "原文：", pump_original)
    _add_identity_paragraph(cell, "Tag No.：", str(tag))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_source_cell(cell: _Cell, source_text: str, *, include_header: bool = False) -> None:
    paragraph = _clear_cell(cell)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if include_header:
        header = paragraph.add_run("来源")
        _format_run(header, size=8.5, bold=True)
        paragraph.add_run().add_break()
    _add_multiline_value(paragraph, source_text, size=7.5)


def _set_cell_grid_widths(table: Table, widths: tuple[int, ...]) -> None:
    grid_columns = list(table._tbl.tblGrid.gridCol_lst)
    if len(grid_columns) != len(widths):
        raise ValueError(f"Expected {len(widths)} grid columns, found {len(grid_columns)}")
    for grid_column, width in zip(grid_columns, widths):
        grid_column.set(qn("w:w"), str(width))

    for row in table.rows:
        grid_index = 0
        for tc in row._tr.tc_lst:
            tc_pr = tc.get_or_add_tcPr()
            grid_span = tc_pr.find(qn("w:gridSpan"))
            span = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
            cell_width = sum(widths[grid_index : grid_index + span])
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_width)
            tc_width.set(qn("w:w"), str(cell_width))
            tc_width.set(qn("w:type"), "dxa")
            grid_index += span


def _ensure_source_column(table: Table) -> None:
    """Accept the public six-column template and retain five-column compatibility."""
    column_count = len(table._tbl.tblGrid.gridCol_lst)
    if column_count == 5:
        table.add_column(Twips(SOURCE_COLUMN_GRID_WIDTHS[-1]))
        _set_cell_grid_widths(table, SOURCE_COLUMN_GRID_WIDTHS)
        return
    if column_count == 6:
        return
    raise ValueError(f"Pump card template must have 5 or 6 columns, found {column_count}")


def _mark_rows_keep_together(table: Table, start: int, count: int) -> None:
    for offset in range(count):
        row = table.rows[start + offset]
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit") is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        keep_next = offset < count - 1
        seen: set[int] = set()
        for cell in row.cells:
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = keep_next


def _fill_card_block(table: Table, start: int, card: dict[str, Any]) -> None:
    fields = card.get("fields") or {}
    _set_identity_cell(table.rows[start].cells[0], card)

    for field_name, relative_row in ROW_FIELD_MAP.items():
        row = table.rows[start + relative_row]
        field = fields.get(field_name)
        cn_text = _field_text(field, language="cn")
        en_text = _field_text(field, language="en")
        cn_label, en_label = LABELS[field_name]
        label_red = field_name == "fluid_name"
        value_size = 7.5 if field_name == "others" else 8.3
        _set_labeled_cell(row.cells[1], cn_label, cn_text, label_red=label_red, value_size=value_size)
        _set_labeled_cell(row.cells[3], en_label, en_text, label_red=label_red, value_size=value_size)
        _set_source_cell(row.cells[5], _field_sources(field), include_header=relative_row == 0)
        if field_name in DEFAULT_UNITS:
            unit = _field_unit(field, DEFAULT_UNITS[field_name], field_name=field_name)
            _set_unit_cell(row.cells[2], unit)
            _set_unit_cell(row.cells[4], unit)

    _mark_rows_keep_together(table, start, CARD_ROW_COUNT)


def _remove_bookmarks(xml: Any) -> None:
    for tag_name in ("w:bookmarkStart", "w:bookmarkEnd"):
        for marker in list(xml.iter(qn(tag_name))):
            parent = marker.getparent()
            if parent is not None:
                parent.remove(marker)


def _replace_title_text(title_xml: Any, project_title: str, body: Any) -> None:
    _remove_bookmarks(title_xml)
    paragraph = Paragraph(title_xml, body)
    if paragraph.runs:
        paragraph.runs[0].text = project_title
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(project_title)
        _format_run(run, size=14, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)


def _add_page_break(body: Any) -> None:
    p_xml = OxmlElement("w:p")
    body.insert(body.index(body.sectPr), p_xml)
    paragraph = Paragraph(p_xml, body)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def render_parameter_cards_docx(
    *,
    template_path: str | Path,
    output_path: str | Path,
    project_title: str,
    cards: list[dict[str, Any]],
) -> dict[str, Any]:
    """Render cards by cloning complete row blocks from a compatible template."""
    template_path = Path(template_path)
    output_path = Path(output_path)
    if not cards:
        raise ValueError("At least one pump card is required")
    document = Document(template_path)
    title_template = deepcopy(document.paragraphs[0]._p)
    table_template = deepcopy(document.tables[0]._tbl)
    body = document._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    page_groups = group_cards_for_pages(cards)
    for page_index, group in enumerate(page_groups):
        title_xml = deepcopy(title_template)
        body.insert(len(body) - 1, title_xml)
        _replace_title_text(title_xml, project_title, body)

        table_xml = deepcopy(table_template)
        _remove_bookmarks(table_xml)
        body.insert(len(body) - 1, table_xml)
        table = Table(table_xml, body)
        tbl_pr = table._tbl.tblPr
        floating_position = tbl_pr.find(qn("w:tblpPr"))
        if floating_position is not None:
            tbl_pr.remove(floating_position)

        keep_rows = CARD_ROW_COUNT * len(group)
        xml_rows = list(table._tbl.tr_lst)
        for row_xml in xml_rows[keep_rows:]:
            table._tbl.remove(row_xml)

        _ensure_source_column(table)

        for card_index, card in enumerate(group):
            _fill_card_block(table, card_index * CARD_ROW_COUNT, card)

        if page_index < len(page_groups) - 1:
            _add_page_break(body)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return {
        "output_path": str(output_path),
        "card_count": len(cards),
        "planned_page_count": len(page_groups),
        "two_card_page_count": sum(1 for group in page_groups if len(group) == 2),
        "single_card_page_count": sum(1 for group in page_groups if len(group) == 1),
        "page_groups": [[card.get("tag_no") for card in group] for group in page_groups],
    }
