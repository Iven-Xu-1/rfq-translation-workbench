from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips


CARD_ROW_COUNT = 11
TEMPLATE_CARD_COUNT = 2
PUBLIC_COLUMN_WIDTHS = (1450, 2320, 800, 2320, 800, 2776)

FIELD_ROWS = (
    ("介质名称：", "Fluid Name:", ""),
    ("密度：", "Density:", "kg/m3"),
    ("粘度：", "Viscosity:", "cP"),
    ("介质温度：", "Fluid Temperature:", "℃"),
    ("环境温度：", "Environment Temperature:", "℃"),
    ("其他：", "Others:", ""),
    ("工艺流量：", "Process Capacity:", "L/h"),
    ("额定流量：", "Rated Capacity:", "L/h"),
    ("入口压力：", "Suction Pressure:", "bar.g"),
    ("出口压力：", "Discharge Pressure:", "bar.g"),
    ("过流部分材质：", "Wetted Parts Material:", ""),
)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    row.height = Twips(610)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.insert(0, tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def _set_cell_margin(cell, *, top: int = 25, bottom: int = 25, left: int = 75, right: int = 75) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _format_run(run, *, size: float = 8.0, bold: bool = False, color: str = "1F2937") -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _set_paragraph(cell, parts: tuple[tuple[str, bool, str], ...], *, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    for text, bold, color in parts:
        _format_run(paragraph.add_run(text), bold=bold, color=color)


def _set_identity(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _shade_cell(cell, "EAF2F8")
    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(3)
    _set_paragraph(cell, (("泵类型：", True, "1F4E78"), ("〔中文值〕", False, "1F2937")))
    for label, placeholder in (("原文：", "[original]"), ("Tag No.：", "[tag]")):
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        _format_run(paragraph.add_run(label), bold=True, color="1F4E78")
        _format_run(paragraph.add_run(placeholder))


def _set_grid(table) -> None:
    grid = table._tbl.tblGrid
    for item in list(grid):
        grid.remove(item)
    for width in PUBLIC_COLUMN_WIDTHS:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, PUBLIC_COLUMN_WIDTHS[index])
            _set_cell_margin(cell)


def _fill_card_block(table, start: int) -> None:
    identity = table.cell(start, 0).merge(table.cell(start + CARD_ROW_COUNT - 1, 0))
    _set_identity(identity)
    for offset, (cn_label, en_label, unit) in enumerate(FIELD_ROWS):
        row = table.rows[start + offset]
        _set_cant_split(row)
        _set_paragraph(row.cells[1], ((cn_label, True, "1F4E78"), ("〔值〕", False, "1F2937")))
        _set_paragraph(row.cells[2], ((unit, False, "A23B3B"),), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_paragraph(row.cells[3], ((en_label, True, "1F4E78"), (" [value]", False, "1F2937")))
        _set_paragraph(row.cells[4], ((unit, False, "A23B3B"),), align=WD_ALIGN_PARAGRAPH.CENTER)
        source_text = "来源\n〔文件/页码〕" if offset == 0 else "〔来源〕"
        _set_paragraph(row.cells[5], ((source_text, offset == 0, "4B5563"),), align=WD_ALIGN_PARAGRAPH.CENTER)
        if offset % 2:
            for cell in row.cells[1:]:
                _shade_cell(cell, "F7F9FB")
    for row in table.rows[start : start + CARD_ROW_COUNT]:
        seen: set[int] = set()
        for cell in row.cells:
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = row is not table.rows[start + CARD_ROW_COUNT - 1]


def create_public_pump_card_template(output_path: str | Path) -> Path:
    """Create the distributable pump-card template from an empty document."""
    output_path = Path(output_path)
    document = Document()
    document.core_properties.title = "通用泵参数卡片模板"
    document.core_properties.subject = "可分发的通用泵参数卡片"
    document.core_properties.author = "RFQ Translation Parameter Tool contributors"
    document.core_properties.keywords = "pump parameter card; synthetic; public template"
    document.core_properties.comments = "Created programmatically from a blank document."

    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(8)
    section.left_margin = Mm(8)
    section.right_margin = Mm(8)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(4)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    _format_run(title.add_run("通用泵参数卡片"), size=13, bold=True, color="17365D")

    table = document.add_table(rows=CARD_ROW_COUNT * TEMPLATE_CARD_COUNT, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    table.allow_autofit = False
    _set_grid(table)
    _fill_card_block(table, 0)
    _fill_card_block(table, CARD_ROW_COUNT)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Create the distributable pump-card DOCX template from a blank document.")
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    print(create_public_pump_card_template(args.output))
